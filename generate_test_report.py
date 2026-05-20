from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os
from datetime import datetime

doc = Document()

# ── Page margins ──────────────────────────────────────────
section = doc.sections[0]
section.left_margin   = Inches(0.8)
section.right_margin  = Inches(0.8)
section.top_margin    = Inches(0.8)
section.bottom_margin = Inches(0.8)

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def add_heading(doc, text, level=1, color='7C3AED'):
    p = doc.add_heading(text, level=level)
    run = p.runs[0] if p.runs else p.add_run(text)
    run.font.color.rgb = RGBColor.from_string(color)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return p

def add_badge(para, text, bg_hex, fg_hex='FFFFFF'):
    run = para.add_run(f'  {text}  ')
    run.font.bold = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor.from_string(fg_hex)
    return run

# ── Title Page ────────────────────────────────────────────
doc.add_paragraph()
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('QuizHub AI — Test Execution Report')
run.font.size = Pt(24)
run.font.bold = True
run.font.color.rgb = RGBColor(124, 58, 237)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = subtitle.add_run('Accenture Hackathon — Level 1 Submission')
run2.font.size = Pt(14)
run2.font.color.rgb = RGBColor(100, 100, 100)

date_p = doc.add_paragraph()
date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
date_p.add_run(f'Test Date: {datetime.now().strftime("%d %B %Y")}    |    Environment: http://localhost:8080')

doc.add_paragraph()

# ── Summary Table ─────────────────────────────────────────
add_heading(doc, 'Test Summary', level=1)

summary_table = doc.add_table(rows=2, cols=5)
summary_table.style = 'Table Grid'
headers = ['Total Scenarios', 'Executed', 'Passed', 'Failed', 'Coverage']
values  = ['242', '1', '1', '0', 'PPT Feature 1 — Login']

hdr_row = summary_table.rows[0]
val_row = summary_table.rows[1]
colors  = ['2D3748', '2563EB', '059669', 'DC2626', '7C3AED']

for i, (h, v, c) in enumerate(zip(headers, values, colors)):
    hc = hdr_row.cells[i]
    vc = val_row.cells[i]
    hc.text = h
    vc.text = v
    set_cell_bg(hc, c)
    hc.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
    hc.paragraphs[0].runs[0].font.bold = True
    hc.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    vc.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    vc.paragraphs[0].runs[0].font.bold = True

doc.add_paragraph()

# ── Test Case Section ─────────────────────────────────────
add_heading(doc, 'TC-01 — Login with Demo Admin Credentials', level=2, color='059669')

# Status badge row
p = doc.add_paragraph()
p.add_run('Status: ').font.bold = True
status_run = p.add_run('✅  PASS')
status_run.font.bold = True
status_run.font.color.rgb = RGBColor(5, 150, 105)
p.add_run('    |    Priority: ').font
p.add_run('HIGH').font.bold = True
p.add_run('    |    Feature Area: PPT — Role-Based Authentication')

doc.add_paragraph()

# Details table
details = doc.add_table(rows=6, cols=2)
details.style = 'Table Grid'
detail_data = [
    ('Test ID',         'TC-01'),
    ('Feature',         'Login — Demo Credentials (PPT Feature 1)'),
    ('Test Type',       'Positive Scenario'),
    ('Precondition',    'App running at localhost:8080. User divya.madhanasekar exists with role ADMIN, password Admin@123'),
    ('Test Steps',      '1. Open http://localhost:3000/login\n2. Click "Use" next to Admin 1 (divya.madhanasekar)\n3. Click "Sign In →"'),
    ('Expected Result', 'User is redirected to Dashboard (/). Role = ADMIN. Sidebar shows Admin-only items (Question Bank, User Management, etc.)'),
]
for row, (label, value) in zip(details.rows, detail_data):
    label_cell = row.cells[0]
    value_cell = row.cells[1]
    label_cell.text = label
    value_cell.text = value
    set_cell_bg(label_cell, 'EDE9FE')
    label_cell.paragraphs[0].runs[0].font.bold = True
    label_cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(124, 58, 237)
    # Set label col width
    label_cell.width = Inches(1.8)

doc.add_paragraph()

# Actual result
add_heading(doc, 'Actual Result', level=3, color='059669')
p = doc.add_paragraph()
p.add_run('✅  PASS — ').font.bold = True
p.add_run('User was redirected to Dashboard. Admin sidebar items visible. Welcome message shows "divya.madhanasekar". Recent activity table shows approved questions with IST-formatted timestamps.')

doc.add_paragraph()

# ── Screenshots ───────────────────────────────────────────
SHOTS_DIR = '/Users/veera.konjeti/Desktop/hack-n-stack/test-screenshots'

add_heading(doc, 'Screenshots — Desktop View (1440 × 900)', level=3, color='2563EB')

images = [
    ('TC01_login_desktop_before.png', 'DESKTOP — Login Page (before sign-in)\nShows: Enterprise ID pre-filled, password filled, demo credentials panel visible'),
    ('TC01_login_desktop_after.png',  'DESKTOP — Dashboard (after sign-in)\nShows: Admin dashboard, stat cards, recent activity, leaderboard, sidebar with admin items'),
]

for filename, caption in images:
    path = os.path.join(SHOTS_DIR, filename)
    if os.path.exists(path):
        doc.add_picture(path, width=Inches(6.0))
        cap = doc.add_paragraph(caption)
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.runs[0].font.size = Pt(9)
        cap.runs[0].font.color.rgb = RGBColor(100, 100, 100)
        cap.runs[0].font.italic = True
        doc.add_paragraph()

add_heading(doc, 'Screenshots — Mobile View (400 × 800)', level=3, color='D97706')

mobile_images = [
    ('TC01_login_mobile_before.png', 'MOBILE — Login Page (before sign-in)\nShows: Compact single-column layout, AI animation panel hidden, demo credentials visible'),
    ('TC01_login_mobile_after.png',  'MOBILE — Dashboard (after sign-in)\nShows: Hamburger menu replaces sidebar, cards stack vertically, responsive layout'),
]

for filename, caption in mobile_images:
    path = os.path.join(SHOTS_DIR, filename)
    if os.path.exists(path):
        doc.add_picture(path, width=Inches(3.0))
        cap = doc.add_paragraph(caption)
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.runs[0].font.size = Pt(9)
        cap.runs[0].font.color.rgb = RGBColor(100, 100, 100)
        cap.runs[0].font.italic = True
        doc.add_paragraph()

# ── Test Observations ─────────────────────────────────────
add_heading(doc, 'Observations & Notes', level=3, color='2D3748')
observations = [
    '✅  Demo credential autofill works correctly for both ADMIN and SME roles.',
    '✅  Role-based password logic: Admin → Admin@123 | SME → Sme@1234.',
    '✅  After login, JWT stored in localStorage (key: token).',
    '✅  Admin sidebar items (Question Bank, User Management, Master Data, Audit Log, Reviewer Metrics) visible.',
    '✅  Recent activity timestamps display in IST relative format (e.g. "2h ago", "Yesterday").',
    '✅  Mobile: sidebar collapses, hamburger menu appears at 400px width.',
    '✅  Mobile: login page shows only form panel (AI animation panel hidden).',
]
for obs in observations:
    p = doc.add_paragraph(obs, style='List Bullet')
    p.runs[0].font.size = Pt(10)

doc.add_paragraph()

# ── PPT Coverage Note ─────────────────────────────────────
add_heading(doc, 'PPT Requirements Coverage', level=2, color='7C3AED')
p = doc.add_paragraph()
p.add_run('This test scenario covers PPT Requirement: ').font.size = Pt(11)
p.add_run('"Role-based authentication with demo credentials for ADMIN and SME roles"').font.bold = True

coverage_table = doc.add_table(rows=4, cols=3)
coverage_table.style = 'Table Grid'
coverage_headers = ['PPT Requirement', 'Implemented', 'Tested']
coverage_data = [
    ('Role-based login (ADMIN / SME)', '✅ Yes', '✅ TC-01'),
    ('Demo credentials auto-fill', '✅ Yes', '✅ TC-01'),
    ('Redirect to dashboard after login', '✅ Yes', '✅ TC-01'),
]

hdr = coverage_table.rows[0]
for i, h in enumerate(coverage_headers):
    hdr.cells[i].text = h
    set_cell_bg(hdr.cells[i], '7C3AED')
    hdr.cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
    hdr.cells[i].paragraphs[0].runs[0].font.bold = True

for row_data, row in zip(coverage_data, coverage_table.rows[1:]):
    for i, val in enumerate(row_data):
        row.cells[i].text = val

doc.add_paragraph()
doc.add_paragraph('— End of TC-01 —').alignment = WD_ALIGN_PARAGRAPH.CENTER

# ── Save ──────────────────────────────────────────────────
output = '/Users/veera.konjeti/Desktop/hack-n-stack/QuizHub_Test_Report.docx'
doc.save(output)
print(f'Saved: {output}')
