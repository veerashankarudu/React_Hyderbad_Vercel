#!/usr/bin/env python3
"""
QuizHub AI — Comprehensive Feature Evidence Word Document
Valkey Hack-N-Stack 2026 | Team Valkey
"""

import os
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import parse_xml

SS_ORIG = Path("/Users/veera.konjeti/Desktop/hack-n-stack/screenshots")
SS_ANN  = Path("/Users/veera.konjeti/Desktop/hack-n-stack/screenshots/annotated")
SS = SS_ANN   # use annotated screenshots by default; fall back to original
OUT = "/Users/veera.konjeti/Desktop/hack-n-stack/QuizHub_AI_All_Features.docx"

PURPLE = RGBColor(0x7C, 0x3A, 0xED)
WHITE  = RGBColor(0xFF, 0xFF, 0xFF)
DARK   = RGBColor(0x1E, 0x1E, 0x2E)
GOLD   = RGBColor(0xF5, 0xA6, 0x23)
GREEN  = RGBColor(0x10, 0xB9, 0x81)

FEATURES = [
    # ─── AUTHENTICATION ───
    {
        "number": "01",
        "title": "Login & Authentication",
        "category": "Core Platform",
        "tagline": "JWT-based enterprise authentication with one-click demo access",
        "highlights": [
            "Enterprise ID + Password login (JWT tokens)",
            "Demo credential autofill panel — click any role to instantly fill fields",
            "Role-based redirect: Admin → Admin Dashboard, SME → SME Dashboard",
            "Forgot Password flow with email-based reset link",
            "New SME self-registration form",
            "Session persistence with secure localStorage token handling",
        ],
        "unique_angle": "Demo credentials panel is a MASSIVE hackathon differentiator — judges can try any role in 1 click without reading docs.",
        "desktop": "login_desktop.png",
        "mobile":  "login_mobile.png",
    },
    # ─── SME DASHBOARD ───
    {
        "number": "02",
        "title": "SME Dashboard",
        "category": "Core Platform",
        "tagline": "Personalized command center for Subject Matter Experts",
        "highlights": [
            "Live stats: Total MCQs, Approved, Under Review, Rejected, Draft counts",
            "MCQ pipeline donut chart — visual distribution by status",
            "Recent Activity timeline with status badges",
            "Platform Insights panel (quality score, SLA breached, approval rate)",
            "Top Reviewers leaderboard widget",
            "Greeting with time-of-day (Good morning/afternoon/evening) + hydration reminders",
        ],
        "unique_angle": "Hydration & wellness reminders embedded into a B2B enterprise tool — judges notice the human touch.",
        "desktop": "sme_dashboard_desktop.png",
        "mobile":  "sme_dashboard_mobile.png",
    },
    # ─── ADMIN DASHBOARD ───
    {
        "number": "03",
        "title": "Admin Dashboard",
        "category": "Admin Features",
        "tagline": "Platform-wide analytics and control center for admins",
        "highlights": [
            "Total MCQs breakdown across ALL users and stacks",
            "MCQs by Tech Stack bar chart (platform-wide)",
            "Performance Overview with approval/review/reject/draft split",
            "Platform Insights: Quality Score, SLA Breached count, Approval Rate",
            "Top Reviewers widget with review counts",
            "Draggable/reorderable dashboard widgets",
            "Customizable dashboard layout saved per user",
        ],
        "unique_angle": "Drag-and-drop customizable dashboard widgets — production-grade UX that most hackathon projects never implement.",
        "desktop": "admin_dashboard_desktop.png",
        "mobile":  "admin_dashboard_mobile.png",
    },
    # ─── MY QUESTIONS ───
    {
        "number": "04",
        "title": "My Questions — Status Tracking",
        "category": "MCQ Management",
        "tagline": "Full lifecycle tracking of every MCQ from draft to approval",
        "highlights": [
            "Status filter tabs with live counts: All / Draft / Ready for Review / Under Review / Approved / Rejected",
            "Search by question text",
            "Filter by Difficulty (Easy/Medium/Hard)",
            "Sortable columns (Date, Difficulty, Stack)",
            "One-click actions: View, Edit, Submit for Review, Delete",
            "Pagination with configurable page size",
        ],
        "unique_angle": "7-state MCQ lifecycle (Draft → Ready → Under Review → Approved/Rejected → Permanently Rejected) mirrors real enterprise content workflows.",
        "desktop": "my_questions_desktop.png",
        "mobile":  "my_questions_mobile.png",
    },
    # ─── CREATE MCQ ───
    {
        "number": "05",
        "title": "Create MCQ — Rich Question Form",
        "category": "MCQ Management",
        "tagline": "Guided MCQ creation with smart dropdowns and dual save options",
        "highlights": [
            "Tech Stack dropdown (admin-managed)",
            "Topic dropdown auto-filtered by selected Tech Stack",
            "Difficulty selector (Easy / Medium / Hard)",
            "4-option MCQ with correct answer radio",
            "Rich explanation field",
            "Tags/keywords input",
            "Save as Draft OR Save & Send for Review in one click",
            "Form validation with inline error messages",
        ],
        "unique_angle": "Topics are dynamically filtered by tech stack — no stale dropdowns, no user errors.",
        "desktop": "create_mcq_form_desktop.png",
        "mobile":  "create_mcq_form_mobile.png",
    },
    # ─── VIEW QUESTION DETAIL ───
    {
        "number": "06",
        "title": "View Question Detail & Comments",
        "category": "MCQ Management",
        "tagline": "Threaded review comments with full question audit trail",
        "highlights": [
            "Full question display: stem, 4 options, correct answer highlighted, explanation",
            "Version history / audit trail",
            "Threaded review comment system",
            "Comment reply with @mention support",
            "Status badge with transition history",
            "Reviewer assignment display",
            "Submit / Withdraw from Review actions",
        ],
        "unique_angle": "GitHub-style threaded comments on MCQs — reviewers can have back-and-forth discussions on a question like a PR review.",
        "desktop": "view_question_detail_desktop.png",
        "mobile":  "view_question_detail_mobile.png",
    },
    # ─── KANBAN BOARD ───
    {
        "number": "07",
        "title": "Kanban Board",
        "category": "MCQ Management",
        "tagline": "Drag-and-drop Kanban view of entire MCQ pipeline",
        "highlights": [
            "5 swim lanes: Draft, Ready for Review, Under Review, Approved, Rejected",
            "Drag cards between lanes to change status",
            "Card shows: question snippet, tech stack chip, difficulty badge, date",
            "Filter by Tech Stack",
            "Animated card transitions",
            "Works on mobile with touch drag",
        ],
        "unique_angle": "Jira-style Kanban for content management — nobody expects this level of UX in a quiz platform hackathon.",
        "desktop": "kanban_board_desktop.png",
        "mobile":  "kanban_board_mobile.png",
    },
    # ─── BULK UPLOAD ───
    {
        "number": "08",
        "title": "Bulk Upload via CSV/Excel",
        "category": "MCQ Management",
        "tagline": "Mass import questions from CSV with AI-powered duplicate detection",
        "highlights": [
            "Drag-and-drop CSV/Excel file upload",
            "Column mapping interface",
            "Row-level validation with error highlighting",
            "AI semantic duplicate detection (flags near-identical questions)",
            "Preview table before committing import",
            "Batch import with success/failure report",
            "Download sample template CSV",
        ],
        "unique_angle": "AI semantic duplicate detection during bulk import — prevents polluting the question bank with near-duplicate questions.",
        "desktop": "bulk_upload_desktop.png",
        "mobile":  "bulk_upload_mobile.png",
    },
    # ─── AI STUDIO ───
    {
        "number": "09",
        "title": "AI Studio — AI-Powered MCQ Generation",
        "category": "AI Features",
        "tagline": "Generate production-ready MCQs from a topic in one click",
        "highlights": [
            "Input: Topic + Tech Stack + Difficulty + Count",
            "AI generates question stem, 4 distractors, correct answer, explanation",
            "Generated questions are editable before saving",
            "One-click save to Draft",
            "Batch generate (up to 10 questions at once)",
            "Uses Spring AI + configurable provider (OpenAI / Ollama)",
        ],
        "unique_angle": "AI-generated MCQs feed directly into the 7-state approval workflow — AI assists humans, humans stay in control.",
        "desktop": "ai_studio_desktop.png",
        "mobile":  "ai_studio_mobile.png",
    },
    # ─── SCREENSHOT MCQ ───
    {
        "number": "10",
        "title": "Screenshot to MCQ (Vision AI)",
        "category": "AI Features",
        "tagline": "Take a screenshot of any technical content — AI converts it to a structured MCQ",
        "highlights": [
            "Upload any image (slides, code screenshots, textbook pages)",
            "Ollama minicpm-v vision model extracts text via OCR",
            "Second AI pass structures the content into MCQ format",
            "Extracted question is pre-filled in the Create MCQ form",
            "Handles code snippets, diagrams, and text",
            "Completely local — no data leaves the machine (Ollama)",
        ],
        "unique_angle": "Fully LOCAL vision AI — no API costs, no data privacy issues. Train material → MCQ in 10 seconds.",
        "desktop": "ai_studio_desktop.png",
        "mobile":  "ai_studio_mobile.png",
    },
    # ─── PENDING REVIEWS ───
    {
        "number": "11",
        "title": "Pending Reviews — Reviewer Workflow",
        "category": "Review Workflow",
        "tagline": "Streamlined review queue with one-click approve/reject/request changes",
        "highlights": [
            "Queue of questions assigned to reviewer",
            "Inline review actions: Approve / Request Changes / Reject",
            "Add review comment before decision",
            "Filter by Tech Stack, Difficulty",
            "SLA countdown per question (days until breach)",
            "Bulk actions (approve/reject selected)",
        ],
        "unique_angle": "SLA countdown timer per question — reviewers can see which questions are about to breach the SLA threshold.",
        "desktop": "pending_reviews_desktop.png",
        "mobile":  "pending_reviews_mobile.png",
    },
    # ─── REVIEWER DASHBOARD ───
    {
        "number": "12",
        "title": "Reviewer Dashboard & SLA Metrics",
        "category": "Review Workflow",
        "tagline": "Personal performance metrics for reviewers with SLA breach tracking",
        "highlights": [
            "Total reviewed, Approved, Rejected, Pending counts",
            "Average review time",
            "SLA breach table: questions stuck beyond threshold (days stuck / SLA limit / since date)",
            "Review velocity chart (reviews per week)",
            "Compared against platform average",
        ],
        "unique_angle": "SLA breach table shows exactly which questions are overdue, for how long, and since when — enterprise-grade accountability.",
        "desktop": "reviewer_dashboard_desktop.png",
        "mobile":  "reviewer_dashboard_mobile.png",
    },
    # ─── ASSIGN REVIEWER ───
    {
        "number": "13",
        "title": "Assign Reviewer (Admin)",
        "category": "Review Workflow",
        "tagline": "Admin assigns questions to specific reviewers with one click",
        "highlights": [
            "Dropdown list of available reviewers (filtered by tech stack expertise)",
            "Batch assign: select multiple questions → assign to reviewer",
            "Reviewer assignment reflected in Pending Reviews queue",
            "Notification sent to assigned reviewer",
        ],
        "unique_angle": "Reviewer assignment respects tech stack expertise — questions go to people who actually know the subject.",
        "desktop": "assign_reviewer_dialog_desktop.png",
        "mobile":  "assign_reviewer_dialog_mobile.png",
    },
    # ─── QUESTION BANK ───
    {
        "number": "14",
        "title": "Question Bank Management (Admin)",
        "category": "Admin Features",
        "tagline": "Platform-wide question repository with full admin controls",
        "highlights": [
            "View ALL questions across all SMEs",
            "Filter by Tech Stack, Topic, Difficulty, Status, Creator",
            "Bulk status changes",
            "Export to CSV",
            "Permanently reject with reason",
            "Full audit trail per question",
        ],
        "unique_angle": "Admin sees everything — full cross-user, cross-stack question visibility with bulk operations.",
        "desktop": "admin_question_bank_desktop.png",
        "mobile":  "admin_question_bank_mobile.png",
    },
    # ─── ANALYTICS ───
    {
        "number": "15",
        "title": "Analytics Dashboard",
        "category": "Analytics & Insights",
        "tagline": "Platform-wide visual analytics for admins",
        "highlights": [
            "MCQ creation trend (line chart, by week/month)",
            "Tech stack distribution (pie/donut chart)",
            "Approval rate over time",
            "Reviewer performance comparison",
            "Question difficulty distribution bar chart",
            "Active vs inactive SME tracking",
        ],
        "unique_angle": "Time-series analytics shows platform health trends — not just snapshots but trends over time.",
        "desktop": "analytics_desktop.png",
        "mobile":  "analytics_mobile.png",
    },
    # ─── RULE BOOK ───
    {
        "number": "16",
        "title": "Rule Book — AI Quality Guidelines",
        "category": "AI Features",
        "tagline": "Admin-configurable AI quality rules that govern MCQ generation and review",
        "highlights": [
            "Define quality rules: e.g., 'No ambiguous answers', 'Exactly one correct option'",
            "Rules are fed into AI generation prompts",
            "Rules displayed to SMEs during question creation",
            "Admin can add/edit/delete rules per tech stack",
            "Rules versioned with activation timestamps",
        ],
        "unique_angle": "AI follows human-defined quality rules — not a black box. This is how responsible AI in enterprise actually works.",
        "desktop": "rulebook_desktop.png",
        "mobile":  "rulebook_mobile.png",
    },
    # ─── QUESTION TYPES ───
    {
        "number": "17",
        "title": "Custom Question Types",
        "category": "Admin Features",
        "tagline": "Beyond MCQ — admin-defined custom question type schemas",
        "highlights": [
            "Admin creates new question type schemas (JSON-based)",
            "Define fields: text, code, image, multi-select, etc.",
            "Question Type Creator UI (no-code schema builder)",
            "SMEs can create questions matching custom schemas",
            "Frontend dynamically renders form based on schema",
        ],
        "unique_angle": "Schema-driven question types — admin defines a new question type and the form auto-generates. No code deployment needed.",
        "desktop": "question_types_desktop.png",
        "mobile":  "question_types_mobile.png",
    },
    # ─── QUIZ BUILDER ───
    {
        "number": "18",
        "title": "Quiz Builder",
        "category": "Assessment Features",
        "tagline": "Build custom quizzes from the approved question bank",
        "highlights": [
            "Select questions manually or auto-populate by criteria",
            "Filter by Tech Stack, Topic, Difficulty",
            "Set time limit, max attempts, passing score",
            "Preview quiz before publishing",
            "Assign quiz to users or roles",
            "Schedule quiz with start/end date",
        ],
        "unique_angle": "Rule-based auto-populate — 'Give me 10 Hard Spring Boot questions' and the system picks them automatically.",
        "desktop": "quiz_builder_desktop.png",
        "mobile":  "quiz_builder_mobile.png",
    },
    # ─── LIVE QUIZ ───
    {
        "number": "19",
        "title": "Live Quiz — Real-Time Multiplayer Quiz",
        "category": "Assessment Features",
        "tagline": "Kahoot-style live quiz with real-time leaderboard and host controls",
        "highlights": [
            "Host creates a session with a room code",
            "Participants join via code on any device",
            "WebSocket-based real-time question sync",
            "Live leaderboard updates after each question",
            "Host can pause, skip, end session",
            "Post-quiz results with participant rankings",
            "Session replay/detail view",
        ],
        "unique_angle": "Real-time multiplayer quiz with WebSockets — turns corporate training into a competitive, engaging experience.",
        "desktop": "live_quiz_desktop.png",
        "mobile":  "live_quiz_mobile.png",
    },
    # ─── LEADERBOARD ───
    {
        "number": "20",
        "title": "Leaderboard",
        "category": "Gamification",
        "tagline": "Platform-wide ranking of top performers with achievement badges",
        "highlights": [
            "Top SMEs ranked by questions created and approved",
            "Top reviewers ranked by reviews completed",
            "Filterable by time period (week/month/all time)",
            "Gold/Silver/Bronze medals for top 3",
            "Profile cards with contribution stats",
        ],
        "unique_angle": "Separate leaderboards for creators AND reviewers — gamifies both sides of the content pipeline.",
        "desktop": "leaderboard_desktop.png",
        "mobile":  "leaderboard_mobile.png",
    },
    # ─── SMART INTERVIEW KIT ───
    {
        "number": "21",
        "title": "Smart Interview Kit",
        "category": "AI Features",
        "tagline": "AI-curated interview prep kit from approved question bank",
        "highlights": [
            "Select tech stack + role + experience level",
            "AI curates a personalized set of interview questions",
            "Mix of difficulty levels based on role seniority",
            "Download as PDF interview prep guide",
            "Questions sourced only from APPROVED bank (quality guaranteed)",
        ],
        "unique_angle": "Interview prep from your own organization's question bank — using questions your own SMEs wrote, not generic internet content.",
        "desktop": "smart_interview_kit_desktop.png",
        "mobile":  "smart_interview_kit_mobile.png",
    },
    # ─── MASTER DATA ───
    {
        "number": "22",
        "title": "Master Data — Tech Stacks & Topics",
        "category": "Admin Features",
        "tagline": "Admin-managed taxonomy of tech stacks, topics, and assignments",
        "highlights": [
            "Create/edit/delete Tech Stacks (e.g., Spring Boot, Core Java)",
            "Create/edit/delete Topics per Tech Stack",
            "Assign Tech Stacks to SMEs (each SME only sees their stacks)",
            "Activate/deactivate stacks",
            "Bulk operations",
        ],
        "unique_angle": "SME scope isolation — each SME only sees and creates questions for their assigned tech stacks. Enterprise access control built in.",
        "desktop": "admin_tech_stacks_desktop.png",
        "mobile":  "admin_tech_stacks_mobile.png",
    },
    # ─── USERS ───
    {
        "number": "23",
        "title": "User Management",
        "category": "Admin Features",
        "tagline": "Full user CRUD with role management and tech stack assignment",
        "highlights": [
            "Create/edit/deactivate users",
            "Assign roles: Admin / SME",
            "Assign tech stacks to SME users",
            "View user activity stats",
            "Reset password (admin-triggered)",
        ],
        "unique_angle": "Tech stack assignment per user drives the entire content visibility model — one setting controls what each SME sees.",
        "desktop": "admin_users_desktop.png",
        "mobile":  "admin_users_mobile.png",
    },
    # ─── NOTIFICATIONS ───
    {
        "number": "24",
        "title": "In-App Notifications & Inbox",
        "category": "Core Platform",
        "tagline": "Real-time notifications with inbox for all workflow events",
        "highlights": [
            "Bell icon with unread count badge",
            "Notification types: Question approved, rejected, under review, comment added, reviewer assigned",
            "Mark as read / mark all as read",
            "Inbox page with full notification history",
            "Email notifications (configurable)",
        ],
        "unique_angle": "Every workflow state change triggers a notification — complete audit trail in the inbox.",
        "desktop": "notifications_panel_desktop.png",
        "mobile":  "notifications_panel_desktop.png",
    },
    # ─── CHATBOT ───
    {
        "number": "25",
        "title": "AI Collab Hub (Chatbot)",
        "category": "AI Features",
        "tagline": "Context-aware AI assistant with RAG over the question bank",
        "highlights": [
            "Floating chatbot available on every page",
            "RAG (Retrieval Augmented Generation) — answers sourced from your own question bank",
            "Ask 'What Spring Boot questions do we have on JPA?'",
            "Conversational multi-turn memory",
            "AI cites which questions it used to answer",
            "Powered by Spring AI + vector embeddings",
        ],
        "unique_angle": "RAG chatbot over your OWN question bank — AI answers are grounded in your organization's content, not hallucinations.",
        "desktop": "chatbot_desktop.png",
        "mobile":  "chatbot_desktop.png",
    },
    # ─── DARK MODE + I18N ───
    {
        "number": "26",
        "title": "Dark Mode + 7-Language i18n",
        "category": "Core Platform",
        "tagline": "Full dark/light theme toggle + 7 language support",
        "highlights": [
            "Dark / Light mode toggle with smooth CSS transitions",
            "7 languages: English, Hindi, Spanish, French, German, Japanese, Chinese",
            "Language selector in navbar",
            "All UI labels, buttons, messages translated",
            "Locale persisted in localStorage",
        ],
        "unique_angle": "7 languages in a Java/Spring hackathon project is almost unheard of — signals global enterprise readiness.",
        "desktop": "dark_mode_desktop.png",
        "mobile":  "dark_mode_mobile.png",
    },
    # ─── AUDIT LOG ───
    {
        "number": "27",
        "title": "Audit Log",
        "category": "Admin Features",
        "tagline": "Immutable audit trail of every admin action on the platform",
        "highlights": [
            "Log of all admin actions: user changes, status overrides, bulk operations",
            "Timestamp, actor, action, target entity",
            "Filter by date range, action type, user",
            "Export to CSV for compliance",
            "Read-only (cannot be deleted)",
        ],
        "unique_angle": "Immutable audit log for compliance — this is what separates a demo from a production-ready enterprise system.",
        "desktop": "admin_dashboard_desktop.png",
        "mobile":  "admin_dashboard_mobile.png",
    },
]


def set_cell_bg(cell, hex_color):
    """Set cell background color."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:val="clear" w:color="auto" w:fill="{hex_color}"/>')
    tcPr.append(shd)


def nsdecls(*prefixes):
    return ' '.join(f'xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"')


def add_cover(doc):
    doc.add_paragraph()
    doc.add_paragraph()

    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("QuizHub AI")
    run.font.size = Pt(42)
    run.font.bold = True
    run.font.color.rgb = PURPLE

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Complete Feature Evidence Document")
    run.font.size = Pt(20)
    run.font.bold = True
    run.font.color.rgb = DARK

    doc.add_paragraph()
    tagline = doc.add_paragraph()
    tagline.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = tagline.add_run("Valkey Hack-N-Stack 2026  ·  Team Valkey  ·  By VEERA & TEJA")
    run.font.size = Pt(13)
    run.font.italic = True
    run.font.color.rgb = RGBColor(0x60, 0x60, 0x80)

    doc.add_paragraph()

    # Stats table
    table = doc.add_table(rows=1, cols=4)
    stats = [
        ("27", "Features"),
        ("2,029", "Tests"),
        ("92.5%", "Coverage"),
        ("235+", "Questions"),
    ]
    for i, (num, label) in enumerate(stats):
        cell = table.cell(0, i)
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r1 = para.add_run(num + "\n")
        r1.font.size = Pt(22)
        r1.font.bold = True
        r1.font.color.rgb = PURPLE
        r2 = para.add_run(label)
        r2.font.size = Pt(10)
        r2.font.color.rgb = RGBColor(0x60, 0x60, 0x80)

    doc.add_page_break()


def add_toc(doc):
    h = doc.add_heading("Table of Contents", level=1)
    h.runs[0].font.color.rgb = PURPLE

    categories = {}
    for f in FEATURES:
        cat = f["category"]
        if cat not in categories:
            categories[cat] = []
        categories[cat].append(f)

    for cat, features in categories.items():
        cat_para = doc.add_paragraph()
        run = cat_para.add_run(f"  {cat}")
        run.font.bold = True
        run.font.size = Pt(12)
        run.font.color.rgb = DARK

        for f in features:
            para = doc.add_paragraph()
            para.paragraph_format.left_indent = Cm(1)
            run = para.add_run(f"  {f['number']}.  {f['title']}")
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0x40, 0x40, 0x60)

    doc.add_page_break()


def add_feature(doc, f):
    # ── Feature header ──
    header = doc.add_paragraph()
    r1 = header.add_run(f"  Feature {f['number']}  ")
    r1.font.size = Pt(11)
    r1.font.bold = True
    r1.font.color.rgb = WHITE

    r2 = header.add_run(f"  {f['category'].upper()}  ")
    r2.font.size = Pt(9)
    r2.font.color.rgb = GOLD

    # ── Title ──
    title = doc.add_heading(f['title'], level=2)
    title.runs[0].font.color.rgb = PURPLE
    title.runs[0].font.size = Pt(18)

    # ── Tagline ──
    tagline = doc.add_paragraph()
    run = tagline.add_run(f['tagline'])
    run.font.size = Pt(12)
    run.font.italic = True
    run.font.color.rgb = RGBColor(0x50, 0x50, 0x70)

    doc.add_paragraph()

    # ── Highlights ──
    h = doc.add_paragraph()
    run = h.add_run("What it does:")
    run.font.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = DARK

    for point in f["highlights"]:
        bullet = doc.add_paragraph(style='List Bullet')
        bullet.paragraph_format.left_indent = Cm(0.5)
        run = bullet.add_run(point)
        run.font.size = Pt(10)

    doc.add_paragraph()

    # ── Unique angle ──
    unique = doc.add_paragraph()
    run = unique.add_run("⭐ Why this stands out:  ")
    run.font.bold = True
    run.font.color.rgb = GOLD
    run2 = unique.add_run(f["unique_angle"])
    run2.font.size = Pt(10)
    run2.font.italic = True

    doc.add_paragraph()

    # ── Screenshots ──
    h2 = doc.add_paragraph()
    run = h2.add_run("Screenshots:")
    run.font.bold = True
    run.font.size = Pt(11)

    # Side by side in a 2-col table
    table = doc.add_table(rows=1, cols=2)
    table.columns[0].width = Inches(3.2)
    table.columns[1].width = Inches(2.8)

    d_file = f.get("desktop", "")
    m_file = f.get("mobile",  "")
    desktop_path = SS_ANN / d_file if (SS_ANN / d_file).exists() else SS_ORIG / d_file
    mobile_path  = SS_ORIG / m_file  # mobile always from originals (no annotated mobile)

    cell_d = table.cell(0, 0)
    cell_d.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    lbl = cell_d.paragraphs[0].add_run("Desktop View\n")
    lbl.font.size = Pt(9)
    lbl.font.bold = True
    lbl.font.color.rgb = RGBColor(0x60, 0x60, 0x80)
    if desktop_path.exists():
        try:
            cell_d.paragraphs[0].add_run().add_picture(str(desktop_path), width=Inches(3.0))
        except Exception as e:
            cell_d.paragraphs[0].add_run(f"[{f['desktop']}]")

    cell_m = table.cell(0, 1)
    cell_m.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    lbl2 = cell_m.paragraphs[0].add_run("Mobile View\n")
    lbl2.font.size = Pt(9)
    lbl2.font.bold = True
    lbl2.font.color.rgb = RGBColor(0x60, 0x60, 0x80)
    if mobile_path.exists():
        try:
            cell_m.paragraphs[0].add_run().add_picture(str(mobile_path), width=Inches(2.6))
        except Exception as e:
            cell_m.paragraphs[0].add_run(f"[{f['mobile']}]")

    doc.add_paragraph()
    doc.add_paragraph("─" * 80)
    doc.add_page_break()


def add_summary(doc):
    h = doc.add_heading("Feature Summary", level=1)
    h.runs[0].font.color.rgb = PURPLE

    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    for cell, text in zip(hdr, ["Category", "Features", "Count"]):
        p = cell.paragraphs[0]
        r = p.add_run(text)
        r.font.bold = True
        r.font.color.rgb = WHITE

    categories = {}
    for f in FEATURES:
        cat = f["category"]
        if cat not in categories:
            categories[cat] = []
        categories[cat].append(f["title"])

    for cat, titles in categories.items():
        row = table.add_row().cells
        row[0].text = cat
        row[1].text = "\n".join(f"• {t}" for t in titles)
        row[2].text = str(len(titles))
        row[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    # Tech stack
    h2 = doc.add_heading("Tech Stack", level=2)
    h2.runs[0].font.color.rgb = PURPLE

    tech = [
        ("Backend", "Spring Boot 3.2.5 · Java 17 · Spring AI 1.0.0 · Spring Security · JPA/Hibernate · MySQL 8"),
        ("Frontend", "React 19 · React Router v6 · Recharts · i18next · CSS Modules"),
        ("AI/ML", "Spring AI · OpenAI GPT-4o-mini · Ollama (qwen2.5:3b + minicpm-v) · Vector Embeddings · RAG"),
        ("Testing", "JUnit 5 · Mockito · Jest · React Testing Library · H2 in-memory DB · JaCoCo"),
        ("DevOps", "Maven · npm · H2 (test) · MySQL (prod) · GitHub"),
        ("Coverage", "92.5% backend (JaCoCo) · 80.37% frontend (Jest) · 2,029 automated tests"),
    ]

    for label, detail in tech:
        p = doc.add_paragraph()
        r1 = p.add_run(f"{label}: ")
        r1.font.bold = True
        r1.font.color.rgb = PURPLE
        r2 = p.add_run(detail)
        r2.font.size = Pt(10)


def main():
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin    = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin   = Cm(2.0)
        section.right_margin  = Cm(2.0)

    # Default font
    doc.styles['Normal'].font.name = 'Calibri'
    doc.styles['Normal'].font.size = Pt(10)

    add_cover(doc)
    add_toc(doc)

    for f in FEATURES:
        print(f"  Adding feature {f['number']}: {f['title']}")
        add_feature(doc, f)

    add_summary(doc)

    doc.save(OUT)
    print(f"\n✅  Saved: {OUT}")
    print(f"   Features: {len(FEATURES)}")


if __name__ == "__main__":
    main()
