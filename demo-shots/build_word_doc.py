#!/usr/bin/env python3
"""
QuizHub AI - Create comprehensive Word document from screenshots.
Each feature shows: Description + Desktop view + Mobile view on same page.
"""

import os
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

SCREENSHOT_DIR = Path("/Users/veera.konjeti/Desktop/hack-n-stack/screenshots")
OUTPUT_PATH = "/Users/veera.konjeti/Desktop/hack-n-stack/demo-shots/QuizHub_AI_Feature_Evidence.docx"

# Feature sections with descriptions and screenshot mappings
FEATURES = [
    {
        "section": "1. Authentication & Login",
        "description": "JWT-based role-based authentication system. Supports SME and Admin roles with separate dashboards. "
                       "Features include: Enterprise ID login, demo credential autofill, password reset, account registration, "
                       "and session management with secure token storage.",
        "what_it_does": "Users enter Enterprise ID and password → JWT token generated → Role-based redirect to appropriate dashboard",
        "screenshots": [
            ("login_desktop.png", "login_mobile.png"),
        ]
    },
    {
        "section": "2. SME Dashboard",
        "description": "Personalized dashboard for Subject Matter Experts showing quick stats (questions created, pending reviews, "
                       "approved count), recent activity timeline, and quick action buttons.",
        "what_it_does": "After login, SME sees their personalized overview: total questions, status breakdown, and quick navigation to all features",
        "screenshots": [
            ("sme_dashboard_desktop.png", "sme_dashboard_mobile.png"),
        ]
    },
    {
        "section": "3. My Questions — Status Tracking & Filters",
        "description": "Paginated table showing all questions created by the logged-in SME. Features status filter tabs with counts "
                       "(All/Draft/Ready for Review/Under Review/Approved/Rejected), search, difficulty filter, sortable columns, "
                       "and action buttons (View/Edit/Submit/Delete).",
        "what_it_does": "SME can track all their questions, filter by status, search, sort, and take actions on each question",
        "screenshots": [
            ("my_questions_desktop.png", "my_questions_mobile.png"),
            ("status_tab_draft_desktop.png", None),
            ("status_tab_ready_for_review_desktop.png", None),
            ("status_tab_approved_desktop.png", None),
            ("status_tab_rejected_desktop.png", None),
        ]
    },
    {
        "section": "4. Create MCQ — Add Single Question",
        "description": "Form to create a new MCQ with fields: Technology Stack (dropdown), Topic (dropdown filtered by tech stack), "
                       "Difficulty (Easy/Medium/Hard), Question Stem, 4 Options (A/B/C/D), Correct Answer selection, and Explanation. "
                       "Two save options: 'Save as Draft' and 'Save & Send for Review'.",
        "what_it_does": "SME fills the form → clicks Save → MCQ saved as Draft | clicks 'Save & Send for Review' → status becomes 'Ready for Review'",
        "screenshots": [
            ("create_mcq_form_desktop.png", "create_mcq_form_mobile.png"),
        ]
    },
    {
        "section": "5. View Question Detail",
        "description": "Full question detail view showing: question stem, all options with correct answer highlighted, "
                       "explanation, metadata (tech stack, topic, difficulty, status, creator, reviewer), version history, "
                       "and comment thread.",
        "what_it_does": "Clicking 'View' on any question shows complete MCQ details including content rendering for all 17 question types",
        "screenshots": [
            ("view_question_detail_desktop.png", "view_question_detail_mobile.png"),
        ]
    },
    {
        "section": "6. Edit MCQ — Update & Resubmit",
        "description": "Pre-filled form for editing Draft or Rejected questions. If Rejected, shows reviewer comments. "
                       "Two actions: 'Save' (keeps current status) and 'Save & Send for Review' (changes status to Ready for Review). "
                       "Includes AI Duplicate Check button to verify uniqueness before submission.",
        "what_it_does": "SME edits question fields → Save (stays Draft) or Save & Send (becomes Ready for Review). "
                       "For Rejected MCQs, reviewer's comments are visible to guide corrections.",
        "screenshots": [
            ("edit_question_desktop.png", "edit_question_mobile.png"),
        ]
    },
    {
        "section": "7. Bulk Upload — CSV/Excel Upload",
        "description": "Upload multiple MCQs at once via CSV/Excel file. Features: Download template button, "
                       "drag-and-drop file upload, field validation (tech stack, topic, difficulty, stem, options, correct answer), "
                       "error report showing which rows failed and why. Successfully validated MCQs are saved as Draft.",
        "what_it_does": "SME downloads template → fills in MCQs → uploads file → system validates each row → valid MCQs saved as Draft, errors reported",
        "screenshots": [
            ("bulk_upload_desktop.png", "bulk_upload_mobile.png"),
        ]
    },
    {
        "section": "8. Kanban Board — Visual Lifecycle",
        "description": "Visual drag-and-drop board showing MCQ lifecycle states as columns: Draft → Ready for Review → "
                       "Under Review → Approved/Rejected. Cards show question preview with metadata.",
        "what_it_does": "Visual representation of the entire MCQ lifecycle. Questions move across columns as they progress through the review workflow.",
        "screenshots": [
            ("kanban_board_desktop.png", "kanban_board_mobile.png"),
        ]
    },
    {
        "section": "9. My Pending Reviews",
        "description": "Shows MCQs assigned to the SME for review with 'Under Review' status. "
                       "SME can view full question, provide feedback comments, and Approve or Reject with mandatory comments for rejection.",
        "what_it_does": "Reviewer SME sees assigned questions → clicks 'View Full Question' → types feedback → clicks Approve/Reject",
        "screenshots": [
            ("pending_reviews_desktop.png", "pending_reviews_mobile.png"),
        ]
    },
    {
        "section": "10. Admin Dashboard",
        "description": "System-wide admin overview with total questions across all creators, status distribution, "
                       "user activity metrics, and system health indicators.",
        "what_it_does": "Admin sees platform-wide statistics: total questions, approval rates, active SMEs, and recent system activity",
        "screenshots": [
            ("admin_dashboard_desktop.png", "admin_dashboard_mobile.png"),
        ]
    },
    {
        "section": "11. Admin Question Bank Management",
        "description": "All MCQs from all creators in one paginated table with creator enterprise ID visible. "
                       "Admin can search, filter by tech stack/status/difficulty, use Semantic Search, Export to Excel, "
                       "and take actions: View, Edit (any status), Assign Reviewer, Delete.",
        "what_it_does": "Admin sees ALL questions regardless of creator → can filter/search → assign reviewers → edit any MCQ",
        "screenshots": [
            ("admin_question_bank_desktop.png", "admin_question_bank_mobile.png"),
            ("assign_reviewer_menu_desktop.png", None),
        ]
    },
    {
        "section": "12. Admin Assign Reviewer",
        "description": "Dialog showing: Question preview (stem, tech stack, topic, creator). "
                       "Reviewer dropdown filtered by tech stack mapping — only SMEs with matching tech stack shown, "
                       "creator excluded to prevent self-review. On assign, status changes to 'Under Review'.",
        "what_it_does": "Admin clicks Assign → dialog shows eligible reviewers (tech stack matched, creator excluded) → selects reviewer → status becomes 'Under Review'",
        "screenshots": [
            ("assign_reviewer_dialog_desktop.png", "assign_reviewer_dialog_mobile.png"),
        ]
    },
    {
        "section": "13. Admin User & Master Data Management",
        "description": "Manage users (SMEs/Admins) and master data. Users page shows all registered users with role and tech stack assignments. "
                       "Tech Stacks page manages the 6 technology stacks and their associated topics.",
        "what_it_does": "Admin manages: User roles, tech stack assignments, topic lists per tech stack. "
                       "Master data (6 tech stacks, 40+ topics) is pre-seeded on startup.",
        "screenshots": [
            ("admin_users_desktop.png", "admin_users_mobile.png"),
            ("admin_tech_stacks_desktop.png", "admin_tech_stacks_mobile.png"),
        ]
    },
    {
        "section": "14. AI Studio — MCQ Generation (Level 2)",
        "description": "AI-powered MCQ generation using Spring AI + OpenAI GPT-4o-mini. Form fields: "
                       "Technology Stack, Topic, Difficulty, Number of Questions, Question Type (17 types supported). "
                       "Auto-duplicate detection during generation — if similarity ≥30%, question is replaced with a new one.",
        "what_it_does": "SME selects tech stack + topic + difficulty + count → clicks Generate → AI creates MCQs as Draft → "
                       "auto-checks for duplicates (≥30% similarity triggers replacement)",
        "screenshots": [
            ("ai_studio_desktop.png", "ai_studio_mobile.png"),
        ]
    },
    {
        "section": "15. Analytics & Reporting",
        "description": "Comprehensive analytics with charts: questions created over time, status distribution pie chart, "
                       "difficulty breakdown, tech stack coverage, approval rates, and reviewer performance metrics.",
        "what_it_does": "Visual insights into question bank health: creation trends, approval rates, coverage gaps, and team productivity",
        "screenshots": [
            ("analytics_desktop.png", "analytics_mobile.png"),
            ("admin_reports_desktop.png", "admin_reports_mobile.png"),
            ("reviewer_dashboard_desktop.png", "reviewer_dashboard_mobile.png"),
        ]
    },
    {
        "section": "16. Assessments & Live Quiz",
        "description": "Build assessments from approved questions. Create quiz sessions with selected questions, "
                       "time limits, and scoring rules. Live Quiz enables real-time quiz sessions with participants.",
        "what_it_does": "Build assessment → select questions → set rules → launch live quiz → participants join → real-time scoring & leaderboard",
        "screenshots": [
            ("quiz_builder_desktop.png", "quiz_builder_mobile.png"),
            ("live_quiz_desktop.png", "live_quiz_mobile.png"),
        ]
    },
    {
        "section": "17. Smart Interview Kit",
        "description": "AI-generated interview preparation material based on tech stack and topics. "
                       "Provides curated question sets for interview preparation.",
        "what_it_does": "Select tech stack & topics → AI generates targeted interview prep questions and study material",
        "screenshots": [
            ("smart_interview_kit_desktop.png", "smart_interview_kit_mobile.png"),
        ]
    },
    {
        "section": "18. Question Types — 17 Formats",
        "description": "Supports 17 question types beyond standard MCQ: SINGLE, MULTI, DRAG_ORDER, MATCH_PAIRS, "
                       "FILL_BLANK, PREDICT_OUTPUT, DEBUG_CODE, RIDDLE, CODE_REARRANGE, SQL_BUILDER, "
                       "ARCH_LAYERS, CODE_REVIEW, PIPELINE_BUILD, FLOWCHART, DEVOPS_PIPE, SECURE_CODE, CODE_OUTPUT.",
        "what_it_does": "Each type has custom rendering: code blocks, drag-drop, pair matching, fill-in blanks, flowcharts, etc.",
        "screenshots": [
            ("question_types_desktop.png", "question_types_mobile.png"),
        ]
    },
    {
        "section": "19. Dark Mode & Internationalization (i18n)",
        "description": "Full dark/light theme toggle with CSS custom properties. Internationalization supports 7 languages: "
                       "English, German, French, Spanish, Hindi, Japanese, Korean. All UI text is translated.",
        "what_it_does": "Toggle theme button switches entire UI. Language selector changes all text to selected language instantly.",
        "screenshots": [
            ("dark_mode_desktop.png", "dark_mode_mobile.png"),
            ("language_selector_desktop.png", None),
        ]
    },
    {
        "section": "20. Notifications, Chatbot & Leaderboard",
        "description": "Real-time notification bell with unread count. AI Chatbot for quick help. "
                       "Leaderboard ranking top contributors by questions created and approved.",
        "what_it_does": "Notifications alert on status changes. Chatbot answers questions about MCQ creation. "
                       "Leaderboard gamifies contribution.",
        "screenshots": [
            ("notifications_panel_desktop.png", None),
            ("chatbot_desktop.png", None),
            ("leaderboard_desktop.png", "leaderboard_mobile.png"),
        ]
    },
    {
        "section": "21. Rule Book",
        "description": "Comprehensive guidelines for MCQ creation: formatting rules, quality standards, "
                       "difficulty calibration criteria, and best practices for each question type.",
        "what_it_does": "Reference document for SMEs to ensure consistent, high-quality question creation",
        "screenshots": [
            ("rulebook_desktop.png", "rulebook_mobile.png"),
        ]
    },
]


def add_shading(cell, color):
    """Add shading to a table cell."""
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)


def create_document():
    doc = Document()
    
    # Set narrow margins
    for section in doc.sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(1.5)
        section.right_margin = Cm(1.5)
    
    # ===== TITLE PAGE =====
    doc.add_paragraph()
    doc.add_paragraph()
    
    title = doc.add_heading("QuizHub AI", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Feature Evidence Document")
    run.font.size = Pt(24)
    run.font.color.rgb = RGBColor(0x66, 0x33, 0x99)
    
    doc.add_paragraph()
    
    team = doc.add_paragraph()
    team.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = team.add_run("Accenture Hack-N-Stack 2026 — Team Bumble Bee")
    run.font.size = Pt(14)
    run.bold = True
    
    doc.add_paragraph()
    
    tech = doc.add_paragraph()
    tech.alignment = WD_ALIGN_PARAGRAPH.CENTER
    lines = [
        "Spring Boot 3.2.5 (Java 17) • React 19 • MySQL 8 • Spring AI • OpenAI GPT-4o-mini",
        "",
        "2,029 Automated Tests (1,072 Backend + 957 Frontend) • 0 Failures",
        "Backend Coverage: 92.5% (JaCoCo) • Frontend Coverage: 80.37% (Jest)",
        "",
        "Every feature tested with Desktop (1440×900) & Mobile (390×844) views"
    ]
    for line in lines:
        if line:
            run = tech.add_run(line + "\n")
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
        else:
            tech.add_run("\n")
    
    doc.add_page_break()
    
    # ===== TABLE OF CONTENTS =====
    doc.add_heading("Table of Contents", level=1)
    for feature in FEATURES:
        p = doc.add_paragraph(feature["section"], style='List Number')
        p.paragraph_format.space_after = Pt(4)
    
    doc.add_page_break()
    
    # ===== FEATURE SECTIONS =====
    for idx, feature in enumerate(FEATURES):
        # Section heading
        doc.add_heading(feature["section"], level=1)
        
        # Description box
        p = doc.add_paragraph()
        run = p.add_run("What it does: ")
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0x00, 0x66, 0x99)
        run = p.add_run(feature["what_it_does"])
        run.font.size = Pt(10)
        
        p = doc.add_paragraph()
        run = p.add_run("Details: ")
        run.bold = True
        run.font.size = Pt(9)
        run = p.add_run(feature["description"])
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
        
        doc.add_paragraph()  # spacing
        
        # Screenshots
        for pair_idx, (desktop, mobile) in enumerate(feature["screenshots"]):
            desktop_path = SCREENSHOT_DIR / desktop if desktop else None
            mobile_path = SCREENSHOT_DIR / mobile if mobile else None
            
            # Desktop + Mobile side by side using a table
            if desktop and mobile and desktop_path.exists() and mobile_path.exists():
                # Create a 2-column table for side-by-side
                table = doc.add_table(rows=2, cols=2)
                table.alignment = WD_TABLE_ALIGNMENT.CENTER
                
                # Headers
                cell_d = table.rows[0].cells[0]
                p = cell_d.paragraphs[0]
                run = p.add_run("🖥️ Desktop View (1440×900)")
                run.bold = True
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0x00, 0x66, 0x99)
                
                cell_m = table.rows[0].cells[1]
                p = cell_m.paragraphs[0]
                run = p.add_run("📱 Mobile View (390×844)")
                run.bold = True
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0x00, 0x99, 0x66)
                
                # Images
                cell_d = table.rows[1].cells[0]
                p = cell_d.paragraphs[0]
                try:
                    run = p.add_run()
                    run.add_picture(str(desktop_path), width=Inches(4.2))
                except:
                    p.add_run(f"[{desktop}]")
                
                cell_m = table.rows[1].cells[1]
                p = cell_m.paragraphs[0]
                try:
                    run = p.add_run()
                    run.add_picture(str(mobile_path), width=Inches(2.0))
                except:
                    p.add_run(f"[{mobile}]")
                
                # Remove table borders for clean look
                for row in table.rows:
                    for cell in row.cells:
                        tc = cell._tc
                        tcPr = tc.get_or_add_tcPr()
                        tcBorders = parse_xml(
                            f'<w:tcBorders {nsdecls("w")}>'
                            '<w:top w:val="nil"/>'
                            '<w:left w:val="nil"/>'
                            '<w:bottom w:val="nil"/>'
                            '<w:right w:val="nil"/>'
                            '</w:tcBorders>'
                        )
                        tcPr.append(tcBorders)
                
            elif desktop and desktop_path and desktop_path.exists():
                # Desktop only
                p = doc.add_paragraph()
                run = p.add_run("🖥️ Desktop View (1440×900)")
                run.bold = True
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0x00, 0x66, 0x99)
                
                try:
                    doc.add_picture(str(desktop_path), width=Inches(6.5))
                    last_p = doc.paragraphs[-1]
                    last_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                except Exception as e:
                    doc.add_paragraph(f"[Screenshot: {desktop} - {e}]")
            
            doc.add_paragraph()  # spacing between pairs
        
        # Page break between sections (except last)
        if idx < len(FEATURES) - 1:
            doc.add_page_break()
    
    # ===== FINAL PAGE - TEST RESULTS =====
    doc.add_page_break()
    doc.add_heading("Automated Test Results", level=1)
    
    p = doc.add_paragraph()
    run = p.add_run("Backend Tests (JUnit 5 + Spring Boot Test)")
    run.bold = True
    p = doc.add_paragraph("• Total Tests: 1,072")
    p = doc.add_paragraph("• Failures: 0")
    p = doc.add_paragraph("• Coverage: 92.5% instruction coverage (JaCoCo)")
    p = doc.add_paragraph("• Database: H2 in-memory (no MySQL needed for tests)")
    
    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run("Frontend Tests (Jest + React Testing Library)")
    run.bold = True
    p = doc.add_paragraph("• Total Tests: 957")
    p = doc.add_paragraph("• Failures: 0")
    p = doc.add_paragraph("• Coverage: 80.37% statement coverage")
    p = doc.add_paragraph("• Mocking: Axios, React Router, i18n")
    
    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run("Total: 2,029 tests • 0 failures • Full CI/CD ready")
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x00, 0x88, 0x00)
    
    # Save
    doc.save(OUTPUT_PATH)
    print(f"✅ Document saved: {OUTPUT_PATH}")
    print(f"   Size: {os.path.getsize(OUTPUT_PATH) / 1024 / 1024:.1f} MB")


if __name__ == "__main__":
    create_document()
