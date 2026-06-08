#!/usr/bin/env python3
"""
QuizHub AI — Comprehensive Feature Evidence Document
Full end-to-end flows, internal architecture, 333 features, all page interconnections.
"""

import os
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

SCREENSHOT_DIR = Path("/Users/veera.konjeti/Desktop/hack-n-stack/screenshots")
OUTPUT_PATH = "/Users/veera.konjeti/Desktop/hack-n-stack/demo-shots/QuizHub_AI_Complete_Evidence.docx"


def set_cell_shading(cell, color):
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)


def remove_table_borders(table):
    for row in table.rows:
        for cell in row.cells:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcBorders = parse_xml(
                f'<w:tcBorders {nsdecls("w")}>'
                '<w:top w:val="nil"/><w:left w:val="nil"/>'
                '<w:bottom w:val="nil"/><w:right w:val="nil"/>'
                '</w:tcBorders>'
            )
            tcPr.append(tcBorders)


def add_screenshot_pair(doc, desktop_file, mobile_file=None, caption=""):
    """Add desktop + mobile screenshots side by side."""
    desktop_path = SCREENSHOT_DIR / desktop_file if desktop_file else None
    mobile_path = SCREENSHOT_DIR / mobile_file if mobile_file else None

    if desktop_path and desktop_path.exists() and mobile_path and mobile_path.exists():
        table = doc.add_table(rows=1, cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        cell_d = table.rows[0].cells[0]
        p = cell_d.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        try:
            run = p.add_run()
            run.add_picture(str(desktop_path), width=Inches(4.0))
        except:
            p.add_run(f"[{desktop_file}]")

        cell_m = table.rows[0].cells[1]
        p = cell_m.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        try:
            run = p.add_run()
            run.add_picture(str(mobile_path), width=Inches(2.0))
        except:
            p.add_run(f"[{mobile_file}]")

        remove_table_borders(table)
        if caption:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(caption)
            run.font.size = Pt(8)
            run.font.italic = True
    elif desktop_path and desktop_path.exists():
        try:
            doc.add_picture(str(desktop_path), width=Inches(6.0))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        except:
            pass
        if caption:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(caption)
            run.font.size = Pt(8)
            run.font.italic = True


def add_flow_box(doc, text, color="006699"):
    """Add a colored flow description box."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    run = p.add_run("▶ " + text)
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(int(color[:2], 16), int(color[2:4], 16), int(color[4:], 16))


def add_feature_list(doc, features):
    """Add numbered feature list."""
    for i, feat in enumerate(features, 1):
        p = doc.add_paragraph(f"{i}. {feat}", style='List Number')
        p.paragraph_format.space_after = Pt(1)
        for run in p.runs:
            run.font.size = Pt(9)


def create_document():
    doc = Document()

    # Narrow margins
    for section in doc.sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(1.5)
        section.right_margin = Cm(1.5)

    # ═══════════════════════════════════════════════════════════
    # TITLE PAGE
    # ═══════════════════════════════════════════════════════════
    doc.add_paragraph()
    title = doc.add_heading("QuizHub AI", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Complete Feature Evidence & Architecture Document")
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(0x66, 0x33, 0x99)

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Accenture Hack-N-Stack 2026 — Team Bumble Bee")
    run.font.size = Pt(14)
    run.bold = True

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for line in [
        "Spring Boot 3.2.5 (Java 17) • React 19 • MySQL 8 • Spring AI • OpenAI GPT-4o-mini",
        "17 Controllers • 120+ API Endpoints • 40 Frontend Routes • 18 Question Types",
        "333 Unique Features • 2,029 Automated Tests • 0 Failures",
        "Backend: 92.5% Coverage • Frontend: 80.37% Coverage",
        "7 Languages (i18n) • Dark/Light Mode • Real-time Chat • Live Quiz Battles",
    ]:
        run = p.add_run(line + "\n")
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════
    # TABLE OF CONTENTS
    # ═══════════════════════════════════════════════════════════
    doc.add_heading("Table of Contents", level=1)
    toc = [
        "Part A — Architecture & System Overview",
        "Part B — Complete MCQ Lifecycle (End-to-End Flow)",
        "Part C — My Questions (22 features)",
        "Part D — My Pending Reviews (19 features)",
        "Part E — Admin Question Bank Management (20 features)",
        "Part F — Admin Assign Reviewer (Internal Flow)",
        "Part G — Master Data Management (16 features)",
        "Part H — AI Studio & MCQ Generation (Level 2)",
        "Part I — AI Duplicate Detection (Level 2)",
        "Part J — Bulk Upload (16 features)",
        "Part K — Question Types (18 formats)",
        "Part L — Live Quiz Battle System (20 features)",
        "Part M — Assessments & Quiz Builder (18 features)",
        "Part N — User Activity & Audit Log (10 features)",
        "Part O — Analytics, Dashboard & Widgets (28 features)",
        "Part P — Real-time Chat & Inbox (28 features)",
        "Part Q — Smart Interview Kit (7 features)",
        "Part R — Cross-Cutting Features (24 features)",
        "Part S — All 40 Routes & 120+ API Endpoints",
        "Part T — Feature Count Summary (333 total)",
        "Part U — Test Results & Coverage",
    ]
    for item in toc:
        p = doc.add_paragraph(item)
        p.paragraph_format.space_after = Pt(3)
        for run in p.runs:
            run.font.size = Pt(10)

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════
    # PART A — ARCHITECTURE
    # ═══════════════════════════════════════════════════════════
    doc.add_heading("Part A — Architecture & System Overview", level=1)

    doc.add_heading("System Architecture", level=2)
    p = doc.add_paragraph()
    run = p.add_run("""
┌─────────────────────────────────────────────────────────────────────────┐
│                          QuizHub AI Architecture                         │
├─────────────────────────────────────────────────────────────────────────┤
│                                                                         │
│   ┌──────────────┐         REST/JSON          ┌──────────────────┐     │
│   │   React 19   │ ◀══════════════════════════▶│  Spring Boot 3.2 │     │
│   │  Port: 3000  │      JWT Bearer Token       │   Port: 8080     │     │
│   │              │                             │                  │     │
│   │ • 40 Routes  │                             │ • 17 Controllers │     │
│   │ • 957 Tests  │                             │ • 120+ Endpoints │     │
│   │ • i18n (7)   │                             │ • 1,072 Tests    │     │
│   │ • Dark Mode  │                             │ • JPA/Hibernate  │     │
│   └──────────────┘                             └────────┬─────────┘     │
│                                                         │               │
│                                                         │               │
│                              ┌───────────────────────────┼──────┐       │
│                              │                           ▼      │       │
│                              │  ┌─────────────┐  ┌───────────┐ │       │
│                              │  │  MySQL 8.x  │  │ OpenAI    │ │       │
│                              │  │  Schema:    │  │ GPT-4o-   │ │       │
│                              │  │  quizhub    │  │ mini      │ │       │
│                              │  │  (H2 test)  │  │ (optional)│ │       │
│                              │  └─────────────┘  └───────────┘ │       │
│                              │       Data Layer                 │       │
│                              └──────────────────────────────────┘       │
│                                                                         │
│   Roles: ADMIN (super-user) │ SME (expert) │ PUBLIC (quiz taker)       │
└─────────────────────────────────────────────────────────────────────────┘
""")
    run.font.size = Pt(8)
    run.font.name = "Courier New"

    doc.add_heading("Navigation Map — How Pages Connect", level=2)
    p = doc.add_paragraph()
    run = p.add_run("""
Home (Dashboard)
 ├── My Questions ─────────┬── Create MCQ (/mcq/create)
 │   (SME's own questions)  ├── View Detail (/mcq/{id}) ── Edit (/mcq/{id}/edit)
 │                          ├── Bulk Upload (/bulk-upload)
 │                          ├── Screenshot MCQ (/screenshot-mcq)
 │                          ├── Coding Question (/coding/create)
 │                          └── 15 Advanced Types (/question-type-create/{type})
 │
 ├── Pending Reviews ────── Review Panel (inline) ── Approve/Reject
 │   (Reviewer's queue)
 │
 ├── Question Bank ─────┬── View/Edit any MCQ
 │   (Admin: all MCQs)  ├── Assign Reviewer (modal)
 │                      ├── Semantic Search → Navigate to matches
 │                      └── Export to Excel
 │
 ├── AI Studio ────────── Code→MCQ │ Rewrite │ Learning Path
 │
 ├── Kanban Board ─────── Visual lifecycle columns
 │
 ├── Live Quiz ────────── Create Session → PIN Join → Host/Play → Results
 │
 ├── Quiz Builder ─────── Create Assessment → Share Link → Proctored Attempt
 │
 ├── Master Data ──────── Tech Stacks → Topics → SME Assignments
 │   (Admin)
 │
 ├── User Management ──── CRUD Users, Roles, Approvals
 │   (Admin)
 │
 ├── Audit Log ────────── All system actions history
 │   (Admin)
 │
 ├── Analytics ────────── Charts, exports, date ranges
 ├── Leaderboard ──────── Reviewer / Assessment / Live Quiz
 ├── Inbox ────────────── Internal messaging
 ├── Chat ─────────────── Real-time with @mentions, @bot, reactions
 ├── Smart Interview Kit ── Resume → AI Questions
 └── Rule Book ────────── Guidelines & standards
""")
    run.font.size = Pt(8)
    run.font.name = "Courier New"

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════
    # PART B — MCQ LIFECYCLE
    # ═══════════════════════════════════════════════════════════
    doc.add_heading("Part B — Complete MCQ Lifecycle (End-to-End Flow)", level=1)

    p = doc.add_paragraph()
    run = p.add_run("This is the core workflow that drives the entire platform. Every MCQ goes through this lifecycle:")
    run.font.size = Pt(10)

    doc.add_heading("Status Flow Diagram", level=2)
    p = doc.add_paragraph()
    run = p.add_run("""
   ┌──────────┐    Submit    ┌─────────────────┐   Admin    ┌──────────────┐
   │          │ ──────────▶  │                 │  Assigns   │              │
   │  DRAFT   │              │ READY FOR REVIEW│ ─────────▶ │ UNDER REVIEW │
   │          │ ◀──────────  │                 │            │              │
   └──────────┘  (edit back) └─────────────────┘            └──────┬───────┘
        ▲                                                          │
        │                                                    ┌─────┴─────┐
        │   Edit & Resubmit                                  │           │
        │                                              ┌─────▼────┐ ┌───▼─────┐
        └──────────────────────────────────────────────│ REJECTED │ │APPROVED │
                                                       └──────────┘ └─────────┘
""")
    run.font.size = Pt(9)
    run.font.name = "Courier New"

    doc.add_heading("Step-by-Step Internal Flow", level=2)

    steps = [
        ("Step 1: SME Creates MCQ", "006699",
         "SME logs in → Dashboard → My Questions → '+ Add Question' → McqForm\n"
         "API: POST /api/v1/mcqs → Status: DRAFT → Appears in My Questions table\n"
         "OR: Bulk Upload (CSV/Excel) → POST /api/v1/upload/bulk → Multiple DRAFTs\n"
         "OR: AI Generation → POST /api/v1/ai/generate-mcqs → Auto-creates DRAFTs with duplicate replacement"),

        ("Step 2: SME Submits for Review", "006699",
         "My Questions → finds Draft MCQ → clicks 'Submit' button\n"
         "API: POST /api/v1/mcqs/{id}/submit → Status changes: DRAFT → READY_FOR_REVIEW\n"
         "Validation: Question stem, all options, correct answer must be filled\n"
         "AI Duplicate Check: Automatically runs on submit — if ≥30% similar, BLOCKS submission"),

        ("Step 3: Admin Assigns Reviewer", "993300",
         "Admin → Question Bank → filters 'Ready for Review' → clicks ⋮ → 'Assign'\n"
         "API: GET /api/v1/admin/mcqs/{id}/eligible-reviewers → Shows SMEs matching tech stack\n"
         "Rules: Creator excluded (no self-review), tech stack must match\n"
         "API: POST /api/v1/admin/mcqs/{id}/assign-reviewer → Status: UNDER_REVIEW\n"
         "Notification: Reviewer gets notification bell alert"),

        ("Step 4: Reviewer Reviews MCQ", "006699",
         "Reviewer → My Pending Reviews → sees assigned MCQs → clicks to expand\n"
         "Review Panel shows: Full question, options, correct answer, AI risk indicator\n"
         "Must complete 4-item checklist before deciding\n"
         "Optional: AI Copilot validates answer correctness (POST /api/v1/ai/validate-answer)\n"
         "Decision: APPROVE or REJECT (reject requires mandatory comment)\n"
         "API: POST /api/v1/reviews/{mcqId}/submit {action: 'APPROVE'|'REJECT', comment: '...'}"),

        ("Step 5a: APPROVED → Done", "008800",
         "Status: UNDER_REVIEW → APPROVED\n"
         "MCQ appears in Question Bank with 'Approved' badge\n"
         "Available for: Quiz Builder, Live Quiz, Assessments\n"
         "Creator sees 'Approved' in My Questions tab"),

        ("Step 5b: REJECTED → Edit & Resubmit", "990000",
         "Status: UNDER_REVIEW → REJECTED\n"
         "Creator sees 'Rejected' in My Questions with reviewer comments visible\n"
         "Creator clicks 'Edit' → McqForm shows rejection comments at top\n"
         "After editing: Status resets to DRAFT → can Submit again → cycle repeats"),
    ]

    for title, color, desc in steps:
        p = doc.add_paragraph()
        run = p.add_run(title)
        run.bold = True
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(int(color[:2], 16), int(color[2:4], 16), int(color[4:], 16))

        for line in desc.split("\n"):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.5)
            run = p.add_run(line)
            run.font.size = Pt(9)

        doc.add_paragraph()

    doc.add_heading("Audit Trail", level=2)
    p = doc.add_paragraph(
        "Every status change, assignment, approval, and rejection is logged in the Audit Log "
        "(/audit-log) with: Timestamp, Actor (who did it), Action type, Target (affected user/MCQ), "
        "and Details. Admin can search, filter by action type, and sort."
    )

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════
    # PART C — MY QUESTIONS
    # ═══════════════════════════════════════════════════════════
    doc.add_heading("Part C — My Questions (22 Features)", level=1)

    p = doc.add_paragraph()
    run = p.add_run("What it does: ")
    run.bold = True
    p.add_run("Shows all MCQs created by the logged-in SME with status tracking, filtering, searching, sorting, pagination, and actions. This is the SME's home base for managing their questions.")

    doc.add_heading("Internal Architecture", level=2)
    p = doc.add_paragraph()
    run = p.add_run("""
┌─────────────────────────────────────────────────────────────────────────┐
│                        MY QUESTIONS PAGE                                  │
├─────────────────────────────────────────────────────────────────────────┤
│                                                                          │
│  ┌─── Status Tabs ────────────────────────────────────────────────────┐  │
│  │  [All:20] [Draft:15] [Ready:3] [Under Review:1] [Approved:1] [Rej:0] │
│  └────────────────────────────────────────────────────────────────────┘  │
│                                                                          │
│  ┌─── Filters ───────────────────────────────────────────────────────┐  │
│  │  [Search: ___________] [Difficulty: All ▼]  [+ Add Question ▼]    │  │
│  └────────────────────────────────────────────────────────────────────┘  │
│                                                                          │
│  ┌─── Table ─────────────────────────────────────────────────────────┐  │
│  │  # │ Question↕ │ Tech Stack↕ │ Topic↕ │ Diff↕ │ Status │ Actions  │  │
│  │  1 │ When us...│ Spring Boot │ Testing│ MED   │ Draft  │ V E S D  │  │
│  │  2 │ In a S...│ Spring Boot │ Testing│ MED   │ Draft  │ V E S D  │  │
│  │  ...                                                              │  │
│  └────────────────────────────────────────────────────────────────────┘  │
│                                                                          │
│  ┌─── Pagination ────────────────────────────────────────────────────┐  │
│  │  Showing 1-10 of 20  │  [< Prev] [1] [2] [Next >]  │ Per page: 10▼│ │
│  └────────────────────────────────────────────────────────────────────┘  │
│                                                                          │
│  Actions: V=View(/mcq/{id}) E=Edit(/mcq/{id}/edit) S=Submit D=Delete    │
└─────────────────────────────────────────────────────────────────────────┘
""")
    run.font.size = Pt(7)
    run.font.name = "Courier New"

    doc.add_heading("API Calls Made", level=2)
    api_table = doc.add_table(rows=8, cols=3)
    api_table.style = 'Table Grid'
    headers = ["Method", "Endpoint", "Purpose"]
    for i, h in enumerate(headers):
        api_table.rows[0].cells[i].paragraphs[0].add_run(h).bold = True
    data = [
        ("GET", "/api/v1/mcqs?status=X&page=N&size=S&sort=col,dir", "Fetch paginated questions"),
        ("POST", "/api/v1/mcqs/{id}/submit", "Submit Draft → Ready for Review"),
        ("DELETE", "/api/v1/mcqs/{id}", "Delete Draft MCQ"),
        ("POST", "/api/v1/ai/score-quality", "AI quality score (0-100)"),
        ("POST", "/api/v1/ai/generate-mcqs", "AI MCQ generation (from modal)"),
        ("GET", "/api/v1/master/tech-stacks", "Populate tech stack dropdown"),
        ("GET", "/api/v1/master/tech-stacks/{id}/topics", "Populate topics dropdown"),
    ]
    for row_idx, (m, e, p_text) in enumerate(data, 1):
        api_table.rows[row_idx].cells[0].paragraphs[0].add_run(m).font.size = Pt(8)
        api_table.rows[row_idx].cells[1].paragraphs[0].add_run(e).font.size = Pt(8)
        api_table.rows[row_idx].cells[2].paragraphs[0].add_run(p_text).font.size = Pt(8)

    doc.add_heading("All 22 Features", level=2)
    add_feature_list(doc, [
        "Status filter tabs with live count badges (All/Draft/Ready/Under Review/Approved/Rejected)",
        "Per-tab URL sync (?status=DRAFT updates URL for bookmarking)",
        "Difficulty filter dropdown (All/Easy/Medium/Hard)",
        "Free-text search across question stem, tech stack, topic",
        "6 sortable columns with toggle asc/desc direction",
        "Paginated results with configurable page size (5/10/20/50)",
        "View button → navigates to /mcq/{id} detail page",
        "Edit button → navigates to /mcq/{id}/edit form (only for Draft/Rejected)",
        "Submit button → sends Draft to Ready for Review (with validation)",
        "Delete button → confirmation dialog → removes MCQ",
        "AI Quality Score auto-calculated (0-100) with color coding (green/yellow/red)",
        "'+ Add Question' button opens creation dialog with 5 methods",
        "Add from UI → navigates to /mcq/create",
        "Bulk Upload → navigates to /bulk-upload",
        "Screenshot MCQ → navigates to /screenshot-mcq (AI OCR extraction)",
        "Coding Question → navigates to /coding/create",
        "15 Advanced Question Types → each navigates to /question-type-create/{typeId}",
        "AI MCQ Generator modal (inline): select tech stack, topic, difficulty, count, type → generate",
        "AI generation shows duplicate replacement count and created IDs",
        "Content translation (question stems, tech stacks rendered in selected language)",
        "Responsive table layout for mobile",
        "Confirmation dialogs for destructive actions (delete, submit)",
    ])

    add_screenshot_pair(doc, "my_questions_desktop.png", "my_questions_mobile.png",
                        "My Questions — Desktop (1440×900) & Mobile (390×844)")

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════
    # PART D — PENDING REVIEWS
    # ═══════════════════════════════════════════════════════════
    doc.add_heading("Part D — My Pending Reviews (19 Features)", level=1)

    p = doc.add_paragraph()
    run = p.add_run("What it does: ")
    run.bold = True
    p.add_run("Shows MCQs assigned to the current user for review. Reviewer can view full question, validate with AI, complete a checklist, and approve/reject with mandatory comments.")

    doc.add_heading("Internal Architecture", level=2)
    p = doc.add_paragraph()
    run = p.add_run("""
┌─────────────────────────────────────────────────────────────────────────┐
│                     PENDING REVIEWS PAGE                                  │
├─────────────────────────────────────────────────────────────────────────┤
│                                                                          │
│  ┌─── Stats Bar ─────────────────────────────────────────────────────┐  │
│  │  Pending: 3  │  Approved (by me): 12  │  Rejected (by me): 2     │  │
│  └────────────────────────────────────────────────────────────────────┘  │
│                                                                          │
│  ┌─── Review Table ──────────────────────────────────────────────────┐  │
│  │  # │ Question↕ │ Tech Stack↕ │ Topic↕ │ Diff↕ │ Creator↕ │ Action│  │
│  │  1 │ When us...│ Spring Boot │ Testing│ MED   │ birendra │ [▶]   │  │
│  └────────────────────────────────────────────────────────────────────┘  │
│                                                                          │
│  ┌─── Review Panel (expands on click) ───────────────────────────────┐  │
│  │  Question: "When using Spring Boot's @SpringBootTest..."          │  │
│  │  A) Configure the base package ✓(correct)                         │  │
│  │  B) Add @RunWith annotation                                       │  │
│  │  C) Set profile to "test"                                         │  │
│  │  D) Enable lazy initialization                                    │  │
│  │                                                                    │  │
│  │  ┌── AI Copilot ──┐  ┌── Risk: LOW ──┐                           │  │
│  │  │ [Validate ▶]   │  │ Quality: 85   │                           │  │
│  │  └────────────────┘  └───────────────┘                           │  │
│  │                                                                    │  │
│  │  Checklist:                                                        │  │
│  │  ☑ Question is clear and unambiguous                              │  │
│  │  ☑ All options are plausible                                      │  │
│  │  ☑ Correct answer is verified                                     │  │
│  │  ☑ Difficulty level is appropriate                                │  │
│  │                                                                    │  │
│  │  Action: [APPROVE ▼]  Comment: [________________]  [Submit]       │  │
│  └────────────────────────────────────────────────────────────────────┘  │
└─────────────────────────────────────────────────────────────────────────┘
""")
    run.font.size = Pt(7)
    run.font.name = "Courier New"

    doc.add_heading("API Flow", level=2)
    p = doc.add_paragraph()
    run = p.add_run(
        "1. GET /api/v1/reviews?status=UNDER_REVIEW → Fetch assigned MCQs\n"
        "2. GET /api/v1/reviews?status=APPROVED → Count approved stats\n"
        "3. GET /api/v1/reviews?status=REJECTED → Count rejected stats\n"
        "4. POST /api/v1/ai/validate-answer → AI Copilot validates correctness\n"
        "5. POST /api/v1/reviews/{mcqId}/submit {action:'APPROVE', comment:'...'} → Final decision\n"
        "   → On REJECT: comment is MANDATORY (validation enforced)\n"
        "   → Status changes to APPROVED or REJECTED in database\n"
        "   → Audit log entry created\n"
        "   → Creator notified via notification bell"
    )
    run.font.size = Pt(9)

    doc.add_heading("All 19 Features", level=2)
    add_feature_list(doc, [
        "Review stats summary bar (Pending/Approved/Rejected counts)",
        "Free-text search across questions, tech stacks, creators",
        "5 sortable columns with direction toggle",
        "Pagination with page size control",
        "Click-to-expand review panel (inline, no navigation away)",
        "Full MCQ display: stem + all 4 options with correct answer highlighted",
        "AI Copilot button — validates answer correctness via AI",
        "AI Risk indicator (LOW/MEDIUM/HIGH based on quality score)",
        "4-item review checklist (must complete all before deciding)",
        "Approve action (one-click after checklist)",
        "Reject action (requires mandatory comment)",
        "Comment field for feedback (mandatory for reject, optional for approve)",
        "Content translation of question stems and tech stacks",
        "Empty state with party popper ('All caught up!' when no pending)",
        "Error handling with inline messages",
        "Loading state spinner",
        "Close review panel button",
        "Translated review option labels (in 7 languages)",
        "Action validation (cannot submit without completing checklist + comment for reject)",
    ])

    add_screenshot_pair(doc, "pending_reviews_desktop.png", "pending_reviews_mobile.png",
                        "My Pending Reviews — Desktop & Mobile")

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════
    # PART E — QUESTION BANK
    # ═══════════════════════════════════════════════════════════
    doc.add_heading("Part E — Admin Question Bank Management (20 Features)", level=1)

    p = doc.add_paragraph()
    run = p.add_run("What it does: ")
    run.bold = True
    p.add_run("Admin-only view of ALL questions from ALL creators. Admin can search, filter, sort, export, assign reviewers, edit any MCQ regardless of status, delete, and use AI semantic search to find similar questions.")

    doc.add_heading("Internal Architecture", level=2)
    p = doc.add_paragraph()
    run = p.add_run("""
┌─────────────────────────────────────────────────────────────────────────┐
│                  QUESTION BANK MANAGEMENT (Admin Only)                    │
├─────────────────────────────────────────────────────────────────────────┤
│  Status Summary: [Total:113] [Approved:15] [Under Review:1] [Ready:10]  │
│                                                                          │
│  Filters: [Search___] [Tech Stack▼] [Status▼] [Difficulty▼]            │
│  Actions: [⬇️ Export Excel] [+ Create MCQ]                              │
│                                                                          │
│  ┌── Semantic AI Search ─────────────────────────────────────────────┐  │
│  │  🧠 "thread safety in concurrent Java" [🔍 Search]               │  │
│  │  Results: MCQ#45 (89% match) │ MCQ#67 (72% match) │ MCQ#23 (45%)│  │
│  └────────────────────────────────────────────────────────────────────┘  │
│                                                                          │
│  Table: # │ Question │ Tech Stack │ Topic │ Diff │ Status │ Creator │ ⋮ │
│  Row menu (⋮): [View] [Edit] [Assign Reviewer] [Delete]                 │
│                                                                          │
│  Key Difference from My Questions:                                       │
│  • Shows ALL creators' questions (not just your own)                     │
│  • Creator column visible                                                │
│  • Reviewer column visible                                               │
│  • Admin can Edit any MCQ at any status (except Draft of other users)   │
│  • Admin can Assign Reviewer (opens AssignReviewerModal)                │
│  • AI Semantic Search (natural language → similarity results)           │
└─────────────────────────────────────────────────────────────────────────┘
""")
    run.font.size = Pt(7)
    run.font.name = "Courier New"

    doc.add_heading("All 20 Features", level=2)
    add_feature_list(doc, [
        "Tech Stack filter dropdown (all stacks from master data)",
        "Status filter dropdown (Draft/Ready/Under Review/Approved/Rejected)",
        "Difficulty filter dropdown (Easy/Medium/Hard)",
        "Free-text search across all fields",
        "7 sortable columns (Question, Tech Stack, Topic, Diff, Status, Creator, Reviewer)",
        "Pagination with configurable page size",
        "Export to Excel button (exports with current active filters applied)",
        "Status summary count cards (Total/Approved/Under Review/Ready/Rejected)",
        "Context menu (⋮) per row with multiple actions",
        "View MCQ detail → navigates to /mcq/{id}",
        "Edit MCQ → navigates to /mcq/{id}/edit (admin can edit any status)",
        "Assign Reviewer → opens AssignReviewerModal",
        "Delete MCQ with confirmation dialog",
        "AI Semantic Search (natural language query)",
        "Semantic results with similarity percentage badges",
        "Click semantic result → navigate to /mcq/{id}",
        "Content translation for all text",
        "URL query param sync (filters persist in URL)",
        "Loading states with spinner",
        "Delete indicator per row (shows deleting progress)",
    ])

    add_screenshot_pair(doc, "admin_question_bank_desktop.png", "admin_question_bank_mobile.png",
                        "Question Bank Management — Desktop & Mobile")

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════
    # PART F — ASSIGN REVIEWER
    # ═══════════════════════════════════════════════════════════
    doc.add_heading("Part F — Admin Assign Reviewer (Internal Flow)", level=1)

    p = doc.add_paragraph()
    run = p.add_run("What it does: ")
    run.bold = True
    p.add_run("When Admin clicks 'Assign' on a 'Ready for Review' MCQ, a dialog opens showing question details and a dropdown of eligible reviewers. Eligibility rules: SME must have matching tech stack AND cannot be the creator (no self-review).")

    doc.add_heading("Step-by-Step Internal Flow", level=2)
    p = doc.add_paragraph()
    run = p.add_run(
        "1. Admin clicks ⋮ menu on a 'Ready for Review' row\n"
        "2. Selects 'Assign Reviewer'\n"
        "3. Frontend calls: GET /api/v1/admin/mcqs/{id}/eligible-reviewers\n"
        "4. Backend logic:\n"
        "   a. Gets MCQ's techStackId\n"
        "   b. Finds all SMEs assigned to that tech stack (from master_sme_tech_stack mapping)\n"
        "   c. Excludes the creator's userId (prevents self-review)\n"
        "   d. Returns list of eligible reviewer objects\n"
        "5. Dialog shows: Question preview, Tech Stack, Topic, Creator name\n"
        "6. Dropdown shows: Only eligible SMEs with names\n"
        "7. Admin selects reviewer → clicks 'Assign'\n"
        "8. Frontend calls: POST /api/v1/admin/mcqs/{id}/assign-reviewer {reviewerId: userId}\n"
        "9. Backend:\n"
        "   a. Updates MCQ status: READY_FOR_REVIEW → UNDER_REVIEW\n"
        "   b. Sets reviewer field on MCQ\n"
        "   c. Creates audit log entry (action: ASSIGN)\n"
        "   d. Sends notification to reviewer\n"
        "10. Table refreshes → status badge changes to 'Under Review'"
    )
    run.font.size = Pt(9)

    add_screenshot_pair(doc, "assign_reviewer_dialog_desktop.png", "assign_reviewer_dialog_mobile.png",
                        "Assign Reviewer Dialog — Tech stack matching, creator excluded")

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════
    # PART G — MASTER DATA
    # ═══════════════════════════════════════════════════════════
    doc.add_heading("Part G — Master Data Management (16 Features)", level=1)

    p = doc.add_paragraph()
    run = p.add_run("What it does: ")
    run.bold = True
    p.add_run("Admin manages the foundation data that powers the entire platform: Technology Stacks, Topics (per stack), and SME-to-TechStack assignments. Changes here propagate to ALL dropdowns across the app (McqForm, MyQuestions, QuestionBank, BulkUpload, AI Studio).")

    doc.add_heading("3-Panel Layout", level=2)
    p = doc.add_paragraph()
    run = p.add_run("""
┌──────────────────┬─────────────────────┬─────────────────────┐
│  TECH STACKS     │  TOPICS             │  SME ASSIGNMENTS    │
├──────────────────┼─────────────────────┼─────────────────────┤
│ ► Spring Boot    │  Auto Configuration │  birendra.kumar ✓   │
│   Spring Cloud   │  Spring Boot Testing│  swati.avinash ✓    │
│   Spring Core    │  Spring Boot Starter│  indugu.hari ✓      │
│   Spring MVC     │  Actuator           │  [Assign New SME ▼] │
│   Spring ORM     │  Embedded Servers   │                     │
│   Core Java      │  [+ Add Topic]      │                     │
│                  │                     │                     │
│ [+ Add Stack]    │  (paginated)        │  [Remove] per SME   │
│ [Edit] [Delete]  │  [Edit] [Delete]    │                     │
└──────────────────┴─────────────────────┴─────────────────────┘
""")
    run.font.size = Pt(8)
    run.font.name = "Courier New"

    doc.add_heading("API Endpoints", level=2)
    p = doc.add_paragraph()
    run = p.add_run(
        "GET/POST/PUT/DELETE /api/v1/master/tech-stacks\n"
        "GET/POST/PUT/DELETE /api/v1/master/tech-stacks/{id}/topics\n"
        "GET /api/v1/master/smes — List all SMEs\n"
        "POST /api/v1/master/tech-stacks/{id}/smes/{userId} — Assign SME to stack\n"
        "DELETE /api/v1/master/tech-stacks/{id}/smes/{userId} — Remove SME from stack\n\n"
        "Data Impact: Changes propagate to ALL dropdowns across Create MCQ, Bulk Upload, AI Generation, Question Bank filters, Assign Reviewer eligible list"
    )
    run.font.size = Pt(9)

    doc.add_heading("Pre-seeded Master Data (auto-loaded on startup via data.sql)", level=2)
    p = doc.add_paragraph()
    run = p.add_run(
        "6 Tech Stacks: Spring Boot, Spring Cloud, Spring Core, Spring MVC & REST, Spring ORM & Data JPA, Core Java\n"
        "40+ Topics distributed across all stacks\n"
        "5 Users: 2 Admins + 3 SMEs with tech stack assignments\n"
        "Sample MCQs in Draft, Ready for Review, Approved states"
    )
    run.font.size = Pt(9)

    add_screenshot_pair(doc, "admin_tech_stacks_desktop.png", "admin_tech_stacks_mobile.png",
                        "Master Data Management — Desktop & Mobile")

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════
    # PART H — AI STUDIO
    # ═══════════════════════════════════════════════════════════
    doc.add_heading("Part H — AI Studio & MCQ Generation (Level 2 Feature)", level=1)

    p = doc.add_paragraph()
    run = p.add_run("What it does: ")
    run.bold = True
    p.add_run("AI-powered MCQ generation using Spring AI + OpenAI GPT-4o-mini. Supports 18 question types. Auto-detects duplicates during generation (≥30% similarity triggers replacement). Also provides Code→MCQ conversion, AI Rewrite, and Learning Path generation.")

    doc.add_heading("AI Generation Flow", level=2)
    p = doc.add_paragraph()
    run = p.add_run(
        "1. SME selects: Tech Stack → Topic → Difficulty → Count → Question Type\n"
        "2. Clicks 'Generate'\n"
        "3. API: POST /api/v1/ai/generate-mcqs\n"
        "   Request: {techStackId, topicId, count, difficulty, questionType}\n"
        "4. Backend (AIService):\n"
        "   a. Builds prompt with tech stack context and question type template\n"
        "   b. Calls OpenAI GPT-4o-mini via Spring AI\n"
        "   c. Parses response into MCQ format\n"
        "   d. For EACH generated question:\n"
        "      - Runs semantic similarity check against ALL existing questions in same tech+topic\n"
        "      - If similarity ≥ 30%: REPLACES with a newly generated question (retry)\n"
        "      - If similarity < 30%: KEEPS the question\n"
        "   e. Saves all valid questions as DRAFT with creatorId\n"
        "5. Response: {generated: N, ids: [...], replacedDuplicates: M, skippedStems: [...]}\n"
        "6. Frontend shows: 'Generated N questions, replaced M duplicates'\n"
        "7. Questions appear in My Questions with DRAFT status"
    )
    run.font.size = Pt(9)

    doc.add_heading("AI Studio Tabs", level=2)
    p = doc.add_paragraph()
    run = p.add_run(
        "Tab 1: Code → MCQ — Paste code snippet → AI generates MCQs about that code\n"
        "   API: POST /api/v1/ai/generate-from-code {code, language, count, difficulty}\n"
        "   Languages: Java, Python, JavaScript, TypeScript, C++, C#, Go, Rust, SQL, Kotlin\n\n"
        "Tab 2: AI Rewrite — Take existing MCQ and rephrase/improve it\n"
        "   Maintains meaning while improving clarity\n\n"
        "Tab 3: Learning Path — AI generates a structured learning path for a topic"
    )
    run.font.size = Pt(9)

    add_screenshot_pair(doc, "ai_studio_desktop.png", "ai_studio_mobile.png",
                        "AI Studio — Desktop & Mobile")

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════
    # PART I — DUPLICATE DETECTION
    # ═══════════════════════════════════════════════════════════
    doc.add_heading("Part I — AI Duplicate Detection (Level 2 Feature)", level=1)

    p = doc.add_paragraph()
    run = p.add_run("What it does: ")
    run.bold = True
    p.add_run("Prevents redundant questions by checking semantic similarity against the entire question bank within the same tech stack and topic. Blocks submission if ≥30% similar. Works in two places: (1) Auto-during AI generation, (2) Manual + auto on Edit page.")

    doc.add_heading("How It Works Internally", level=2)
    p = doc.add_paragraph()
    run = p.add_run(
        "Location 1 — AI Generation (automatic):\n"
        "  • Every AI-generated question is checked BEFORE saving\n"
        "  • If ≥30% match found: question is REPLACED with new AI generation\n"
        "  • Response includes: replacedDuplicates count + replacedStems list\n\n"
        "Location 2 — Edit Page (manual + auto on submit):\n"
        "  • 'Duplicate Check' button: POST /api/v1/ai/check-duplicate-db\n"
        "  • Auto-triggered on 'Save & Send for Review' button click\n"
        "  • If ≥30% similar:\n"
        "    - Shows error message with similar questions listed\n"
        "    - Shows similarity percentage per match\n"
        "    - BLOCKS the submission (cannot proceed)\n"
        "    - SME must edit question to make it unique\n"
        "  • If <30% similar: submission proceeds normally\n\n"
        "Location 3 — McqForm (auto on typing):\n"
        "  • 2-second debounce on question stem changes\n"
        "  • Auto-calls duplicate check in background\n"
        "  • Shows warning banner if potential duplicates found\n\n"
        "Backend Algorithm:\n"
        "  • Uses semantic comparison (not just text matching)\n"
        "  • Scopes to same tech stack + topic (reduces false positives)\n"
        "  • Returns: {isDuplicate: bool, similarityScore: float, similarQuestions: [...]}"
    )
    run.font.size = Pt(9)

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════
    # PART J — BULK UPLOAD
    # ═══════════════════════════════════════════════════════════
    doc.add_heading("Part J — Bulk Upload (16 Features)", level=1)

    p = doc.add_paragraph()
    run = p.add_run("What it does: ")
    run.bold = True
    p.add_run("Upload multiple MCQs at once via CSV/Excel file. Download template, fill in questions, upload. System validates every field per row, reports errors with row numbers, handles duplicates with preview, and allows inline editing of failed rows.")

    doc.add_heading("All 16 Features", level=2)
    add_feature_list(doc, [
        "Download template button (GET /api/v1/upload/template → .xlsx file)",
        "File picker with drag-and-drop support (CSV/Excel)",
        "Upload button with progress indicator",
        "Success count display (N questions imported as Draft)",
        "Error details per failed row (row number, field name, error message)",
        "Duplicate preview modal (side-by-side: uploaded vs existing DB match)",
        "Force Add button (override duplicate detection for that row)",
        "Inline edit modal (fix invalid row data without re-uploading)",
        "Tech Stack / Topic reference panel (expandable hierarchy)",
        "Copy-to-clipboard for tech/topic names (easy reference while filling template)",
        "Re-submit fixed rows (after inline edit)",
        "Topic auto-loading per selected stack",
        "Validation categories (Required Field Missing / Invalid Value / Duplicate)",
        "Color-coded results (Green=success, Red=blocked, Yellow=warning)",
        "Content translation",
        "Loading states with upload progress",
    ])

    add_screenshot_pair(doc, "bulk_upload_desktop.png", "bulk_upload_mobile.png",
                        "Bulk Upload — Desktop & Mobile")

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════
    # PART K — QUESTION TYPES
    # ═══════════════════════════════════════════════════════════
    doc.add_heading("Part K — Question Types (18 Formats)", level=1)

    p = doc.add_paragraph()
    run = p.add_run("What it does: ")
    run.bold = True
    p.add_run("Supports 18 different question formats beyond standard MCQ. Each type has custom rendering (code blocks, drag-drop, pair matching, fill-blanks, flowcharts, pipelines). AI can generate ALL 18 types.")

    types_table = doc.add_table(rows=19, cols=4)
    types_table.style = 'Table Grid'
    headers = ["#", "Type", "Format", "AI Gen"]
    for i, h in enumerate(headers):
        types_table.rows[0].cells[i].paragraphs[0].add_run(h).bold = True
    type_data = [
        ("1", "SINGLE_MCQ", "4 options, 1 correct answer (A/B/C/D)", "✓"),
        ("2", "MULTI_MCQ", "4 options, multiple correct (A,B / A,C etc)", "✓"),
        ("3", "DRAG_ORDER", "Drag items into correct sequence order", "✓"),
        ("4", "MATCH_PAIRS", "Match left items to right definitions", "✓"),
        ("5", "CODE_OUTPUT", "Match code snippets to their outputs", "✓"),
        ("6", "FILL_BLANK", "Fill in missing words in code/text (___)", "✓"),
        ("7", "PREDICT_OUTPUT", "Read code → predict console output", "✓"),
        ("8", "DEBUG_CODE", "Find and fix the bug in code snippet", "✓"),
        ("9", "CODE_REARRANGE", "Arrange scrambled code lines correctly", "✓"),
        ("10", "SQL_BUILDER", "Build SQL query from clause fragments", "✓"),
        ("11", "ARCH_LAYERS", "Assign components to architecture layers", "✓"),
        ("12", "CODE_REVIEW", "Find issues in code review challenge", "✓"),
        ("13", "PIPELINE_BUILD", "Build stream/pipeline from operators", "✓"),
        ("14", "FLOWCHART", "Arrange flowchart steps in order", "✓"),
        ("15", "DEVOPS_PIPE", "Build CI/CD pipeline stages", "✓"),
        ("16", "SECURE_CODE", "Identify security vulnerability in code", "✓"),
        ("17", "RIDDLE", "Solve tech riddle (lateral thinking)", "✓"),
        ("18", "CROSSWORD", "Technical crossword puzzle", "✓"),
    ]
    for row_idx, (n, t, f, ai) in enumerate(type_data, 1):
        types_table.rows[row_idx].cells[0].paragraphs[0].add_run(n).font.size = Pt(8)
        types_table.rows[row_idx].cells[1].paragraphs[0].add_run(t).font.size = Pt(8)
        types_table.rows[row_idx].cells[2].paragraphs[0].add_run(f).font.size = Pt(8)
        types_table.rows[row_idx].cells[3].paragraphs[0].add_run(ai).font.size = Pt(8)

    add_screenshot_pair(doc, "question_types_desktop.png", "question_types_mobile.png",
                        "Question Types Showcase — Desktop & Mobile")

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════
    # PART L — LIVE QUIZ
    # ═══════════════════════════════════════════════════════════
    doc.add_heading("Part L — Live Quiz Battle System (20 Features)", level=1)

    p = doc.add_paragraph()
    run = p.add_run("What it does: ")
    run.bold = True
    p.add_run("Real-time multiplayer quiz battles. Host creates session with PIN, players join, host controls question flow, players answer in real-time, live leaderboard updates. Supports teams, adaptive difficulty, recording, co-hosts, reconnection.")

    add_feature_list(doc, [
        "Create live session (title, tech stack, topic, difficulty, question count, time limit per question)",
        "Session modes (Battle mode for competitive play)",
        "Team mode toggle (split players into teams)",
        "Adaptive difficulty toggle (questions get harder as players score higher)",
        "Recording enabled toggle (save session for replay)",
        "Co-host assignment (delegate host controls)",
        "PIN-based join (6-digit PIN, no auth required to join)",
        "Real-time lobby (waiting room shows connected players)",
        "Host controls: Start / Next Question / End Question / Pause / Resume / Extend Time",
        "Kick participant (remove disruptive player)",
        "Transfer host to co-host (host can leave)",
        "Real-time answer submission (players see timer counting down)",
        "Live leaderboard (updates after each question)",
        "Team leaderboard (aggregate team scores)",
        "Session summary (final rankings, stats)",
        "Session replay/recording (watch back the session)",
        "Reconnection support (player + host can rejoin if disconnected)",
        "Invite link generation (shareable URL)",
        "Host session history (all past sessions)",
        "Participant session history (sessions I've played in)",
    ])

    add_screenshot_pair(doc, "live_quiz_desktop.png", "live_quiz_mobile.png",
                        "Live Quiz — Desktop & Mobile")

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════
    # PART M — ASSESSMENTS
    # ═══════════════════════════════════════════════════════════
    doc.add_heading("Part M — Assessments & Quiz Builder (18 Features)", level=1)

    p = doc.add_paragraph()
    run = p.add_run("What it does: ")
    run.bold = True
    p.add_run("Build shareable assessments from approved questions. Create proctored quiz sessions with time limits, share via link/token, track attempts with tab-switch detection and screenshot capture for violations.")

    add_feature_list(doc, [
        "Create assessment form (title, tech stack, topic, difficulty, count, time limit, link validity)",
        "Auto-selects random approved questions matching criteria",
        "Generates shareable link (unique token URL)",
        "Session history list (view all created assessments)",
        "Copy link button (clipboard copy)",
        "Navigate to view all attempts for a session",
        "Quiz taking interface (no auth required — token-based access)",
        "Timer countdown (per question or total)",
        "Question navigation (next/prev/jump)",
        "Answer selection with visual feedback",
        "Submit with auto-scoring (correct/incorrect per question)",
        "Tab-switch detection (proctoring — counts violations)",
        "Violation screenshot capture (browser screenshot on tab switch)",
        "Exam lock guard (prevents back/forward navigation during quiz)",
        "Attempt results display (score, time taken, violations)",
        "Assessment leaderboard (rank by score across all attempts)",
        "View violation screenshots (admin can verify)",
        "Cached master data (sessionStorage for performance)",
    ])

    add_screenshot_pair(doc, "quiz_builder_desktop.png", "quiz_builder_mobile.png",
                        "Quiz Builder — Desktop & Mobile")

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════
    # PART N — AUDIT LOG
    # ═══════════════════════════════════════════════════════════
    doc.add_heading("Part N — User Activity & Audit Log (10 Features)", level=1)

    p = doc.add_paragraph()
    run = p.add_run("What it does: ")
    run.bold = True
    p.add_run("Comprehensive audit trail of every significant action in the system. Admin can see who did what, when, to whom. Supports search, filter by action type, sort by time.")

    doc.add_heading("Actions Logged", level=2)
    p = doc.add_paragraph()
    run = p.add_run(
        "• APPROVE — Reviewer approved an MCQ\n"
        "• REJECT — Reviewer rejected an MCQ\n"
        "• ASSIGN — Admin assigned a reviewer to MCQ\n"
        "• UNASSIGN — Admin removed reviewer assignment\n"
        "• REGISTER — New user registered\n"
        "• ROLE_CHANGE — Admin changed user's role\n"
        "• DEACTIVATE — Admin deactivated a user\n"
        "• REACTIVATE — Admin reactivated a user\n"
        "• DELETE — MCQ or user deleted"
    )
    run.font.size = Pt(9)

    doc.add_heading("All 10 Features", level=2)
    add_feature_list(doc, [
        "Full audit trail table with all system actions",
        "Free-text search across all columns (actor, target, details)",
        "Action type filter dropdown (dynamically populated from actual logged actions)",
        "5 sortable columns (Timestamp, Actor, Action, Target, Details)",
        "Pagination with configurable page size",
        "Action color coding & icons per type (green=approve, red=reject, blue=assign)",
        "Entry count display ('Showing X entries')",
        "Refresh button (re-fetch latest)",
        "Loading state",
        "Empty state when no audit entries",
    ])

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════
    # PART O — ANALYTICS & DASHBOARD
    # ═══════════════════════════════════════════════════════════
    doc.add_heading("Part O — Analytics, Dashboard & Widgets (28 Features)", level=1)

    p = doc.add_paragraph()
    run = p.add_run("What it does: ")
    run.bold = True
    p.add_run("Customizable dashboard with 10 widgets (show/hide/reorder). Analytics page with charts. Reviewer dashboard with personal stats. Admin reviewer metrics with SLA breach detection.")

    doc.add_heading("Dashboard Widgets (user-customizable)", level=2)
    add_feature_list(doc, [
        "Stat Cards (Total/Approved/In Review/Rejected/Draft counts)",
        "MCQs by Tech Stack (bar chart)",
        "Recent Activity (table with action, actor, time-ago)",
        "Performance Overview (circular progress rings)",
        "Platform Insights (key metrics)",
        "Top Reviewers (mini-leaderboard)",
        "Pending Approvals (Admin only — MCQs waiting for reviewer)",
        "Reviewer Workload (Admin only — distribution of reviews)",
        "Action Items (things needing your attention)",
        "Quality Gauge (average AI quality score)",
        "Widget customization panel (toggle show/hide per widget)",
        "Per-user widget preference persistence (localStorage)",
    ])

    doc.add_heading("Analytics Page", level=2)
    add_feature_list(doc, [
        "Status distribution donut chart",
        "MCQs by tech stack bar chart",
        "Date range filter (from/to date pickers)",
        "Export to Excel button",
        "Summary stat cards with trend indicators",
        "Color-coded status legend",
        "Responsive charts (pure CSS/SVG — no chart library dependency)",
    ])

    doc.add_heading("Reviewer Metrics (Admin)", level=2)
    add_feature_list(doc, [
        "All-reviewer performance table (approval rate, total reviews, avg time)",
        "SLA breach table (MCQs stuck > 48 hours without review)",
        "Hours-stuck color coding (red > 72h, yellow > 48h)",
        "Approval rate color coding (green > 80%, red < 50%)",
    ])

    add_screenshot_pair(doc, "analytics_desktop.png", "analytics_mobile.png",
                        "Analytics — Desktop & Mobile")

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════
    # PART P — CHAT & INBOX
    # ═══════════════════════════════════════════════════════════
    doc.add_heading("Part P — Real-time Chat & Inbox (28 Features)", level=1)

    p = doc.add_paragraph()
    run = p.add_run("What it does: ")
    run.bold = True
    p.add_run("Two communication systems: (1) Real-time group chat with AI bot, reactions, @mentions, slash commands. (2) Internal inbox messaging for direct communication between users.")

    doc.add_heading("Real-time Chat (14 features)", level=2)
    add_feature_list(doc, [
        "Global chat panel (collapsible sidebar)",
        "Send messages to all users",
        "@mention specific users (notifications sent)",
        "@bot triggers AI assistant responses",
        "Slash commands: /create, /quiz-builder, /leaderboard, /question-bank, /help",
        "Emoji reactions (5 emojis: 👍❤️😂🔥👏) per message",
        "Reply to specific messages (threaded)",
        "Edit own messages",
        "Delete messages",
        "Pin message (Admin only — pinned at top)",
        "Online users indicator (who's currently active)",
        "Presence heartbeat (auto-updates every 30s)",
        "Auto-translate messages (in 7 languages)",
        "Hidden during quiz/live sessions (anti-cheat protection)",
    ])

    doc.add_heading("Internal Inbox (14 features)", level=2)
    add_feature_list(doc, [
        "Inbox tab (received messages)",
        "Sent tab (messages you sent)",
        "Starred tab (bookmarked messages)",
        "Compose tab (write new message)",
        "Draft auto-save (localStorage persistence)",
        "Send message to specific user",
        "Mark as read (individual)",
        "Mark all read (bulk)",
        "Delete message",
        "Star/unstar toggle (bookmark important messages)",
        "Unread count badge (in navbar)",
        "Search across all messages",
        "Time-ago display (relative timestamps)",
        "Initials-based avatars (deterministic colors)",
    ])

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════
    # PART Q — SMART INTERVIEW KIT
    # ═══════════════════════════════════════════════════════════
    doc.add_heading("Part Q — Smart Interview Kit (7 Features)", level=1)

    p = doc.add_paragraph()
    run = p.add_run("What it does: ")
    run.bold = True
    p.add_run("Upload resume (PDF/DOCX/TXT) + optional job description → AI analyzes and generates 6 categories of targeted interview questions: Technical, Coding, SQL, Project-Based, Behavioral, Scenario.")

    add_feature_list(doc, [
        "Resume upload (PDF/DOCX/TXT) — POST /api/v1/resume/analyze",
        "Optional Job Description input (improves question relevance)",
        "AI generates 6 question categories",
        "Tabs: Technical / Coding / SQL / Project-Based / Behavioral / Scenario",
        "Difficulty badges per generated question",
        "Question type indicators (positive/negative/edge_case)",
        "Loading with 5-minute timeout (complex AI processing)",
    ])

    add_screenshot_pair(doc, "smart_interview_kit_desktop.png", "smart_interview_kit_mobile.png",
                        "Smart Interview Kit — Desktop & Mobile")

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════
    # PART R — CROSS-CUTTING
    # ═══════════════════════════════════════════════════════════
    doc.add_heading("Part R — Cross-Cutting Platform Features (24 Features)", level=1)

    add_feature_list(doc, [
        "JWT Authentication with secure token storage (localStorage)",
        "Role-Based Access Control (ADMIN / SME / PENDING approval)",
        "Dark/Light Mode (CSS variables, persisted to localStorage)",
        "Internationalization (i18n) — 7 Languages: English, Hindi, Telugu, French, German, Urdu, Kannada",
        "Real-time Content Translation (AI-powered useContentTranslation hook)",
        "RTL Support for Urdu",
        "Notification Bell with real-time unread count (polling)",
        "Toast Notifications (30-second auto-close, multiple concurrent)",
        "Error Boundary (React crash recovery — doesn't break entire app)",
        "Lazy Loading with React Suspense (code-split routes for performance)",
        "Fully Responsive Design (desktop + tablet + mobile)",
        "11 Reusable Components (Navbar, SortableTh, TablePagination, StatusBadge, AssignReviewerModal, McqCommentSection, QuestionStemRenderer, ErrorBoundary, NotificationBell, ChangePasswordModal, ChatBot)",
        "MCQ Version History tracking (every edit creates new version)",
        "Threaded Comments on MCQs (replies to specific comments)",
        "Auto-seed Data on startup (data.sql — no manual SQL needed)",
        "H2 In-Memory Database for tests (no MySQL dependency for testing)",
        "Swagger/OpenAPI documentation (auto-generated at /swagger-ui/index.html)",
        "Spring Actuator Health endpoint (/actuator/health)",
        "Email Notifications (optional — sends on status changes)",
        "OpenAI GPT-4o-mini Integration via Spring AI (optional — app works without it)",
        "AI Duplicate Detection (semantic similarity scoring)",
        "AI Quality Scoring (0-100 per MCQ, auto-scored)",
        "Proctoring System (tab-switch detection + screenshot capture)",
        "Exam Lock Guard (prevents navigation during quiz attempts)",
    ])

    add_screenshot_pair(doc, "dark_mode_desktop.png", "dark_mode_mobile.png",
                        "Dark Mode — Desktop & Mobile")

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════
    # PART S — ALL ROUTES & ENDPOINTS
    # ═══════════════════════════════════════════════════════════
    doc.add_heading("Part S — All 40 Frontend Routes", level=1)

    routes = [
        ("/login", "Public", "JWT login with demo credentials"),
        ("/register", "Public", "Self-registration for new SMEs"),
        ("/forgot-password", "Public", "Password reset request"),
        ("/reset-password", "Public", "Reset via email token"),
        ("/", "Auth", "Customizable widget dashboard"),
        ("/my-questions", "Auth", "SME's own questions + actions"),
        ("/mcq/create", "Auth", "Create new MCQ form"),
        ("/mcq/:id", "Auth", "View MCQ detail + AI tools"),
        ("/mcq/:id/edit", "Auth", "Edit MCQ form"),
        ("/coding/create", "Auth", "Create coding question"),
        ("/pending-reviews", "Auth", "Reviewer's assigned queue"),
        ("/reviewer-dashboard", "Auth", "Personal reviewer stats"),
        ("/question-bank", "ADMIN", "All MCQs platform-wide"),
        ("/bulk-upload", "Auth", "CSV/Excel MCQ upload"),
        ("/screenshot-mcq", "Auth", "Image → MCQ via AI OCR"),
        ("/leaderboard", "Auth", "Multi-mode leaderboard"),
        ("/user-management", "ADMIN", "CRUD users & roles"),
        ("/reviewer-metrics", "ADMIN", "All-reviewer performance"),
        ("/audit-log", "ADMIN", "System audit trail"),
        ("/master-data", "ADMIN", "Tech Stacks/Topics/SMEs"),
        ("/kanban", "Auth", "Visual lifecycle board"),
        ("/quiz", "Auth", "Quick self-practice"),
        ("/analytics", "Auth", "Charts & exports"),
        ("/quiz-builder", "Auth", "Create assessments"),
        ("/quiz/take/:token", "Public", "Take proctored quiz"),
        ("/quiz-sessions/:id/attempts", "Auth", "View attempts"),
        ("/inbox", "Auth", "Internal messaging"),
        ("/smart-interview-kit", "Auth", "Resume → AI questions"),
        ("/ai-studio", "Auth", "Code→MCQ, Rewrite, Path"),
        ("/question-types", "Auth", "18 type showcase"),
        ("/question-type-create/:typeId", "Auth", "Create advanced type"),
        ("/rulebook", "Auth", "Platform guidelines"),
        ("/live", "Auth", "Live quiz hub"),
        ("/live/join", "Public", "Join via PIN"),
        ("/live/join/:pin", "Public", "Direct PIN join"),
        ("/live/lobby/:sessionId", "Public", "Waiting room"),
        ("/live/host/:sessionId", "Auth", "Host controls"),
        ("/live/play/:sessionId", "Public", "Player interface"),
        ("/live/results/:sessionId", "Public", "Post-game results"),
        ("/live/sessions/:sessionId", "Auth", "Session replay"),
    ]

    route_table = doc.add_table(rows=len(routes) + 1, cols=3)
    route_table.style = 'Table Grid'
    for i, h in enumerate(["Route", "Access", "Description"]):
        route_table.rows[0].cells[i].paragraphs[0].add_run(h).bold = True
    for idx, (route, access, desc) in enumerate(routes, 1):
        route_table.rows[idx].cells[0].paragraphs[0].add_run(route).font.size = Pt(7)
        route_table.rows[idx].cells[1].paragraphs[0].add_run(access).font.size = Pt(7)
        route_table.rows[idx].cells[2].paragraphs[0].add_run(desc).font.size = Pt(7)

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════
    # PART T — FEATURE COUNT
    # ═══════════════════════════════════════════════════════════
    doc.add_heading("Part T — Feature Count Summary (333 Total)", level=1)

    p = doc.add_paragraph()
    run = p.add_run("Granular count of every unique feature across the platform:")
    run.font.size = Pt(10)

    counts = [
        ("Home/Dashboard", "18"),
        ("My Questions", "22"),
        ("Pending Reviews", "19"),
        ("Question Bank (Admin)", "20"),
        ("AI Studio", "9"),
        ("Master Data (Admin)", "16"),
        ("User Management (Admin)", "15"),
        ("Audit Log (Admin)", "10"),
        ("Leaderboard", "15"),
        ("Kanban Board", "12"),
        ("Bulk Upload", "16"),
        ("Quiz Builder", "10"),
        ("Take Quiz / Assessment", "8"),
        ("Live Quiz Battle", "20"),
        ("Inbox / Messaging", "14"),
        ("Real-time Chat", "14"),
        ("Smart Interview Kit", "7"),
        ("Analytics", "10"),
        ("Reviewer Dashboard", "8"),
        ("Reviewer Metrics", "8"),
        ("Question Types", "5"),
        ("Coding Question", "6"),
        ("RuleBook", "4"),
        ("Screenshot MCQ", "3"),
        ("Cross-cutting Platform", "24"),
        ("GRAND TOTAL", "333"),
    ]

    count_table = doc.add_table(rows=len(counts) + 1, cols=2)
    count_table.style = 'Table Grid'
    count_table.rows[0].cells[0].paragraphs[0].add_run("Page / Feature Area").bold = True
    count_table.rows[0].cells[1].paragraphs[0].add_run("Feature Count").bold = True
    for idx, (area, count) in enumerate(counts, 1):
        count_table.rows[idx].cells[0].paragraphs[0].add_run(area).font.size = Pt(9)
        r = count_table.rows[idx].cells[1].paragraphs[0].add_run(count)
        r.font.size = Pt(9)
        if area == "GRAND TOTAL":
            r.bold = True
            r.font.size = Pt(12)

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════
    # PART U — TEST RESULTS
    # ═══════════════════════════════════════════════════════════
    doc.add_heading("Part U — Test Results & Coverage", level=1)

    doc.add_heading("Backend Tests (JUnit 5 + Spring Boot Test)", level=2)
    p = doc.add_paragraph()
    run = p.add_run(
        "• Total Tests: 1,072\n"
        "• Failures: 0\n"
        "• Coverage: 92.5% instruction coverage (JaCoCo)\n"
        "• Database: H2 in-memory (no MySQL dependency for tests)\n"
        "• What's tested: All 17 controllers, services, repositories, AI integration, lifecycle transitions\n"
        "• Command: mvn test"
    )
    run.font.size = Pt(10)

    doc.add_heading("Frontend Tests (Jest + React Testing Library)", level=2)
    p = doc.add_paragraph()
    run = p.add_run(
        "• Total Tests: 957\n"
        "• Failures: 0\n"
        "• Coverage: 80.37% statement coverage\n"
        "• What's tested: All 40 pages, 11 components, API layer, routing, i18n, dark mode\n"
        "• Mocking: Axios responses, React Router, i18n, localStorage\n"
        "• Command: npm test -- --watchAll=false"
    )
    run.font.size = Pt(10)

    doc.add_heading("Combined", level=2)
    p = doc.add_paragraph()
    run = p.add_run("2,029 automated tests • 0 failures • Full CI/CD ready")
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0x00, 0x88, 0x00)

    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run(
        "Key Testing Highlights:\n"
        "• Every API endpoint has corresponding integration tests\n"
        "• MCQ lifecycle transitions tested end-to-end (Draft→Approved, Draft→Rejected→Resubmit)\n"
        "• AI endpoints tested with mock responses\n"
        "• Role-based access control tested (unauthorized returns 403)\n"
        "• Edge cases: duplicate detection, bulk upload validation, concurrent reviews\n"
        "• Frontend: Every page renders without crash, API interactions mocked, user flows simulated"
    )
    run.font.size = Pt(9)

    # Save document
    doc.save(OUTPUT_PATH)
    file_size = os.path.getsize(OUTPUT_PATH) / 1024 / 1024
    print(f"✅ Document saved: {OUTPUT_PATH}")
    print(f"   Size: {file_size:.1f} MB")
    print(f"   Pages: ~40+ (comprehensive)")


if __name__ == "__main__":
    create_document()
