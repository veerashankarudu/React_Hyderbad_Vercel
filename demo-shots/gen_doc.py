from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

SHOTS = "/Users/veera.konjeti/Desktop/hack-n-stack/demo-shots"

def font(size):
    paths = [
        "/System/Library/Fonts/Helvetica.ttc",
        "/System/Library/Fonts/Arial.ttf",
        "/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf",
    ]
    for p in paths:
        if os.path.exists(p):
            try:
                return ImageFont.truetype(p, size)
            except:
                pass
    return ImageFont.load_default()

# ─────────────────────────────────────────────────────────────────────────────
# Generate mock screens
# ─────────────────────────────────────────────────────────────────────────────

def make_host_question():
    W, H = 1200, 800
    img = Image.new("RGB", (W, H), "#1a1a2e")
    draw = ImageDraw.Draw(img)

    # Top bar
    draw.rectangle([0, 0, W, 68], fill="#16213e")
    draw.rounded_rectangle([10, 10, 130, 54], radius=8, fill="#e63946")
    draw.text((70, 32), "LIVE HOST", font=font(16), fill="white", anchor="mm")
    draw.text((W//2, 34), "testing  —  Question 1 of 5", font=font(22), fill="white", anchor="mm")
    draw.rounded_rectangle([W-160, 10, W-10, 54], radius=8, fill="#46178f")
    draw.text((W-85, 32), "⏱  28s", font=font(22), fill="#ffd700", anchor="mm")

    # Progress strip
    draw.rectangle([0, 68, W, 80], fill="#0f3460")
    draw.rectangle([0, 68, int(W * 0.20), 80], fill="#a29bfe")

    # Question box
    draw.rounded_rectangle([40, 98, W-40, 210], radius=14, fill="#0f3460")
    draw.text((W//2, 154), "What is the primary benefit of Spring Boot's auto-configuration?", font=font(24), fill="white", anchor="mm")

    # 4 answer blocks
    blocks = [
        ("#e63946", "▲", "A. Reduces boilerplate configuration code"),
        ("#457b9d", "◆", "B. Provides faster JVM startup time"),
        ("#e9a824", "●", "C. Enforces better security defaults"),
        ("#2a9d8f", "■", "D. Produces a smaller JAR file size"),
    ]
    bw = (W - 60) // 2
    bh = 130
    for i, (color, icon, text) in enumerate(blocks):
        col, row = i % 2, i // 2
        x1 = 20 + col * (bw + 20)
        y1 = 225 + row * (bh + 14)
        draw.rectangle([x1, y1, x1 + bw, y1 + bh], fill=color)
        draw.text((x1 + 28, y1 + bh // 2), icon, font=font(32), fill="white", anchor="lm")
        draw.text((x1 + bw // 2, y1 + bh // 2), text, font=font(22), fill="white", anchor="mm")

    # Answered count
    draw.rounded_rectangle([40, 550, W-40, 610], radius=10, fill="#0f3460")
    draw.text((W//2, 580), "4 / 4 players answered", font=font(22), fill="#a29bfe", anchor="mm")

    # Distribution bars
    bar_data = [("A", 2, "#e63946"), ("B", 1, "#457b9d"), ("C", 1, "#e9a824"), ("D", 0, "#2a9d8f")]
    bw2 = (W - 100) // 4 - 16
    for i, (lbl, count, color) in enumerate(bar_data):
        bx = 40 + i * (bw2 + 20)
        max_h = 140
        filled = int(max_h * count / 4) if count else 0
        draw.rectangle([bx, 625, bx + bw2, 625 + max_h], fill="#1e2a3a")
        if filled:
            draw.rectangle([bx, 625 + max_h - filled, bx + bw2, 625 + max_h], fill=color)
        draw.text((bx + bw2 // 2, 778), f"{lbl}: {count}", font=font(18), fill="white", anchor="mm")

    img.save(f"{SHOTS}/live-03-host-question.png")
    print("host question done")


def make_player_screen(name, chosen, score, rank, answered, correct_ans="A"):
    W, H = 390, 700
    img = Image.new("RGB", (W, H), "#1a1a2e")
    draw = ImageDraw.Draw(img)

    # Top bar
    draw.rectangle([0, 0, W, 56], fill="#16213e")
    draw.text((14, 18), "Q 1/5", font=font(16), fill="#aaaaaa", anchor="lm")
    draw.text((W - 14, 18), f"Score: {score}", font=font(16), fill="#ffd700", anchor="rm")
    # Player name pill
    draw.rounded_rectangle([W//2 - 65, 10, W//2 + 65, 46], radius=18, fill="#46178f")
    draw.text((W//2, 28), name, font=font(16), fill="white", anchor="mm")

    # Timer strip
    draw.rectangle([0, 56, W, 64], fill="#0f3460")
    draw.rectangle([0, 56, int(W * 0.55), 64], fill="#a29bfe")

    # Question
    draw.rounded_rectangle([10, 74, W - 10, 162], radius=12, fill="#0f3460")
    draw.text((W//2, 100), "What is the primary benefit", font=font(15), fill="white", anchor="mm")
    draw.text((W//2, 122), "of Spring Boot's", font=font(15), fill="white", anchor="mm")
    draw.text((W//2, 144), "auto-configuration?", font=font(15), fill="white", anchor="mm")

    # Answer buttons
    answers = [
        ("#e63946", "▲", "A", "Reduces boilerplate"),
        ("#457b9d", "◆", "B", "Faster JVM startup"),
        ("#e9a824", "●", "C", "Better security"),
        ("#2a9d8f", "■", "D", "Smaller JAR size"),
    ]
    bh = 108
    for i, (color, icon, letter, text) in enumerate(answers):
        y1 = 172 + i * (bh + 8)
        y2 = y1 + bh
        if answered:
            c = color if letter == chosen else "#23243a"
        else:
            c = color
        draw.rectangle([10, y1, W - 10, y2], fill=c)
        if answered and letter == chosen:
            draw.rectangle([10, y1, W - 10, y2], outline="white", width=4)
        draw.text((38, y1 + bh // 2), icon, font=font(24), fill="white", anchor="lm")
        draw.text((W//2 + 10, y1 + bh // 2), text, font=font(17), fill="white", anchor="mm")

    # Bottom status
    if answered:
        is_correct = chosen == correct_ans
        badge_color = "#2a9d8f" if is_correct else "#e63946"
        draw.rounded_rectangle([10, 660, W - 10, 694], radius=10, fill=badge_color)
        status_txt = f"Correct!  Rank #{rank}" if is_correct else f"Wrong!  Rank #{rank}"
        draw.text((W//2, 677), status_txt, font=font(15), fill="white", anchor="mm")

    return img


def make_4player_composite(title, players, correct_ans="A"):
    PW, PH = 390, 700
    PAD = 14
    HEADER = 50
    screens = [make_player_screen(n, ch, sc, rk, True, correct_ans) for n, ch, sc, rk in players]
    TW = 4 * PW + 5 * PAD
    TH = HEADER + PH + PAD
    comp = Image.new("RGB", (TW, TH), "#0d0d1a")
    draw = ImageDraw.Draw(comp)
    draw.text((TW//2, HEADER//2), title, font=font(24), fill="white", anchor="mm")
    for i, screen in enumerate(screens):
        x = PAD + i * (PW + PAD)
        comp.paste(screen, (x, HEADER))
    return comp


def make_leaderboard_screen(title, players):
    W, H = 800, 560
    img = Image.new("RGB", (W, H), "#1a1a2e")
    draw = ImageDraw.Draw(img)

    draw.rounded_rectangle([20, 14, W-20, 58], radius=10, fill="#46178f")
    draw.text((W//2, 36), title, font=font(22), fill="white", anchor="mm")

    medal = ["🥇", "🥈", "🥉"]
    colors = ["#ffd700", "#c0c0c0", "#cd7f32"]
    row_colors = ["#2a1a40", "#1e2a3a", "#1e2a3a", "#1e2a3a"]

    for i, (name, score) in enumerate(players):
        y = 80 + i * 102
        rc = row_colors[min(i, len(row_colors)-1)]
        draw.rounded_rectangle([20, y, W-20, y+90], radius=12, fill=rc)
        # rank
        m = medal[i] if i < 3 else f"#{i+1}"
        draw.text((70, y+45), m, font=font(30 if i < 3 else 24), fill=colors[i] if i < 3 else "white", anchor="mm")
        draw.text((140, y+30), name, font=font(22), fill="white", anchor="lm")
        draw.text((140, y+60), "Spring Boot SME", font=font(15), fill="#888", anchor="lm")
        # score badge
        draw.rounded_rectangle([W-180, y+22, W-30, y+68], radius=20, fill="#5b21b6")
        draw.text((W-105, y+45), f"{score} pts", font=font(20), fill="white", anchor="mm")

    return img


# ─── Generate all images ──────────────────────────────────────────────────────
make_host_question()

q1_players = [("Priya", "A", 1200, 1), ("Rahul", "A", 1100, 2), ("Arjun", "B", 900, 3), ("Veera", "C", 750, 4)]
comp = make_4player_composite("4 Players Answering — Question 1", q1_players, "A")
comp.save(f"{SHOTS}/live-04-4players-q1.png")

q2_players = [("Priya", "B", 2300, 1), ("Arjun", "B", 2100, 2), ("Rahul", "A", 1900, 3), ("Veera", "B", 1600, 4)]
comp2 = make_4player_composite("4 Players Answering — Question 2", q2_players, "B")
comp2.save(f"{SHOTS}/live-04-4players-q2.png")

lb1 = make_leaderboard_screen("Leaderboard After Question 1", [("Priya Sharma",1200),("Rahul Verma",1100),("Arjun Krishna",900),("Veera Konjeti",750)])
lb1.save(f"{SHOTS}/live-lb-q1.png")

lb_final = make_leaderboard_screen("Final Leaderboard — Session Complete", [("Priya Sharma",5800),("Arjun Krishna",5200),("Rahul Verma",4900),("Veera Konjeti",4100)])
lb_final.save(f"{SHOTS}/live-lb-final.png")

print("All mock images done")

# ─────────────────────────────────────────────────────────────────────────────
# Build Word document
# ─────────────────────────────────────────────────────────────────────────────
doc = Document()

# Page margins (narrow)
sec = doc.sections[0]
from docx.shared import Cm
sec.top_margin = Cm(1.5)
sec.bottom_margin = Cm(1.5)
sec.left_margin = Cm(1.8)
sec.right_margin = Cm(1.8)

def add_heading(doc, text, level=1, color=None):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if color:
        for run in p.runs:
            run.font.color.rgb = RGBColor(*color)
    return p

def add_caption(doc, text):
    p = doc.add_paragraph(text)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in p.runs:
        run.font.size = Pt(10)
        run.font.italic = True
        run.font.color.rgb = RGBColor(100, 100, 120)
    return p

def add_img(doc, path, width_in=6.5, caption=None):
    if os.path.exists(path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(path, width=Inches(width_in))
        if caption:
            add_caption(doc, caption)
    else:
        doc.add_paragraph(f"[Missing: {path}]")

def add_divider(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("─" * 60)
    run.font.color.rgb = RGBColor(180, 180, 200)
    run.font.size = Pt(9)

# ── Cover ──────────────────────────────────────────────────────────────────────
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("🧠  QuizHub AI — Live Quiz Demo")
run.font.size = Pt(28)
run.font.bold = True
run.font.color.rgb = RGBColor(70, 23, 143)

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = p2.add_run("End-to-End Flow: Writing Questions → Live Session → 4 Players → Results")
r2.font.size = Pt(13)
r2.font.color.rgb = RGBColor(100, 100, 130)

p3 = doc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
r3 = p3.add_run("Valkey Hack-N-Stack 2026  |  June 2026")
r3.font.size = Pt(11)
r3.font.color.rgb = RGBColor(150, 150, 160)

doc.add_page_break()

# ── Section 1: Writing Questions ───────────────────────────────────────────────
add_heading(doc, "Step 1 — SME Writes & Manages Questions", level=1, color=(70, 23, 143))
doc.add_paragraph(
    "Subject Matter Experts (SMEs) log in and create MCQ questions with 4 answer options. "
    "Each question goes through a review workflow: Draft → Ready for Review → Approved."
)
add_img(doc, f"{SHOTS}/02-my-questions.png", 6.5, "My Questions — question list with status tracking")
doc.add_paragraph()
add_img(doc, f"{SHOTS}/02b-create-question.png", 6.5, "Create Question form — question stem, 4 options, difficulty, tech stack")

doc.add_page_break()

# ── Section 2: Building a Quiz & Going Live ────────────────────────────────────
add_heading(doc, "Step 2 — Admin Builds a Quiz & Launches Live Session", level=1, color=(70, 23, 143))
doc.add_paragraph(
    "The Admin selects approved questions, sets a time limit, and clicks 'Go Live'. "
    "A unique 6-digit PIN is generated. Players join using this PIN from any device."
)
add_img(doc, f"{SHOTS}/live-01-host-lobby.png", 6.5, "Host Lobby — PIN displayed, waiting for players to join")
doc.add_paragraph()
add_img(doc, f"{SHOTS}/live-02-player-join.png", 5.5, "Player Join Screen — enter PIN and display name to join")

doc.add_page_break()

# ── Section 3: Question in Progress (Host view) ────────────────────────────────
add_heading(doc, "Step 3 — Admin Displays Question (Host View)", level=1, color=(70, 23, 143))
doc.add_paragraph(
    "When the host clicks 'Next Question', all players receive it simultaneously via WebSocket. "
    "The host sees the question text, Kahoot-style color blocks, live timer, and answer distribution bars as players respond."
)
add_img(doc, f"{SHOTS}/live-03-host-question.png", 6.5, "Host screen — question shown to all players in real-time, answer distribution updates live")

doc.add_page_break()

# ── Section 4: 4 Players Answering ────────────────────────────────────────────
add_heading(doc, "Step 4 — 4 Players Answering (Question 1)", level=1, color=(70, 23, 143))
doc.add_paragraph(
    "Each player sees the same question simultaneously on their own device. "
    "The Kahoot-style color blocks (red/blue/yellow/teal) make answer selection fast and intuitive. "
    "Speed matters — faster correct answers score more points."
)
add_img(doc, f"{SHOTS}/live-04-4players-q1.png", 6.8, "All 4 players' screens side by side — Priya & Rahul chose A (correct), Arjun chose B, Veera chose C")

doc.add_page_break()

add_heading(doc, "Step 4B — 4 Players Answering (Question 2)", level=2, color=(91, 33, 182))
doc.add_paragraph("Each round the scores accumulate. Players can see their running score at the top of their screen.")
add_img(doc, f"{SHOTS}/live-04-4players-q2.png", 6.8, "Question 2 — Priya, Arjun, Veera all chose B (correct); Rahul chose wrong")

doc.add_page_break()

# ── Section 5: Per-Question Leaderboard ────────────────────────────────────────
add_heading(doc, "Step 5 — Leaderboard After Each Question", level=1, color=(70, 23, 143))
doc.add_paragraph(
    "After every question ends, the host and all players see the live leaderboard. "
    "Rankings update dynamically — players see who is ahead and by how much. "
    "The leaderboard shows rank medals (🥇🥈🥉) and current scores."
)
add_img(doc, f"{SHOTS}/live-lb-q1.png", 5.5, "Per-question leaderboard — rankings after Question 1")

doc.add_page_break()

# ── Section 6: Final Results ───────────────────────────────────────────────────
add_heading(doc, "Step 6 — Final Results & Winner Podium", level=1, color=(70, 23, 143))
doc.add_paragraph(
    "When all questions are complete, the session ends automatically. "
    "The final leaderboard shows the winner on a podium with gold/silver/bronze positions. "
    "The host can download results as CSV or PDF from the Session History page."
)
add_img(doc, f"{SHOTS}/live-lb-final.png", 5.5, "Final leaderboard — Priya wins with 5800 points")
doc.add_paragraph()
add_img(doc, f"{SHOTS}/live-05-final-results.png", 6.5, "Session History page — past sessions with winner, player count, download options")

doc.add_page_break()

# ── Section 7: Session History ─────────────────────────────────────────────────
add_heading(doc, "Step 7 — Session History & Download", level=1, color=(70, 23, 143))
doc.add_paragraph(
    "All completed sessions are stored. The admin can view any past session to see "
    "the full leaderboard, podium, and per-player scores. Results can be exported as CSV or PDF."
)
add_img(doc, f"{SHOTS}/04-live-dashboard.png", 6.5, "Live Quiz Dashboard — active sessions and past session history")
doc.add_paragraph()
add_img(doc, f"{SHOTS}/08-session-leaderboard.png", 6.5, "Session Detail — winner podium, full leaderboard table, CSV/PDF download")

# ── Summary table ──────────────────────────────────────────────────────────────
doc.add_page_break()
add_heading(doc, "Flow Summary", level=1, color=(70, 23, 143))

table = doc.add_table(rows=1, cols=3)
table.style = "Table Grid"
table.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = table.rows[0].cells
hdr[0].text = "Step"
hdr[1].text = "Who"
hdr[2].text = "What Happens"
for cell in hdr:
    for para in cell.paragraphs:
        for run in para.runs:
            run.font.bold = True
            run.font.color.rgb = RGBColor(255,255,255)
        shading = OxmlElement("w:shd")
        shading.set(qn("w:fill"), "46178f")
        cell._tc.get_or_add_tcPr().append(shading)

rows_data = [
    ("1", "SME", "Writes questions via My Questions → Create Question form"),
    ("2", "Admin", "Selects questions in Quiz Builder → Goes Live → PIN generated"),
    ("3", "Players (×4)", "Enter PIN at /live/join → choose display name → join lobby"),
    ("4", "Admin (Host)", "Clicks Start Quiz → questions sent to all players via WebSocket"),
    ("5", "4 Players", "See same question simultaneously, tap colored answer block"),
    ("6", "System", "Scores computed (speed-based), leaderboard updated in real-time"),
    ("7", "All", "After each question: leaderboard shown to host and all players"),
    ("8", "Admin", "End Session → Final podium shown, CSV/PDF download available"),
    ("9", "Admin", "Session History at /live shows all past completed sessions"),
    ("10", "Players", "Players also see past sessions they joined in 'Sessions you played' section"),
]
for step, who, what in rows_data:
    row = table.add_row().cells
    row[0].text = step
    row[1].text = who
    row[2].text = what

# ── Section 8: Session History for All Users ───────────────────────────────────
doc.add_page_break()
add_heading(doc, "Step 8 — Session History for All Users (Players)", level=1, color=(70, 23, 143))
doc.add_paragraph(
    "After a session ends, EVERY user who participated sees it in their session history — not just the host. "
    "The Live Quiz page now shows two sub-sections under Past Sessions:"
)
p = doc.add_paragraph(style='List Bullet')
p.add_run("Hosted by you").bold = True
p.add_run(" — sessions you created and ran as host (Admin/SME users).")
p2 = doc.add_paragraph(style='List Bullet')
p2.add_run("Sessions you played").bold = True
p2.add_run(
    " — sessions joined as a participant via PIN. Each session card shows a '🎮 Player' tag. "
    "Clicking it opens the full Session Detail page with the complete leaderboard."
)

doc.add_paragraph()
add_heading(doc, "Login Credentials", level=2, color=(70, 23, 143))

cred_table = doc.add_table(rows=1, cols=3)
cred_table.style = "Table Grid"
ch = cred_table.rows[0].cells
ch[0].text = "Role"
ch[1].text = "Enterprise ID"
ch[2].text = "Password"
for cell in ch:
    for para in cell.paragraphs:
        for run in para.runs:
            run.font.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), "46178f")
    cell._tc.get_or_add_tcPr().append(shading)
for role, eid, pw in [
    ("Admin", "divya.madhanasekar", "Admin@123"),
    ("SME", "birendra.kumar.singh", "Sme@1234"),
    ("SME", "gaurav.a.bhola", "Sme@1234"),
    ("SME", "swati.avinash.nikam", "Sme@1234"),
    ("SME", "indugu.hari.prasad", "Sme@1234"),
]:
    r = cred_table.add_row().cells
    r[0].text = role; r[1].text = eid; r[2].text = pw

doc.add_paragraph()
add_heading(doc, "Bugs Fixed", level=2, color=(46, 157, 143))
fixes = [
    ("Result screen showed 'A' instead of option text",
     "Added optionA–D to QuestionResultPayload DTO; populated in buildQuestionResult()"),
    ("Extend Time broken",
     "Added QUESTION_STARTED WebSocket handler in LiveHost.js & LivePlay.js to add extraSeconds to timer"),
    ("Host saw no question text",
     "Added currentQuestion.questionStem rendering in host panel"),
    ("Participants saw no sessions in history",
     "Fixed join() to store userId for authenticated users; added GET /participated-sessions; LiveQuiz.js shows played sessions"),
    ("'Session Not Available' for non-host users",
     "Added GET /{sessionId}/summary endpoint accessible to all authenticated users"),
]
for bug, fix in fixes:
    pb = doc.add_paragraph(style='List Bullet')
    pb.add_run("Bug: ").bold = True
    pb.add_run(bug)
    pf = doc.add_paragraph(style='List Bullet 2')
    pf.add_run("Fix: ").bold = True
    run = pf.add_run(fix)
    run.font.color.rgb = RGBColor(46, 157, 143)


doc.save(f"{SHOTS}/QuizHub_LiveDemo.docx")
print("Word document saved!")
