#!/usr/bin/env python3
"""Add KT front-matter then merge Doc1 + Doc2 with images preserved."""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

OUT = "/Users/veera.konjeti/Desktop/hack-n-stack/QuizHub_AI_KT_Document.docx"

def h(doc, text, size=14, color=RGBColor(0x1a,0x73,0xe8), bold=True):
    p = doc.add_paragraph()
    r = p.add_run(text); r.bold=bold; r.font.size=Pt(size); r.font.color.rgb=color
    p.paragraph_format.space_before = Pt(10)
    return p

def body(doc, text, size=10):
    p = doc.add_paragraph()
    r = p.add_run(text); r.font.size=Pt(size)
    return p

def rule(doc):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    b = OxmlElement('w:bottom')
    b.set(qn('w:val'),'single'); b.set(qn('w:sz'),'6')
    b.set(qn('w:space'),'1'); b.set(qn('w:color'),'1a73e8')
    pBdr.append(b); pPr.append(pBdr)

# Build front matter
front = Document()
sec = front.sections[0]
sec.top_margin=Cm(1.5); sec.bottom_margin=Cm(1.5)
sec.left_margin=Cm(2.0); sec.right_margin=Cm(2.0)

# COVER
p = front.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("QuizHub AI"); r.bold=True; r.font.size=Pt(38); r.font.color.rgb=RGBColor(0x1a,0x73,0xe8)

p2 = front.add_paragraph(); p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = p2.add_run("Knowledge Transfer Document"); r2.font.size=Pt(20); r2.font.color.rgb=RGBColor(0x44,0x44,0x44)

front.add_paragraph()
p3 = front.add_paragraph(); p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
r3 = p3.add_run("Valkey Hack-N-Stack 2026  ·  Team Valkey"); r3.font.size=Pt(13); r3.font.color.rgb=RGBColor(0x88,0x88,0x88)

front.add_paragraph()
p4 = front.add_paragraph(); p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
r4 = p4.add_run("445 Features  ·  2,029 Tests (0 failures)  ·  92.5% Backend Coverage  ·  80.37% Frontend Coverage")
r4.font.size=Pt(11); r4.font.color.rgb=RGBColor(0x33,0x33,0x33)
rule(front)
front.add_paragraph()

# 1. TECH STACK
h(front, "1.  Tech Stack")
rows_data = [
    ("Backend",    "Spring Boot 3.2.5  ·  Java 17  ·  Spring Security + JWT  ·  Spring AI  ·  JPA/Hibernate  ·  WebSocket STOMP/SockJS"),
    ("Frontend",   "React 19  ·  JavaScript  ·  Axios  ·  i18next (7 languages)  ·  Recharts"),
    ("Database",   "MySQL 8.x (production)  ·  H2 in-memory (tests — no MySQL needed for tests)"),
    ("AI/Search",  "OpenAI GPT-4o-mini (optional)  ·  Ollama local profile  ·  pgVector embedding search"),
    ("Tests",      "JUnit 5 + Mockito — 1,072 backend  ·  Jest + RTL — 957 frontend  ·  0 failures"),
]
t = front.add_table(rows=len(rows_data), cols=2); t.style='Table Grid'
for i,(k,v) in enumerate(rows_data):
    t.cell(i,0).paragraphs[0].add_run(k).bold=True
    t.cell(i,1).paragraphs[0].add_run(v).font.size=Pt(9.5)
front.add_paragraph()

# 2. ARCHITECTURE
h(front, "2.  Architecture Overview")
body(front, "React 19 (SPA : 3000)  ──REST/WS──►  Spring Boot (API : 8080)  ──JPA──►  MySQL 8.x")
body(front, "                                              └──Spring AI──►  OpenAI / Ollama")
front.add_paragraph()
for item in [
    "All API calls from frontend go through src/api.js (single source). JWT auto-injected via Axios interceptor.",
    "Spring Security validates JWT role on EVERY API endpoint — @PreAuthorize annotations.",
    "WebSocket (STOMP/SockJS) powers Live Quiz Battle — real-time events broadcast to all participants.",
    "19 database entities auto-created by Hibernate. data.sql seeds 6 tech stacks + 5 users on first start.",
    "React PrivateRoute guards ALL frontend pages by role. Backend re-checks role — double protection.",
]:
    p = front.add_paragraph(style="List Bullet"); p.add_run(item).font.size=Pt(10)
front.add_paragraph()

# 3. LOCAL SETUP
h(front, "3.  Local Setup (5 Steps)")
steps = [
    ("Prerequisites",    "Java 17+,  Maven 3.8+,  Node 18+,  MySQL 8.x"),
    ("Create database",  'mysql -u root -e "CREATE DATABASE IF NOT EXISTS quizhub;"'),
    ("Backend install",  "cd backend  &&  mvn clean install -DskipTests"),
    ("Start backend",    "cd backend  &&  mvn spring-boot:run        →  http://localhost:8080"),
    ("Start frontend",   "cd frontend  &&  npm install  &&  npm start  →  http://localhost:3000"),
]
for i,(label,cmd) in enumerate(steps):
    p = front.add_paragraph()
    p.add_run(f"Step {i+1}: {label}  — ").bold = True
    r = p.add_run(cmd); r.font.color.rgb=RGBColor(0x1a,0x73,0xe8); r.font.size=Pt(10)
front.add_paragraph()

# 4. CREDENTIALS
h(front, "4.  Pre-Seeded Login Credentials  (auto-created on first startup)")
ct = front.add_table(rows=6, cols=3); ct.style='Table Grid'
for i,hdr in enumerate(["Role","Enterprise ID","Password"]):
    ct.cell(0,i).paragraphs[0].add_run(hdr).bold=True
creds = [
    ("Admin","divya.madhanasekar","Admin@123"),
    ("Admin","gaurav.a.bhola","Admin@123"),
    ("SME","birendra.kumar.singh","Sme@1234"),
    ("SME","swati.avinash.nikam","Sme@1234"),
    ("SME","indugu.hari.prasad","Sme@1234"),
]
for ri,(role,eid,pwd) in enumerate(creds):
    ct.cell(ri+1,0).paragraphs[0].add_run(role)
    r = ct.cell(ri+1,1).paragraphs[0].add_run(eid)
    r.font.color.rgb=RGBColor(0x1a,0x73,0xe8)
    ct.cell(ri+1,2).paragraphs[0].add_run(pwd)
front.add_paragraph()

# 5. KEY URLs
h(front, "5.  Key URLs")
urls = [
    ("React App",         "http://localhost:3000"),
    ("Spring Boot API",   "http://localhost:8080"),
    ("Swagger API Docs",  "http://localhost:8080/swagger-ui/index.html"),
    ("Health Check",      "http://localhost:8080/actuator/health"),
    ("JaCoCo Coverage",   "backend/target/site/jacoco/index.html  (run: mvn verify)"),
    ("Jest Coverage",     "frontend/coverage/lcov-report/index.html  (run: npm test --coverage)"),
]
ut = front.add_table(rows=len(urls)+1, cols=2); ut.style='Table Grid'
ut.cell(0,0).paragraphs[0].add_run("Page").bold=True
ut.cell(0,1).paragraphs[0].add_run("URL").bold=True
for i,(desc,url) in enumerate(urls):
    ut.cell(i+1,0).paragraphs[0].add_run(desc)
    ut.cell(i+1,1).paragraphs[0].add_run(url).font.color.rgb=RGBColor(0x1a,0x73,0xe8)
front.add_paragraph()

# 6. RUN TESTS
h(front, "6.  Running Tests")
tests = [
    ("Backend Tests",     "cd backend  &&  mvn test",                           "1,072 tests — 0 failures — H2 in-memory (no MySQL)"),
    ("Backend Coverage",  "cd backend  &&  mvn verify",                         "92.5% instruction coverage — JaCoCo report generated"),
    ("Frontend Tests",    "cd frontend  &&  npm test -- --watchAll=false",       "957 tests — 0 failures"),
    ("Frontend Coverage", "cd frontend  &&  npm test -- --watchAll=false --coverage", "80.37% statement coverage — Jest report generated"),
]
tt = front.add_table(rows=len(tests)+1, cols=3); tt.style='Table Grid'
for i,hdr in enumerate(["Suite","Command","Expected"]):
    tt.cell(0,i).paragraphs[0].add_run(hdr).bold=True
for i,(name,cmd,result) in enumerate(tests):
    tt.cell(i+1,0).paragraphs[0].add_run(name)
    tt.cell(i+1,1).paragraphs[0].add_run(cmd).font.color.rgb=RGBColor(0x1a,0x73,0xe8)
    tt.cell(i+1,2).paragraphs[0].add_run(result)
front.add_paragraph()

# 7. PROJECT STRUCTURE
h(front, "7.  Project Structure")
structure = [
    "hack-n-stack/",
    "  backend/src/main/java/          ← Controllers, Services, Entities, Repositories",
    "  backend/src/main/resources/     ← application.yml  +  data.sql (auto-seed on startup)",
    "  backend/src/test/java/          ← 1,072 JUnit 5 tests",
    "  frontend/src/pages/             ← React page components + co-located Jest tests",
    "  frontend/src/components/        ← Reusable UI components (Navbar, Sidebar, Cards...)",
    "  frontend/src/locales/           ← i18n JSON translations (EN, HI, TA, TE, KN, ML, JA)",
    "  frontend/src/api.js             ← ALL Axios API calls in one file",
    "  frontend/src/AuthContext.js     ← JWT + role global context (useAuth hook)",
    "  frontend/src/PrivateRoute.js    ← Role-based route guard component",
]
for line in structure:
    p = front.add_paragraph()
    r = p.add_run(line); r.font.size=Pt(9.5); r.font.name="Courier New"
front.add_paragraph()

# 8. IMPORTANT NOTES
h(front, "8.  Important Notes for New Team Members")
notes = [
    "AI is OPTIONAL — all 445 features work without an OpenAI key (set spring.profiles.active=ollama for local AI).",
    "Email is OPTIONAL — set MAIL_USERNAME + MAIL_PASSWORD env vars to enable email notifications.",
    "All test data auto-seeds on first backend run — no manual SQL needed.",
    "JWT in localStorage (not cookies) — stateless auth, no CSRF vulnerability.",
    "RBAC is double-enforced: React PrivateRoute (frontend) + @PreAuthorize (backend).",
    "Dark mode + language preference persisted in localStorage across sessions.",
    "Live Quiz Battle requires WebSocket connections — check proxy/firewall allows WS.",
    "Service worker cache is intentionally excluded from JS/CSS bundles — new deploys load immediately.",
    "For production deploy: set spring.profiles.active=azure and provide Azure credentials.",
    "Swagger UI at /swagger-ui/index.html shows all 140+ API endpoints with request/response schemas.",
]
for note in notes:
    p = front.add_paragraph(style="List Bullet"); p.add_run(note).font.size=Pt(10)

# SECTION DIVIDER
front.add_page_break()
p_div = front.add_paragraph(); p_div.alignment=WD_ALIGN_PARAGRAPH.CENTER
r_div = p_div.add_run("━━━  FEATURE DOCUMENTATION  (445 Features)  ━━━")
r_div.bold=True; r_div.font.size=Pt(16); r_div.font.color.rgb=RGBColor(0x1a,0x73,0xe8)
front.add_paragraph()
p_sub = front.add_paragraph(); p_sub.alignment=WD_ALIGN_PARAGRAPH.CENTER
r_sub = p_sub.add_run("One page per feature  ·  Annotated screenshot  ·  Workflow order: Register → Login → Create MCQ → Review → Approve/Reject → Notifications → Inbox → AI → Quiz → Leaderboard")
r_sub.font.size=Pt(10); r_sub.font.color.rgb=RGBColor(0x55,0x55,0x55)
front.add_page_break()

# MERGE using docxcompose — properly carries image binary data
FRONT_TMP = "/tmp/quizhub_frontmatter_tmp.docx"
front.save(FRONT_TMP)

from docxcompose.composer import Composer

composer = Composer(Document(FRONT_TMP))
DOC1 = "/Users/veera.konjeti/Desktop/hack-n-stack/QuizHub_Doc1_Features_1to220.docx"
DOC2 = "/Users/veera.konjeti/Desktop/hack-n-stack/QuizHub_Doc2_Features_221to440.docx"
composer.append(Document(DOC1))
print("  ✔ Doc1 appended (features 1-220)")
composer.append(Document(DOC2))
print("  ✔ Doc2 appended (features 221-445)")

composer.save(OUT)
print(f"✅ QuizHub_AI_KT_Document.docx ready — {OUT}")
