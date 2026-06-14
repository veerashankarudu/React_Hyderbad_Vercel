#!/usr/bin/env python3
"""
QuizHub AI — 440 Per-Feature Word Documents
One page per feature. Annotated screenshot + description.
Doc1: features list[:220]  |  Doc2: features list[220:]
"""
import os, io, math
from pathlib import Path
from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

BASE = Path("/Users/veera.konjeti/Desktop/hack-n-stack")
SS   = BASE / "screenshots/fresh"
OUT1 = BASE / "QuizHub_Doc1_Features_1to220.docx"
OUT2 = BASE / "QuizHub_Doc2_Features_221to440.docx"

RED   = (220, 30, 30, 255)
WHITE = (255, 255, 255, 255)

# ─── drawing helpers ─────────────────────────────────────────────────────────
def load_font(size=22):
    for p in ["/System/Library/Fonts/Helvetica.ttc"]:
        if os.path.exists(p):
            return ImageFont.truetype(p, size)
    return ImageFont.load_default()

def thick_arrow(draw, x1, y1, x2, y2, width=9, head=36):
    draw.line([(x1, y1), (x2, y2)], fill=RED, width=width)
    angle = math.atan2(y2 - y1, x2 - x1)
    for side in [+0.45, -0.45]:
        ax = x2 - head * math.cos(angle - side)
        ay = y2 - head * math.sin(angle - side)
        draw.line([(x2, y2), (ax, ay)], fill=RED, width=width)

def tag_box(draw, x, y, text, font):
    bbox = draw.textbbox((0, 0), text, font=font)
    tw, th = bbox[2] - bbox[0], bbox[3] - bbox[1]
    pad = 8
    draw.rectangle([x - pad, y - pad, x + tw + pad, y + th + pad], fill=RED)
    draw.text((x, y), text, fill=WHITE, font=font)

def make_annotated(screenshot_file, feat_num, label_text, ax_pct, ay_pct):
    """Return annotated BytesIO or None."""
    p = SS / screenshot_file
    if not p.exists():
        p = SS / screenshot_file.replace("_mobile", "_desktop")
    if not p.exists():
        return None
    img = Image.open(p).convert("RGBA")
    W, H = img.size
    ov  = Image.new("RGBA", (W, H), (0, 0, 0, 0))
    d   = ImageDraw.Draw(ov)
    f   = load_font(22)

    ax = int(W * ax_pct)
    ay = int(H * ay_pct)

    text = f"#{feat_num}: {label_text}"
    bbox = d.textbbox((0, 0), text, font=f)
    tw = bbox[2] - bbox[0];  th = bbox[3] - bbox[1]
    pad = 8;  lw = tw + 2*pad;  lh = th + 2*pad

    # Place label in opposite quadrant from arrow tip
    if ax_pct > 0.5:
        lx = max(pad + 5, 8)
    else:
        lx = min(W - lw - pad - 8, int(W * 0.55))

    if ay_pct > 0.5:
        ly = max(pad + 5, 8)
    else:
        ly = min(H - lh - pad - 20, int(H * 0.80))

    # Arrow from nearest label edge → arrow tip
    lcx = lx + lw // 2;  lcy = ly + lh // 2
    ex  = (lx + lw) if ax > lcx else lx
    ey  = (ly + lh) if ay > lcy else ly

    thick_arrow(d, ex, ey, ax, ay)
    tag_box(d, lx, ly, text, f)

    out = Image.alpha_composite(img, ov).convert("RGB")
    buf = io.BytesIO()
    out.save(buf, "PNG")
    buf.seek(0)
    return buf

# ─── Word doc builder ─────────────────────────────────────────────────────────
def set_margins(doc):
    sec = doc.sections[0]
    sec.top_margin    = Cm(1.2)
    sec.bottom_margin = Cm(1.0)
    sec.left_margin   = Cm(1.5)
    sec.right_margin  = Cm(1.5)

def build_doc(features, output_path, doc_num):
    doc = Document()
    set_margins(doc)

    # Cover
    cp = doc.add_paragraph()
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cp.add_run("QuizHub AI — Feature Documentation")
    r.bold = True;  r.font.size = Pt(22)
    r.font.color.rgb = RGBColor(0x1a, 0x73, 0xe8)

    cp2 = doc.add_paragraph()
    cp2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = cp2.add_run(f"Document {doc_num}  |  {len(features)} Features")
    r2.font.size = Pt(14);  r2.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    doc.add_paragraph()
    cp3 = doc.add_paragraph()
    cp3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = cp3.add_run("Accenture Hack-N-Stack 2026  ·  Team BumbleBee")
    r3.font.size = Pt(11);  r3.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    doc.add_page_break()

    total = len(features)
    for i, feat in enumerate(features):
        feat_num, short_label, full_desc, screenshot, ax_pct, ay_pct = feat
        print(f"  [{i+1:>3}/{total}] #{feat_num}: {short_label}")

        # Heading
        hp = doc.add_paragraph()
        hr = hp.add_run(f"Feature {feat_num} — {short_label}")
        hr.bold = True;  hr.font.size = Pt(14)
        hr.font.color.rgb = RGBColor(0x1a, 0x73, 0xe8)
        hp.paragraph_format.space_after = Pt(4)

        # Annotated screenshot
        buf = make_annotated(screenshot, feat_num, short_label, ax_pct, ay_pct)
        if buf:
            ip = doc.add_paragraph()
            ip.alignment = WD_ALIGN_PARAGRAPH.CENTER
            ip.paragraph_format.space_before = Pt(2)
            ip.paragraph_format.space_after  = Pt(4)
            ip.add_run().add_picture(buf, width=Inches(5.5))
        else:
            doc.add_paragraph(f"[Screenshot not found: {screenshot}]")

        # Description
        dp = doc.add_paragraph()
        dr = dp.add_run(f"→  {full_desc}")
        dr.font.size = Pt(10)
        dr.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
        dp.paragraph_format.space_after = Pt(2)

        if i < total - 1:
            doc.add_page_break()

    doc.save(str(output_path))
    print(f"\n✅ Saved: {output_path}  ({len(features)} pages)")

# ═══════════════════════════════════════════════════════════════════════════════
# ALL FEATURES — workflow order: register → login → create MCQ → review → …
# (feat_num, short_label, full_description, screenshot_file, arrow_x%, arrow_y%)
# ═══════════════════════════════════════════════════════════════════════════════
FEATURES = [

    # ── REGISTER FIRST (workflow starts here) ─────────────────────────────────
    (15, "Register Page",         "Register page with all fields (enterprise ID, full name, email, password, tech stacks)", "03_register_desktop.png", 0.30, 0.30),
    (16, "Password Masking",      "Password masking on input with toggle show/hide visibility", "03_register_desktop.png", 0.47, 0.52),
    (17, "Pending on Submit",     "Submit → account in PENDING status (cannot login yet)", "03_register_desktop.png", 0.50, 0.75),
    (18, "Admin Approval Flow",   "Admin must approve account before the user can login", "03_register_desktop.png", 0.50, 0.80),
    (19, "Duplicate ID Block",    "Duplicate enterprise ID → rejected with clear error message", "03_register_desktop.png", 0.30, 0.28),
    (20, "Password Policy",       "Weak password policy enforcement — complexity rules applied", "03_register_desktop.png", 0.30, 0.57),
    (21, "Change Password Modal", "Change password modal accessible from navbar profile menu", "08_profile_menu_desktop.png", 0.85, 0.40),
    (22, "Wrong Password Error",  "Wrong current password → validation error shown immediately", "03_register_desktop.png", 0.30, 0.60),
    (23, "Password Mismatch",     "Password mismatch (new ≠ confirm) → error shown inline", "03_register_desktop.png", 0.30, 0.65),
    (24, "Password Changed OK",   "Correct change password flow → session maintained after update", "08_profile_menu_desktop.png", 0.85, 0.45),

    # ── LOGIN ─────────────────────────────────────────────────────────────────
    (1,  "Enterprise ID Login",   "Login page with enterprise ID + password fields", "01_login_desktop.png", 0.07, 0.22),
    (2,  "1-Click Demo Autofill", "Demo panel with 5 one-click login users (2 Admin, 3 SME)", "01_login_desktop.png", 0.18, 0.65),
    (3,  "JWT Authentication",    "JWT-based authentication — stateless, no server sessions", "01_login_desktop.png", 0.30, 0.52),
    (4,  "JWT localStorage",      "JWT stored in localStorage with auto-injection via Axios interceptor", "01_login_desktop.png", 0.30, 0.48),
    (5,  "Admin Full Sidebar",    "Admin login → full sidebar showing all pages", "05_admin_dashboard_desktop.png", 0.05, 0.40),
    (6,  "SME Restricted Sidebar","SME login → restricted sidebar (no admin pages visible)", "76_sme_dashboard_desktop.png", 0.05, 0.35),
    (7,  "Invalid Credentials",   "Wrong password → 'Invalid credentials' error message shown", "02_login_error_desktop.png", 0.30, 0.47),
    (8,  "Empty Field Validation","Empty fields → client-side validation fires, no API call made", "01_login_desktop.png", 0.30, 0.55),
    (9,  "Forgot Password Page",  "Forgot password page with enterprise ID / email input field", "04_forgot_password_desktop.png", 0.30, 0.40),
    (10, "Password Reset Email",  "Password reset email flow: SMTP + DB-stored expiring token (one-time use)", "04_forgot_password_desktop.png", 0.30, 0.50),
    (11, "Reset Token Validation","Reset token validation — expired/used token shows clear error", "04_forgot_password_desktop.png", 0.30, 0.55),
    (12, "Login Rate Limiting",   "Login rate limiting — 10 attempts per IP per 60s, HTTP 429 response", "01_login_desktop.png", 0.30, 0.38),
    (13, "Concurrent Sessions",   "Concurrent session support — multiple devices allowed", "01_login_desktop.png", 0.30, 0.35),
    (14, "Logout Clears Session", "Logout clears JWT + localStorage, redirects to login page", "08_profile_menu_desktop.png", 0.85, 0.60),

    # ── DASHBOARD ─────────────────────────────────────────────────────────────
    (25, "Live Stat Cards",       "Stat cards showing live DB data (Total MCQs, Approved, Under Review, Rejected)", "05_admin_dashboard_desktop.png", 0.21, 0.21),
    (26, "Dark Mode Toggle",      "Dark mode / Light mode toggle — preference persisted in localStorage", "06_dark_mode_desktop.png", 0.73, 0.03),
    (27, "Language Switcher",     "Language switcher with 7 locales and flag icons", "07_language_switcher_desktop.png", 0.76, 0.03),
    (28, "IST Timestamps",        "UTC → IST timestamp display with '2h ago' relative format", "05_admin_dashboard_desktop.png", 0.70, 0.65),
    (29, "Recent Activity Table", "Recent activity table showing latest MCQ status updates", "05_admin_dashboard_desktop.png", 0.50, 0.65),
    (30, "Top Reviewers Widget",  "Leaderboard widget on dashboard showing top reviewers", "05_admin_dashboard_desktop.png", 0.90, 0.68),
    (31, "SME Own Stats Only",    "SME sees only their own stats — not system-wide counts", "76_sme_dashboard_desktop.png", 0.30, 0.20),
    (32, "Admin System Stats",    "Admin sees system-wide counts across all users", "05_admin_dashboard_desktop.png", 0.55, 0.20),
    (33, "Mobile Dashboard",      "Mobile-responsive dashboard — hamburger menu + stacked cards", "05_admin_dashboard_mobile.png", 0.05, 0.10),
    (34, "App Branding",          "App logo and 'QuizHub AI' name visible in topbar", "05_admin_dashboard_desktop.png", 0.07, 0.04),

    # ── MCQ CREATE / EDIT ─────────────────────────────────────────────────────
    (35, "Question Stem Input",   "MCQ form with multiline question stem textarea", "14_create_mcq_blank_desktop.png", 0.50, 0.22),
    (36, "4 Answer Options",      "4 answer option inputs (A, B, C, D) with text fields", "14_create_mcq_blank_desktop.png", 0.50, 0.45),
    (37, "Correct Answer Radio",  "Radio selector to mark one option as the correct answer", "14_create_mcq_blank_desktop.png", 0.25, 0.42),
    (38, "Tech Stack Dropdown",   "Subject/Tech Stack dropdown linked to Master Data", "14_create_mcq_blank_desktop.png", 0.25, 0.30),
    (39, "Dynamic Topic Dropdown","Topic dropdown — dynamically populated from selected tech stack", "14_create_mcq_blank_desktop.png", 0.50, 0.30),
    (40, "Difficulty Selector",   "Difficulty selector: Easy / Medium / Hard", "14_create_mcq_blank_desktop.png", 0.75, 0.30),
    (41, "Bloom's Taxonomy",      "Bloom's Taxonomy level selector — 6 levels (Remember→Create)", "14_create_mcq_blank_desktop.png", 0.85, 0.30),
    (42, "Code Block Button",     "Code Block button — inserts formatted code block into question", "15_create_mcq_filled_desktop.png", 0.80, 0.22),
    (43, "XSS Protection",        "Rich text renders safely — XSS prevented via QuestionStemRenderer + DOMPurify", "15_create_mcq_filled_desktop.png", 0.50, 0.20),
    (44, "AI Generate MCQ",       "AI-assisted MCQ generation — 🤖 Generate with AI button", "17_mcq_form_ai_section_desktop.png", 0.88, 0.22),
    (45, "Save as DRAFT",         "Save as Draft → MCQ saved in DRAFT status", "14_create_mcq_blank_desktop.png", 0.40, 0.88),
    (46, "Send for Review",       "Save & Send for Review → status becomes READY_FOR_REVIEW", "14_create_mcq_blank_desktop.png", 0.60, 0.88),
    (47, "Edit MCQ Pre-filled",   "Edit draft → form pre-filled with all existing MCQ data", "20_edit_mcq_desktop.png", 0.50, 0.30),
    (48, "Delete Draft MCQ",      "Delete draft → MCQ removed with success toast notification", "20_edit_mcq_desktop.png", 0.85, 0.88),
    (49, "Empty Stem Validation", "Empty question stem → validation error, submit blocked", "16_mcq_form_validation_desktop.png", 0.50, 0.25),
    (50, "No Answer Validation",  "No correct answer selected → validation error shown", "16_mcq_form_validation_desktop.png", 0.25, 0.42),

    # ── MY QUESTIONS ──────────────────────────────────────────────────────────
    (51, "Paginated MCQ Table",   "Paginated table: Question, Tech Stack, Topic, Difficulty, Status, Actions", "10_my_questions_all_desktop.png", 0.40, 0.35),
    (52, "Status Filter Tabs",    "Status filter tabs (All/Draft/Ready/Under Review/Approved/Rejected) with counts", "10_my_questions_all_desktop.png", 0.30, 0.14),
    (53, "Real-time Search",      "Real-time search/filter across all questions instantly", "12_my_questions_search_desktop.png", 0.70, 0.14),
    (54, "Column Sort",           "Column sort ascending/descending with visual arrow indicator", "10_my_questions_all_desktop.png", 0.40, 0.24),
    (55, "Page Size Config",      "Pagination with configurable page size: 5 / 10 / 15 / 20", "10_my_questions_all_desktop.png", 0.50, 0.90),
    (56, "Edit DRAFT/REJECTED",   "Edit button shown only for DRAFT and REJECTED MCQs", "11_my_questions_draft_desktop.png", 0.85, 0.35),
    (57, "Resubmit Rejected",     "Resubmit REJECTED MCQ → status back to READY_FOR_REVIEW", "11_my_questions_rejected_desktop.png", 0.85, 0.35),
    (58, "Empty State",           "New SME with no questions → empty state 'No questions yet'", "10_my_questions_all_desktop.png", 0.50, 0.55),
    (59, "Export CSV",            "Export filtered question list to CSV file", "10_my_questions_all_desktop.png", 0.88, 0.14),
    (60, "Export Excel",          "Export filtered question list to Excel (.xlsx) file", "10_my_questions_all_desktop.png", 0.92, 0.14),
    (61, "View Full Question",    "View Full Question link per row — opens MCQ Detail page", "10_my_questions_all_desktop.png", 0.78, 0.35),
    (62, "Status Colour Badges",  "Status badges: DRAFT=grey, READY=blue, UNDER_REVIEW=yellow, APPROVED=green, REJECTED=red", "10_my_questions_all_desktop.png", 0.70, 0.35),

    # ── MCQ DETAIL ────────────────────────────────────────────────────────────
    (63, "Full Detail View",      "Full detail: question stem, all 4 options, metadata (tech stack, topic, difficulty, creator, dates)", "18_mcq_detail_desktop.png", 0.50, 0.20),
    (64, "Correct Answer Green",  "Correct answer highlighted green for admin/reviewer viewing", "18_mcq_detail_desktop.png", 0.35, 0.38),
    (65, "Reviewer Feedback",     "Reviewer feedback panel shown for REJECTED MCQs with reason", "18_mcq_detail_desktop.png", 0.80, 0.45),
    (66, "Comment Thread",        "Discussion comment thread with threaded replies (McqCommentSection)", "18_mcq_detail_desktop.png", 0.50, 0.70),
    (67, "Post Comment",          "Post comment → visible immediately with timestamp + author avatar", "18_mcq_detail_desktop.png", 0.50, 0.82),
    (68, "Chronological Order",   "Comments in chronological order with @mention support", "18_mcq_detail_desktop.png", 0.50, 0.74),
    (69, "Delete Own Comment",    "User can delete their own comment from the thread", "18_mcq_detail_desktop.png", 0.80, 0.74),
    (70, "Back Navigation",       "Back navigation returns to correct referring page (breadcrumb)", "18_mcq_detail_desktop.png", 0.05, 0.10),
    (71, "Print/PDF Export",      "Print/PDF export button on MCQ detail page", "18_mcq_detail_desktop.png", 0.90, 0.10),
    (72, "Status on Detail",      "MCQ status badge clearly visible on detail page", "18_mcq_detail_desktop.png", 0.85, 0.15),
    (73, "IST Comment Timestamp", "IST timestamps shown on every comment", "18_mcq_detail_desktop.png", 0.70, 0.74),
    (74, "Rich Text Renders",     "Rich text question renders correctly — code blocks and formatting preserved", "18_mcq_detail_desktop.png", 0.50, 0.15),

    # ── PENDING REVIEWS (workflow: MCQ assigned → reviewer sees it) ───────────
    (75, "Assigned Only Shown",   "Only assigned UNDER_REVIEW questions shown to this reviewer", "25_pending_reviews_desktop.png", 0.35, 0.25),
    (76, "Pending Status Badge",  "'Pending' label on each review card with status badge", "25_pending_reviews_desktop.png", 0.75, 0.25),
    (77, "4-Checkbox Checklist",  "Pre-submission checklist — 4 checkboxes, all must be ticked before submitting verdict", "25_pending_reviews_desktop.png", 0.35, 0.45),
    (78, "Approve → APPROVED",   "Approve action → MCQ becomes APPROVED and leaves reviewer list", "25_pending_reviews_desktop.png", 0.70, 0.65),
    (79, "Reject + Comment",      "Reject with mandatory comment → REJECTED status, reason stored", "25_pending_reviews_desktop.png", 0.80, 0.65),
    (80, "Comment No Verdict",    "Post comment without verdict → question stays UNDER_REVIEW", "25_pending_reviews_desktop.png", 0.60, 0.72),
    (81, "Reviewer Isolation",    "Reviewer A and B see only their own assigned questions (isolated)", "25_pending_reviews_desktop.png", 0.35, 0.20),
    (82, "All Caught Up State",   "No reviews assigned → empty state 'All caught up!' displayed", "25_pending_reviews_desktop.png", 0.50, 0.50),
    (83, "Navbar Pending Badge",  "Navbar badge shows live pending review count", "25_pending_reviews_desktop.png", 0.82, 0.03),
    (84, "SME Auto Notified",     "SME automatically notified (bell + email) when reviewer submits decision", "74_notifications_panel_desktop.png", 0.50, 0.45),

    # ── QUESTION BANK (admin view) ────────────────────────────────────────────
    (85, "All MCQs Visible",      "Admin sees all MCQs from all users — full platform view, paginated", "26_question_bank_desktop.png", 0.50, 0.40),
    (86, "Tech Stack Filter",     "Subject/Tech Stack filter dropdown in question bank", "27_question_bank_filters_desktop.png", 0.30, 0.12),
    (87, "Status Filter Dropdown","Status filter dropdown — filter by any MCQ status", "27_question_bank_filters_desktop.png", 0.55, 0.12),
    (88, "Semantic Search",       "Semantic search by keyword — embedding-based relevance matching", "26_question_bank_desktop.png", 0.40, 0.12),
    (89, "Export Filtered CSV",   "Export filtered results to CSV or Excel", "26_question_bank_desktop.png", 0.90, 0.12),
    (90, "Assign Reviewer Btn",   "Assign Reviewer button — only shown for READY_FOR_REVIEW MCQs", "26_question_bank_desktop.png", 0.85, 0.35),
    (91, "Assign Dialog Info",    "Assign dialog shows: Tech Stack, Topic, Creator ID, reviewer dropdown", "28_assign_reviewer_dialog_desktop.png", 0.50, 0.35),
    (92, "Stack-Filtered Dropdown","Reviewer dropdown filtered by tech stack — excludes MCQ creator", "28_assign_reviewer_dialog_desktop.png", 0.50, 0.55),
    (93, "Admin Can Review",      "Admin can be assigned as a reviewer for any MCQ", "28_assign_reviewer_dialog_desktop.png", 0.50, 0.60),
    (94, "Assign → UNDER_REVIEW", "Assign reviewer → MCQ status changes to UNDER_REVIEW", "28_assign_reviewer_dialog_desktop.png", 0.70, 0.75),
    (95, "Bulk Checkbox Select",  "Bulk checkbox select + bulk actions on multiple MCQs", "26_question_bank_desktop.png", 0.10, 0.35),
    (96, "Admin Edit Any Status", "Admin can edit any MCQ regardless of current status", "26_question_bank_desktop.png", 0.88, 0.40),

    # ── BULK UPLOAD ───────────────────────────────────────────────────────────
    (97,  "Bulk Upload Page",     "Bulk upload page with drag-and-drop zone for CSV/XLSX files", "23_bulk_upload_desktop.png", 0.50, 0.45),
    (98,  "Download Template",    "Download blank Template_MCQs.xlsx template button", "23_bulk_upload_desktop.png", 0.25, 0.20),
    (99,  "Upload Drop Zone",     "Drag file onto drop zone or click to browse and upload", "24_bulk_upload_dropzone_desktop.png", 0.50, 0.45),
    (100, "Preview Table",        "Preview table shows all parsed rows before saving", "23_bulk_upload_desktop.png", 0.50, 0.60),
    (101, "Row Validation",       "Validates each row: tech stack, topic, difficulty, stem, 4 options, correct answer", "23_bulk_upload_desktop.png", 0.75, 0.65),
    (102, "Valid Rows Saved",     "Valid rows saved as DRAFT in My Questions page", "23_bulk_upload_desktop.png", 0.50, 0.70),
    (103, "Partial File Errors",  "Partial file: valid rows saved, invalid rows listed with per-row error detail", "23_bulk_upload_desktop.png", 0.80, 0.65),
    (104, "Wrong File Rejected",  "Wrong file type (.pdf, .jpg etc.) → rejected with clear error message", "23_bulk_upload_desktop.png", 0.50, 0.50),
    (105, "Empty File Handled",   "Empty file handled gracefully — upload progress bar shown", "23_bulk_upload_desktop.png", 0.50, 0.48),

    # ── USER MANAGEMENT ───────────────────────────────────────────────────────
    (106, "User Table",           "User table with roles, status, tech stack assignment info (paginated)", "29_user_management_desktop.png", 0.50, 0.30),
    (107, "Approve Pending User", "Approve pending user → account activated, user can now login", "29_user_management_desktop.png", 0.75, 0.35),
    (108, "Reject Registration",  "Reject user registration → account permanently blocked", "29_user_management_desktop.png", 0.85, 0.35),
    (109, "Change Role SME↔Admin","Change role SME ↔ ADMIN takes effect immediately", "29_user_management_desktop.png", 0.60, 0.35),
    (110, "Search Users",         "Search users by name or enterprise ID in the user table", "29_user_management_desktop.png", 0.30, 0.12),
    (111, "Deactivate User",      "Deactivate an active user — blocks login without deletion", "29_user_management_desktop.png", 0.88, 0.40),
    (112, "Self-Protection",      "Cannot delete or deactivate your own account", "29_user_management_desktop.png", 0.80, 0.40),
    (113, "User Count Matches",   "User count matches live dashboard stat cards", "29_user_management_desktop.png", 0.25, 0.15),
    (114, "Master Data Page",     "Master Data page — manage tech stacks and topics", "30_master_data_desktop.png", 0.50, 0.30),
    (115, "Add Tech Stack",       "Add new tech stack → immediately available in MCQ form dropdown", "32_master_data_add_stack_desktop.png", 0.50, 0.45),
    (116, "Edit Stack Name",      "Edit existing tech stack/subject name", "30_master_data_desktop.png", 0.80, 0.35),
    (117, "Delete with Check",    "Delete tech stack — blocked if MCQs are linked (dependency check)", "30_master_data_desktop.png", 0.85, 0.35),
    (118, "Add Topic",            "Add new topic under a tech stack", "31_master_data_topics_desktop.png", 0.80, 0.55),
    (119, "Duplicate Name Block", "Duplicate tech stack name → rejected with error", "30_master_data_desktop.png", 0.50, 0.45),
    (120, "SME RBAC Block",       "SME cannot access /master-data — RBAC enforced (redirected)", "82_sme_rbac_blocked_user_mgmt_desktop.png", 0.50, 0.50),
    (121, "Spring Cache Sync",    "Dropdown data syncs instantly via Spring Cache + @CacheEvict", "30_master_data_desktop.png", 0.50, 0.20),

    # ── ANALYTICS + KANBAN ────────────────────────────────────────────────────
    (122, "Analytics Charts",     "Analytics dashboard with donut chart + bar chart", "33_analytics_desktop.png", 0.40, 0.40),
    (123, "Date Range Filter",    "Date range filter changes chart data dynamically", "34_analytics_charts_desktop.png", 0.85, 0.15),
    (124, "Export Analytics",     "Export analytics report as Excel or print to PDF", "33_analytics_desktop.png", 0.92, 0.15),
    (125, "SME Own Analytics",    "SME sees only their own data — not platform-wide stats", "33_analytics_desktop.png", 0.30, 0.15),
    (126, "Reviewer Performance", "Reviewer performance chart — approval rate and review count", "34_analytics_charts_desktop.png", 0.60, 0.65),
    (127, "Approval Rate %",      "Approval rate % calculated per reviewer and displayed", "34_analytics_charts_desktop.png", 0.65, 0.70),
    (128, "5-Column Kanban",      "5-column Kanban board: DRAFT / READY / UNDER_REVIEW / APPROVED / REJECTED", "21_kanban_board_desktop.png", 0.50, 0.15),
    (129, "Kanban User Filter",   "SME sees only own questions in Kanban; Admin sees all", "21_kanban_board_desktop.png", 0.05, 0.12),
    (130, "Card Click → Detail",  "Click any Kanban card → opens MCQ Detail page", "21_kanban_board_desktop.png", 0.25, 0.35),
    (131, "Live Column Counts",   "Column card counts are live — update on status changes", "21_kanban_board_desktop.png", 0.10, 0.15),
    (132, "Kanban Stack Filter",  "Filter Kanban by tech stack + keyword search bar", "22_kanban_board_scroll_desktop.png", 0.70, 0.10),

    # ── KANBAN VISUAL ─────────────────────────────────────────────────────────
    ("K1", "DRAFT Column",        "Kanban DRAFT column — all draft MCQs shown as cards with count badge", "21_kanban_board_desktop.png", 0.10, 0.35),
    ("K2", "READY Column",        "Kanban READY column — MCQs ready for review assignment", "21_kanban_board_desktop.png", 0.28, 0.35),
    ("K3", "APPROVED Column",     "Kanban APPROVED column — approved and published MCQs", "21_kanban_board_desktop.png", 0.64, 0.35),
    ("K4", "Kanban Search Bar",   "Search bar at top filters cards across ALL columns simultaneously", "22_kanban_board_scroll_desktop.png", 0.50, 0.10),
    ("K5", "Kanban Tech Stack",   "Tech Stack dropdown narrows cards to selected stack only", "22_kanban_board_scroll_desktop.png", 0.72, 0.10),

    # ── QUIZ BUILDER + PROCTORED ──────────────────────────────────────────────
    (133, "Quiz Builder Page",    "Quiz Builder page — create and manage proctored assessment sessions", "44_quiz_builder_desktop.png", 0.50, 0.20),
    (134, "Create Quiz Form",     "Create quiz from approved MCQs with filters: tech stack, difficulty, count, time limit", "45_quiz_builder_create_form_desktop.png", 0.50, 0.35),
    (135, "Quiz Attempts History","Quiz attempts history page — all past attempts per session", "46_quiz_list_desktop.png", 0.50, 0.55),
    (136, "Quiz Landing Page",    "Quiz landing page with name + email entry and 'Continue →' start button", "46_quiz_list_desktop.png", 0.50, 0.35),
    (137, "Countdown Timer",      "Countdown timer: purple → orange at 5 min → red at 1 min remaining", "44_quiz_builder_desktop.png", 0.70, 0.40),
    (138, "Timer Auto-Submit",    "Timer expires → quiz auto-submitted with current answers", "44_quiz_builder_desktop.png", 0.75, 0.45),
    (139, "Tab Switch Warning",   "Tab switch → violation warning toast + screenshot captured via html2canvas", "44_quiz_builder_desktop.png", 0.50, 0.50),
    (140, "Fullscreen Exit",      "Fullscreen exit → violation counted via fullscreenchange event", "44_quiz_builder_desktop.png", 0.80, 0.40),
    (141, "Copy-Paste Disabled",  "Copy-paste disabled during quiz — clipboard events blocked", "44_quiz_builder_desktop.png", 0.50, 0.45),
    (142, "Score Displayed",      "Submit → score displayed immediately with detailed results per question", "46_quiz_list_desktop.png", 0.75, 0.55),
    (143, "3 Strikes = TERMINATED","3 tab-switch strikes → auto-submit with TERMINATED status", "46_quiz_list_desktop.png", 0.75, 0.50),
    (144, "Non-Registered User",  "Non-registered users can take quiz — name + email entry, one attempt per email", "46_quiz_list_desktop.png", 0.50, 0.45),
    (145, "Quiz Link Expiry",     "Quiz link expiry configurable in hours — enforced on attempt start", "46_quiz_list_desktop.png", 0.65, 0.40),
    (146, "Exam Lock Guard",      "Exam lock guard — blocks opening app in 2nd tab via sessionStorage", "44_quiz_builder_desktop.png", 0.50, 0.55),

    # ── LIVE QUIZ BATTLE Part 1 ───────────────────────────────────────────────
    (147, "Live Battle Hub",      "Live Quiz Battle page — host real-time multiplayer quiz sessions", "61_live_quiz_home_desktop.png", 0.50, 0.12),
    (148, "6-Digit PIN",          "Generate unique 6-digit PIN code for participants to join session", "61_live_quiz_home_desktop.png", 0.35, 0.45),
    (149, "Join a Game",          "'Join a Game' button — participants enter PIN code to join", "65_live_quiz_join_desktop.png", 0.65, 0.45),
    (150, "Active Sessions List", "Active sessions list showing game codes and live participant count", "61_live_quiz_home_desktop.png", 0.50, 0.60),
    (151, "Past Sessions History","Past sessions history tabs — hosted and participated sessions", "61_live_quiz_home_desktop.png", 0.50, 0.70),
    (152, "Live Leaderboard",     "Real-time competition with live leaderboard updating per answer", "61_live_quiz_home_desktop.png", 0.80, 0.50),
    (153, "Create Live Quiz Form","Create Live Quiz form: title, tech stack, topic, difficulty, question count, time per question", "64_live_quiz_create_form_desktop.png", 0.50, 0.35),
    (154, "BATTLE Mode Select",   "Session Mode selector — BATTLE mode for competitive play", "64_live_quiz_create_form_desktop.png", 0.35, 0.54),
    (155, "Team Mode Toggle",     "Team Mode toggle — enable team-based competition", "64_live_quiz_create_form_desktop.png", 0.65, 0.54),
    (156, "Adaptive Difficulty",  "Adaptive Difficulty — question difficulty adjusts dynamically per player", "64_live_quiz_create_form_desktop.png", 0.35, 0.62),
    (157, "Recording Toggle",     "Recording enabled — session recorded for later replay", "64_live_quiz_create_form_desktop.png", 0.65, 0.62),
    (158, "Co-host Assign",       "Co-host support — assign co-host by enterprise ID", "64_live_quiz_create_form_desktop.png", 0.50, 0.70),
    (159, "WebSocket STOMP",      "WebSocket real-time — STOMP over SockJS for instant event broadcast", "61_live_quiz_home_desktop.png", 0.50, 0.35),
    (160, "Host Controls Panel",  "Live Host Controls — start, pause, resume, extend time, end session, next question", "61_live_quiz_home_desktop.png", 0.80, 0.35),
    (161, "Kick Participant",     "Kick participant — host can remove a player mid-session", "61_live_quiz_home_desktop.png", 0.85, 0.55),
    (162, "Transfer Host",        "Transfer host role to another participant mid-session", "61_live_quiz_home_desktop.png", 0.75, 0.55),
    (163, "Live Lobby",           "Live Lobby — waiting room showing joined participants before start", "61_live_quiz_home_desktop.png", 0.50, 0.40),
    (164, "3-2-1 Countdown",      "3-2-1 animated countdown before each question is revealed", "61_live_quiz_home_desktop.png", 0.50, 0.50),
    (165, "Keyboard Answer Keys", "Live Play — participants answer with keyboard shortcuts A/B/C/D", "61_live_quiz_home_desktop.png", 0.50, 0.65),

    # ── LIVE QUIZ BATTLE Part 2 ───────────────────────────────────────────────
    (166, "Score Formula",        "Score = base points (Easy 1000 / Medium 1500 / Hard 2000) + time bonus", "64_live_quiz_create_form_desktop.png", 0.80, 0.40),
    (167, "Podium + Medals",      "Live Results page — animated podium (top 3 with gold/silver/bronze medals)", "50_leaderboard_live_battle_desktop.png", 0.50, 0.30),
    (168, "Certificate Gen",      "Downloadable certificate for top 3 + participation certificate for all", "50_leaderboard_live_battle_desktop.png", 0.80, 0.70),
    (169, "Invite Link",          "Shareable invite link generated per session", "64_live_quiz_create_form_desktop.png", 0.80, 0.60),
    (170, "Session Replay",       "Session replay — replay any past session question-by-question", "61_live_quiz_home_desktop.png", 0.50, 0.72),
    (171, "Team Leaderboard",     "Separate team leaderboard for team mode sessions", "50_leaderboard_live_battle_desktop.png", 0.50, 0.60),
    (172, "Session State Restore","Session state restore — host/players can reconnect and resume", "61_live_quiz_home_desktop.png", 0.50, 0.75),
    (173, "Host Reconnect",       "Host can recover full session control after accidental page refresh", "61_live_quiz_home_desktop.png", 0.80, 0.40),
    (174, "Session Detail Page",  "Live session detail page — full summary + leaderboard for ended sessions", "50_leaderboard_live_battle_desktop.png", 0.50, 0.40),
    (175, "Participated Sessions","View history of all sessions you participated in", "61_live_quiz_home_desktop.png", 0.50, 0.70),
    (176, "PIN Validation",       "PIN validation endpoint — verifies PIN validity before allowing join", "65_live_quiz_join_desktop.png", 0.50, 0.45),
    (177, "Real-time Answer",     "Answer submission — real-time answer with time-based scoring", "65_live_quiz_join_desktop.png", 0.50, 0.60),
    (178, "Question Result Broadcast","All participants see correct answer + response stats simultaneously", "65_live_quiz_join_desktop.png", 0.50, 0.65),
    (179, "Auto-advance Timer",   "Question timer expires → auto-advance to results screen", "64_live_quiz_create_form_desktop.png", 0.65, 0.46),
    (180, "Result Countdown",     "After showing result — auto-advances to next question with countdown", "65_live_quiz_join_desktop.png", 0.50, 0.70),
    (181, "Extend Time",          "Host can add extra seconds to a live question mid-stream", "61_live_quiz_home_desktop.png", 0.80, 0.45),
    (182, "Duplicate Session Guard","Prevents creating duplicate live session for the same quiz", "64_live_quiz_create_form_desktop.png", 0.70, 0.82),
    (183, "Participated Dedup",   "Participated sessions exclude sessions already counted as hosted", "61_live_quiz_home_desktop.png", 0.50, 0.68),
    (184, "Session Status Badges","Session status: WAITING → ACTIVE → ENDED with colour badges", "61_live_quiz_home_desktop.png", 0.70, 0.60),

    # ── LEADERBOARD ───────────────────────────────────────────────────────────
    (185, "Leaderboard Podium",   "Rankings with animated podium (top 3) + full scores table below", "48_leaderboard_desktop.png", 0.50, 0.30),
    (186, "Filter by Stack",      "Filter leaderboard by subject/tech stack dropdown", "48_leaderboard_desktop.png", 0.80, 0.12),
    (187, "My Rank Highlighted",  "Current user's rank highlighted — 'YOUR RANK #X' badge", "49_leaderboard_table_desktop.png", 0.50, 0.65),
    (188, "Leaderboard Updates",  "Leaderboard score updates automatically after each quiz attempt", "50_leaderboard_assessment_desktop.png", 0.50, 0.40),
    (189, "SME Leaderboard Tab",  "SME MCQ contributions leaderboard tab", "50_leaderboard_sme_desktop.png", 0.20, 0.12),
    (190, "Assessment Tab",       "Assessment quiz leaderboard tab", "50_leaderboard_assessment_desktop.png", 0.40, 0.12),
    (191, "Live Battle Tab",      "Live Battle leaderboard tab — ranking from live sessions", "50_leaderboard_live_battle_desktop.png", 0.60, 0.12),

    # ── INBOX ─────────────────────────────────────────────────────────────────
    (192, "Inbox All Tab",        "Inbox 'All Messages' tab showing all received messages", "52_inbox_all_desktop.png", 0.10, 0.15),
    (193, "Compose Message",      "Compose new message — To / Subject / Body fields", "54_inbox_compose_desktop.png", 0.50, 0.40),
    (194, "Message Thread View",  "Click message → full thread view shown on right panel", "52_inbox_all_desktop.png", 0.70, 0.45),
    (195, "Reply to Message",     "Reply to message directly from the detail view", "52_inbox_all_desktop.png", 0.70, 0.80),
    (196, "Starred Messages",     "Star message → appears in dedicated Starred tab", "53_inbox_starred_desktop.png", 0.20, 0.15),
    (197, "Drafts Tab",           "Draft message auto-saved → visible in Drafts tab", "53_inbox_drafts_desktop.png", 0.30, 0.15),
    (198, "Sent Tab",             "Sent messages tab shows all outgoing messages", "53_inbox_sent_desktop.png", 0.40, 0.15),
    (199, "Trash Tab",            "Delete message → moved to Trash tab", "53_inbox_trash_desktop.png", 0.50, 0.15),

    # ── NOTIFICATIONS + CHATBOT ───────────────────────────────────────────────
    (200, "Notification Bell",    "Notification panel opens on bell icon click with unread badge count", "74_notifications_panel_desktop.png", 0.82, 0.03),
    (201, "Approved Notification","APPROVED notification: 'Your MCQ was approved' with MCQ title", "75_notifications_approved_desktop.png", 0.50, 0.35),
    (202, "Assigned Notification","ASSIGNED notification: 'A new MCQ is assigned for your review'", "75_notifications_assigned_desktop.png", 0.50, 0.35),
    (203, "Rejected Notification","REJECTED notification: 'Your MCQ was rejected' with reviewer reason", "75_notifications_rejected_desktop.png", 0.50, 0.35),
    (204, "Mark All Read",        "Mark all notifications as read in one click", "74_notifications_panel_desktop.png", 0.80, 0.15),
    (205, "Bell Badge Count",     "Bell icon shows real-time unread notification count badge", "74_notifications_panel_desktop.png", 0.82, 0.03),
    (206, "Notification Filter",  "Filter notifications by type (approved/rejected/assigned)", "74_notifications_panel_desktop.png", 0.70, 0.15),
    (207, "IST Notification Time","IST timestamp shown on every notification entry", "74_notifications_panel_desktop.png", 0.70, 0.40),
    (208, "Real-time Push",       "Real-time notification push — no page refresh needed", "74_notifications_panel_desktop.png", 0.50, 0.25),
    (209, "Click to Navigate",    "Click notification → directly navigates to relevant MCQ or page", "74_notifications_panel_desktop.png", 0.50, 0.40),
    (210, "Email Notification",   "Email notification sent for all status changes (Spring Mail + SMTP)", "74_notifications_panel_desktop.png", 0.50, 0.65),
    (211, "AI Chatbot Widget",    "AI chatbot widget permanently visible at bottom-right corner", "70_chatbot_open_desktop.png", 0.90, 0.93),
    (212, "Ask Anything ChatBot", "ChatBot responds to quiz/MCQ-related questions with AI answers", "70_chatbot_open_desktop.png", 0.70, 0.70),
    (213, "Chatbot Streaming",    "Chatbot responses stream word-by-word with typing indicator", "70_chatbot_open_desktop.png", 0.70, 0.75),
    (214, "Chatbot History",      "Chatbot conversation history preserved within the session", "70_chatbot_open_desktop.png", 0.70, 0.65),
    (215, "Sound Effects System", "Sound effects system — useSoundEffects hook, toggle on/off + volume", "05_admin_dashboard_desktop.png", 0.72, 0.03),
    (216, "Wellness Reminder",    "Wellness reminder popups — periodic break reminders as overlay", "05_admin_dashboard_desktop.png", 0.50, 0.50),

    # ── AI MCQ FEATURES ───────────────────────────────────────────────────────
    (217, "AI Generate Button",   "AI-assisted MCQ generation — 🤖 Generate with AI button on create form", "17_mcq_form_ai_section_desktop.png", 0.88, 0.22),
    (218, "Topic-Based AI Gen",   "Generate MCQ by entering a topic description into AI input", "17_mcq_form_ai_section_desktop.png", 0.50, 0.25),
    (219, "AI Fills All Fields",  "AI fills: question stem, all 4 options, correct answer, difficulty, Bloom's level", "17_mcq_form_ai_section_desktop.png", 0.50, 0.40),
    (220, "Edit AI Output",       "AI-generated MCQ is fully editable by user before saving", "17_mcq_form_ai_section_desktop.png", 0.50, 0.45),

    # ════════════════════ DOC 2 STARTS HERE ════════════════════════════════════

    (221, "Code AI Generation",   "Paste code snippet → AI generates a matching MCQ about that code", "17_mcq_form_ai_section_desktop.png", 0.85, 0.40),
    (222, "Scenario AI MCQ",      "Scenario description → AI writes a scenario-based MCQ", "17_mcq_form_ai_section_desktop.png", 0.80, 0.45),
    (223, "AI Bloom's Auto-pick", "AI auto-selects the appropriate Bloom's Taxonomy level", "17_mcq_form_ai_section_desktop.png", 0.80, 0.30),
    (224, "Difficulty AI Suggest","Difficulty auto-suggested by AI based on question complexity", "17_mcq_form_ai_section_desktop.png", 0.70, 0.30),
    (225, "One-Click Generate",   "One-click MCQ generation from existing question context/topic", "17_mcq_form_ai_section_desktop.png", 0.85, 0.35),
    (226, "AI Regenerate",        "Regenerate button — re-runs AI if first result is unsatisfactory", "17_mcq_form_ai_section_desktop.png", 0.85, 0.55),
    (227, "Export AI MCQ",        "Export AI-generated MCQ directly to My Questions page", "17_mcq_form_ai_section_desktop.png", 0.85, 0.50),
    (228, "Chain Generation",     "Chain generation — generate multiple related MCQs in one sequence", "17_mcq_form_ai_section_desktop.png", 0.50, 0.55),

    # ── AI STUDIO ─────────────────────────────────────────────────────────────
    (229, "AI Studio Hub",        "AI Studio — central hub for all advanced AI-powered tools", "35_ai_studio_main_desktop.png", 0.50, 0.10),
    (230, "AI Rewrite Tab",       "AI Rewrite tab — improve/rephrase existing MCQ text", "36_ai_studio_ai_rewrite_desktop.png", 0.20, 0.18),
    (231, "Code to MCQ Tab",      "Code to MCQ tab — paste code → AI generates MCQ question", "36_ai_studio_code_to_mcq_desktop.png", 0.40, 0.18),
    (232, "Learning Path Tab",    "Learning Path tab — AI generates structured learning roadmap per tech stack", "36_ai_studio_learning_path_desktop.png", 0.60, 0.18),
    (233, "AI Input Text Area",   "Text/code input area for AI processing in Studio", "35_ai_studio_main_desktop.png", 0.50, 0.40),
    (234, "AI Generate Button",   "Generate button triggers AI processing with current input", "35_ai_studio_main_desktop.png", 0.80, 0.50),
    (235, "AI Output Panel",      "Output panel shows AI-generated result in real time", "35_ai_studio_main_desktop.png", 0.50, 0.65),
    (236, "Export to Word",       "Export AI output to Word document (.docx)", "35_ai_studio_main_desktop.png", 0.85, 0.80),
    (237, "Tech Stack Context",   "Learning path uses selected tech stack as AI context", "36_ai_studio_learning_path_desktop.png", 0.30, 0.30),
    (238, "MCQ Count Config",     "Configure number of MCQs to generate in batch mode", "35_ai_studio_main_desktop.png", 0.70, 0.35),
    (239, "AI Rating Feedback",   "Rate the AI output quality — feedback improves future results", "35_ai_studio_main_desktop.png", 0.75, 0.70),
    (240, "Generation History",   "View history of past AI generations in left panel", "35_ai_studio_main_desktop.png", 0.15, 0.40),
    (241, "Screenshot to MCQ",    "Screenshot MCQ — upload image → AI extracts and creates MCQ", "37_screenshot_mcq_desktop.png", 0.50, 0.40),
    (242, "Bulk AI Generation",   "Bulk AI MCQ generation — generate many MCQs at once from list", "35_ai_studio_main_desktop.png", 0.80, 0.25),
    (243, "AI Chain Mode",        "AI chain mode — each generated output feeds the next generation step", "35_ai_studio_main_desktop.png", 0.50, 0.55),

    # ── CODING QUESTIONS ──────────────────────────────────────────────────────
    (244, "Coding Question Page", "Coding question page with full Monaco editor interface", "67_coding_question_desktop.png", 0.60, 0.45),
    (245, "Language Selector",    "Programming language selector — Java, Python, JavaScript, etc.", "67_coding_question_desktop.png", 0.35, 0.20),
    (246, "Problem Statement",    "Problem statement with description, constraints and examples", "67_coding_question_desktop.png", 0.25, 0.40),
    (247, "Monaco Code Editor",   "Monaco code editor with syntax highlighting and autocomplete", "68_coding_question_editor_desktop.png", 0.60, 0.50),
    (248, "Run Code Button",      "Run code button executes code against all test cases", "67_coding_question_desktop.png", 0.85, 0.20),
    (249, "Output Console",       "Output console shows execution results and error messages", "67_coding_question_desktop.png", 0.80, 0.70),
    (250, "Test Cases Panel",     "Test cases panel — expected vs actual output per test case", "69_coding_question_testcases_desktop.png", 0.80, 0.45),
    (251, "Submit Code",          "Submit code — all test cases must pass for full credit", "67_coding_question_desktop.png", 0.88, 0.25),
    (252, "Pass/Fail Indicator",  "Per-test-case pass/fail indicator with colour coding", "69_coding_question_testcases_desktop.png", 0.85, 0.65),
    (253, "Difficulty Badge",     "Difficulty badge displayed on coding question card", "67_coding_question_desktop.png", 0.25, 0.20),
    (254, "Example I/O Section",  "Example input/output section with sample test cases", "67_coding_question_desktop.png", 0.25, 0.60),
    (255, "Coding in Assessment", "Coding questions can be included in proctored quiz sessions", "68_coding_question_editor_desktop.png", 0.50, 0.20),

    # ── 17 INTERACTIVE QUESTION TYPES ─────────────────────────────────────────
    (256, "MCQ Type",             "Standard MCQ — 4 options single correct answer question type", "40_question_types_desktop.png", 0.15, 0.25),
    (257, "True/False Type",      "True/False question type — binary answer options", "40_question_types_desktop.png", 0.30, 0.25),
    (258, "Fill in the Blank",    "Fill in the blank — text input answer question type", "40_question_types_desktop.png", 0.45, 0.25),
    (259, "Match Columns",        "Match the following — drag columns to pair items", "40_question_types_desktop.png", 0.60, 0.25),
    (260, "Ordering Type",        "Ordering/sequence — arrange items in correct order", "40_question_types_desktop.png", 0.75, 0.25),
    (261, "Multi-Select Type",    "Multi-select — multiple correct answers from option list", "41_question_types_more_desktop.png", 0.15, 0.25),
    (262, "Coding Question Type", "Coding challenge question type with live code execution", "41_question_types_more_desktop.png", 0.30, 0.25),
    (263, "Hotspot Type",         "Hotspot — click on correct regions of an image", "41_question_types_more_desktop.png", 0.45, 0.25),
    (264, "Scenario-Based Type",  "Scenario-based — question with context paragraph and options", "41_question_types_more_desktop.png", 0.60, 0.25),
    (265, "Drag and Drop Type",   "Drag and drop — drag items into correct target zones", "41_question_types_more_desktop.png", 0.75, 0.25),
    (266, "Audio Response Type",  "Audio recording response — answer by speaking", "42_question_types_more2_desktop.png", 0.15, 0.25),
    (267, "Video Response Type",  "Video recording response — answer by recording video", "42_question_types_more2_desktop.png", 0.30, 0.25),
    (268, "Flowchart Type",       "Flowchart completion — fill missing steps in a process flow", "42_question_types_more2_desktop.png", 0.45, 0.25),
    (269, "Diagram Labelling",    "Diagram labelling — click/type labels on diagram regions", "42_question_types_more2_desktop.png", 0.60, 0.25),
    (270, "Code Trace Type",      "Code trace — predict output by tracing code execution", "42_question_types_more2_desktop.png", 0.75, 0.25),
    (271, "Type Creator Form",    "Custom question type creator form — define your own type", "43_question_type_creator_desktop.png", 0.50, 0.35),
    (272, "Live Preview Panel",   "Live preview panel updates as you build the custom question type", "43_question_type_creator_desktop.png", 0.80, 0.35),
    (273, "17 Types Badge",       "17 interactive question types — count badge on type catalogue page", "40_question_types_desktop.png", 0.90, 0.12),
    (274, "Add Custom Type",      "Add custom question type beyond the 17 built-in types", "43_question_type_creator_desktop.png", 0.80, 0.85),
    (275, "Export Question Type", "Export custom question type configuration as JSON", "43_question_type_creator_desktop.png", 0.85, 0.88),

    # ── RULEBOOK + SMART INTERVIEW KIT ────────────────────────────────────────
    (276, "RuleBook Page",        "RuleBook — centralized rule management for the entire platform", "38_rulebook_main_desktop.png", 0.50, 0.12),
    (277, "AI Rules Tab",         "AI Rules tab — rules governing AI MCQ generation behaviour", "39_rulebook_ai_rules_desktop.png", 0.20, 0.20),
    (278, "Lifecycle Rules Tab",  "Lifecycle tab — MCQ status transition rules and constraints", "39_rulebook_lifecycle_desktop.png", 0.35, 0.20),
    (279, "Roles Permissions Tab","Roles tab — RBAC role definitions and allowed permissions per role", "39_rulebook_roles_desktop.png", 0.50, 0.20),
    (280, "Workflow Rules Tab",   "Workflow tab — approval workflow rules and step definitions", "39_rulebook_workflow_desktop.png", 0.65, 0.20),
    (281, "Quiz Rules Tab",       "Quiz Rules tab — proctoring rules and exam conduct rules", "39_rulebook_quiz_rules_desktop.png", 0.80, 0.20),
    (282, "Rule Cards Display",   "Individual rule cards each showing title and full description", "38_rulebook_main_desktop.png", 0.50, 0.45),
    (283, "Edit Rule (Admin)",    "Edit rule action — Admin only, not visible to SME", "38_rulebook_main_desktop.png", 0.80, 0.45),
    (284, "Add New Rule",         "Add new rule button — Admin can add rules to any category", "38_rulebook_main_desktop.png", 0.85, 0.30),
    (285, "Rules Enforced",       "Rules are enforced in backend code — not just documentation", "38_rulebook_main_desktop.png", 0.50, 0.60),
    (286, "Duplicate MCQ Rules",  "Duplicate MCQ detection rules and threshold configuration", "39_rulebook_duplicate_desktop.png", 0.35, 0.50),
    (287, "Export Rules PDF",     "Export all rules as a PDF document for offline reference", "38_rulebook_main_desktop.png", 0.85, 0.80),
    (288, "Admin Only Actions",   "Admin-only badge on edit/delete rule actions", "38_rulebook_main_desktop.png", 0.80, 0.15),
    (289, "Rule Version Track",   "Rule version tracking — each rule update is versioned", "38_rulebook_main_desktop.png", 0.20, 0.50),
    (290, "Smart Interview Kit",  "Smart Interview Kit — AI-generated interview question sets per tech stack", "51_smart_interview_kit_desktop.png", 0.50, 0.35),
    (291, "Interview Set Export", "Export generated interview question set as PDF or Word document", "51_smart_interview_kit_desktop.png", 0.85, 0.80),

    # ── QUIZ ATTEMPTS + REAL-TIME ─────────────────────────────────────────────
    (292, "Quiz List Page",       "Quiz list page — all available assessment sessions", "46_quiz_list_desktop.png", 0.50, 0.35),
    (293, "Take Quiz Button",     "Take Quiz button launches the proctored quiz session", "46_quiz_list_desktop.png", 0.80, 0.35),
    (294, "Share Quiz Link",      "Share/copy quiz link button for distribution", "46_quiz_list_desktop.png", 0.85, 0.35),
    (295, "Attempt History",      "Past attempt history shown per session with scores and dates", "46_quiz_list_desktop.png", 0.50, 0.55),
    (296, "Score Per Attempt",    "Score displayed per attempt — percentage and raw points", "46_quiz_list_desktop.png", 0.75, 0.55),
    (297, "Real-time Quiz Chat",  "Real-time chat during live quiz session via WebSocket", "46_quiz_list_desktop.png", 0.90, 0.60),
    (298, "Session Expiry Guard", "Session expiry check enforced on quiz start attempt", "46_quiz_list_desktop.png", 0.65, 0.40),
    (299, "Non-Reg User Attempt", "Non-registered users take quiz — name + email, one attempt per email", "46_quiz_list_desktop.png", 0.50, 0.45),
    (300, "3 Strikes TERMINATED", "3 tab-switch strikes → quiz auto-submitted with TERMINATED badge", "46_quiz_list_desktop.png", 0.80, 0.45),
    (301, "Per-Question Timer",   "Per-question countdown timer displayed prominently during quiz", "46_quiz_list_desktop.png", 0.70, 0.30),
    (302, "TERMINATED Status",    "TERMINATED status badge shown on forced auto-submissions", "46_quiz_list_desktop.png", 0.75, 0.50),
    (303, "Per-Question Result",  "Per-question correct/wrong result shown after final submission", "46_quiz_list_desktop.png", 0.50, 0.65),
    (304, "Detailed Score Card",  "Detailed score card with complete question-level breakdown", "46_quiz_list_desktop.png", 0.50, 0.70),
    (305, "Anti-cheat Violations","Anti-cheat violation log stored per attempt in database", "46_quiz_list_desktop.png", 0.80, 0.40),
    (306, "Quiz Discussion",      "Discussion thread on quiz session page for post-quiz discussion", "46_quiz_list_desktop.png", 0.50, 0.75),
    (307, "Rate This Quiz",       "Rate the quiz session 1-5 stars after completion", "46_quiz_list_desktop.png", 0.70, 0.75),
    (308, "Report Question",      "Flag/report a problematic question from within the quiz", "46_quiz_list_desktop.png", 0.85, 0.75),
    (309, "Auto Leaderboard",     "Quiz result automatically updates main leaderboard standings", "46_quiz_list_desktop.png", 0.50, 0.80),
    (310, "Download Certificate", "Download completion certificate after finishing the quiz", "46_quiz_list_desktop.png", 0.80, 0.80),
    (311, "Export Attempts CSV",  "Export all quiz attempt data as CSV/Excel", "46_quiz_list_desktop.png", 0.88, 0.15),
    (312, "Filter Attempts",      "Filter quiz attempts by date range or completion status", "46_quiz_list_desktop.png", 0.50, 0.15),

    # ── INBOX ADVANCED ────────────────────────────────────────────────────────
    (313, "Inbox All Messages",   "Inbox — view all received messages in one place", "52_inbox_all_desktop.png", 0.10, 0.15),
    (314, "Search Messages",      "Search messages by keyword or sender name", "52_inbox_all_desktop.png", 0.50, 0.10),
    (315, "Compose Message",      "Compose new message with To / Subject / Body fields", "54_inbox_compose_desktop.png", 0.50, 0.40),
    (316, "Auto-Draft Save",      "Auto-draft saving with debounce — saves every 2s of inactivity", "54_inbox_compose_desktop.png", 0.70, 0.85),
    (317, "Debounce Save Indicator","'Saving...' indicator shown during debounce auto-save", "54_inbox_compose_desktop.png", 0.70, 0.87),
    (318, "Star Message",         "Star/unstar message — starred messages appear in Starred folder", "53_inbox_starred_desktop.png", 0.20, 0.15),
    (319, "Trash Recovery",       "Recover deleted message from Trash folder", "53_inbox_trash_desktop.png", 0.80, 0.35),
    (320, "Sent Tab View",        "Sent messages tab — view all outgoing messages with timestamps", "53_inbox_sent_desktop.png", 0.40, 0.15),
    (321, "Mail Unread Badge",    "Unread message count badge on mail icon in topbar", "52_inbox_all_desktop.png", 0.82, 0.03),
    (322, "Reply in Thread",      "Reply in-thread — maintains full conversation context", "52_inbox_all_desktop.png", 0.70, 0.80),
    (323, "Permanent Delete",     "Permanently delete message from Trash (no recovery)", "53_inbox_trash_desktop.png", 0.85, 0.35),
    (324, "Inbox Pagination",     "Inbox message list paginated for large inboxes", "52_inbox_all_desktop.png", 0.50, 0.90),

    # ── REVIEWER METRICS + RBAC ───────────────────────────────────────────────
    (325, "Reviewer Dashboard",   "Reviewer metrics dashboard — complete performance overview", "57_reviewer_dashboard_desktop.png", 0.50, 0.15),
    (326, "Assigned Count Card",  "Assigned questions count card (total assigned to this reviewer)", "57_reviewer_dashboard_desktop.png", 0.25, 0.20),
    (327, "Reviewed Count Card",  "Reviewed questions count card (total decisions made)", "57_reviewer_dashboard_desktop.png", 0.50, 0.20),
    (328, "Approval Rate Card",   "Approval rate % card — approved / total reviewed", "57_reviewer_dashboard_desktop.png", 0.75, 0.20),
    (329, "Assignment Table",     "Full assignment table listing MCQ details per assignment", "57_reviewer_dashboard_desktop.png", 0.50, 0.40),
    (330, "SLA Per Assignment",   "SLA indicator per assignment showing days remaining before breach", "57_reviewer_dashboard_desktop.png", 0.80, 0.40),
    (331, "SLA Breach Table",     "SLA breach table — overdue assignments with days exceeded", "59_reviewer_metrics_sla_table_desktop.png", 0.80, 0.55),
    (332, "Performance Chart",    "Reviewer performance chart — approval rate trend over time", "58_reviewer_metrics_desktop.png", 0.50, 0.60),
    (333, "Filter by Reviewer",   "Filter metrics by selecting a specific reviewer from dropdown", "57_reviewer_dashboard_desktop.png", 0.80, 0.15),
    (334, "Date Range Metrics",   "Date range filter — view metrics for 7d / 30d / 90d period", "58_reviewer_metrics_desktop.png", 0.70, 0.15),
    (335, "Export Metrics Excel", "Export reviewer metrics data as Excel spreadsheet", "57_reviewer_dashboard_desktop.png", 0.90, 0.15),
    (336, "Admin Role Badge",     "ADMIN role badge visible next to username in topbar", "05_admin_dashboard_desktop.png", 0.92, 0.03),
    (337, "SME Sidebar Only",     "SME sidebar — restricted to allowed pages, no admin sections", "76_sme_dashboard_desktop.png", 0.05, 0.35),
    (338, "RBAC Page Block",      "RBAC page block — unauthorized role gets 403/redirect immediately", "82_sme_rbac_blocked_user_mgmt_desktop.png", 0.50, 0.50),
    (339, "PrivateRoute Guard",   "PrivateRoute React component guards every protected page", "82_sme_rbac_blocked_user_mgmt_desktop.png", 0.50, 0.55),
    (340, "JWT Role Verification","JWT role claim verified server-side on every API request", "01_login_desktop.png", 0.30, 0.52),
    (341, "403 Redirect",         "403 unauthorized page shown clearly for role violations", "82_sme_rbac_blocked_user_mgmt_desktop.png", 0.50, 0.60),
    (342, "Audit Log Page",       "Audit log — records all admin actions with full timestamps", "55_audit_log_desktop.png", 0.50, 0.30),
    (343, "Admin Settings Page",  "Admin settings page — system-wide platform configuration", "60_admin_settings_desktop.png", 0.50, 0.30),
    (344, "Role Change Instant",  "Changing user's role takes effect on their next request", "29_user_management_desktop.png", 0.60, 0.35),
    (345, "Self Role Guard",      "Cannot change your own role (prevents accidental lockout)", "29_user_management_desktop.png", 0.80, 0.40),
    (346, "Audit Log Table",      "Audit log table with: action, actor, target entity, timestamp", "56_audit_log_table_desktop.png", 0.50, 0.45),
    (347, "Filter Audit Log",     "Filter audit log by action type or user performing action", "55_audit_log_desktop.png", 0.80, 0.15),
    (348, "Export Audit Log CSV", "Export audit log as CSV file for compliance/reporting", "55_audit_log_desktop.png", 0.90, 0.15),
    (349, "Metrics Per Decision", "Detailed reviewer metrics broken down per question decision", "58_reviewer_metrics_desktop.png", 0.50, 0.40),
    (350, "Metrics Period Select","Metrics period selector — 7d / 30d / 90d quick tabs", "58_reviewer_metrics_desktop.png", 0.70, 0.60),
    (351, "SLA Days Calculation", "SLA days = stuck date → SLA limit shown as 'Xd / Nd limit'", "59_reviewer_metrics_sla_table_desktop.png", 0.82, 0.55),
    (352, "Active Reviewers Count","Active reviewers count card in metrics dashboard", "57_reviewer_dashboard_desktop.png", 0.40, 0.30),
    (353, "Avg Review Time",      "Average review time per reviewer (hours/days to decision)", "58_reviewer_metrics_desktop.png", 0.60, 0.25),
    (354, "Bottleneck Detection", "Bottleneck detection — slowest reviewer highlighted for admin", "57_reviewer_dashboard_desktop.png", 0.50, 0.45),
    (355, "SME Reviewer Metrics", "SME can view their own reviewer performance metrics only", "78_sme_pending_reviews_desktop.png", 0.50, 0.30),

    # ── MCQ LIFECYCLE INTEGRITY ────────────────────────────────────────────────
    (356, "DRAFT Status Start",   "MCQ lifecycle begins at DRAFT status after creation", "21_kanban_board_desktop.png", 0.10, 0.35),
    (357, "→ READY_FOR_REVIEW",   "Submit → status changes to READY_FOR_REVIEW (awaiting admin assign)", "21_kanban_board_desktop.png", 0.28, 0.35),
    (358, "→ UNDER_REVIEW",       "Admin assigns reviewer → status becomes UNDER_REVIEW", "21_kanban_board_desktop.png", 0.46, 0.35),
    (359, "→ APPROVED",           "Reviewer approves → APPROVED status (final positive outcome)", "21_kanban_board_desktop.png", 0.64, 0.35),
    (360, "→ REJECTED",           "Reviewer rejects → REJECTED status with mandatory reason stored", "21_kanban_board_desktop.png", 0.82, 0.35),
    (361, "REJECTED → Resubmit",  "REJECTED MCQ: edit and resubmit → back to READY_FOR_REVIEW", "11_my_questions_rejected_desktop.png", 0.85, 0.35),
    (362, "PERMANENTLY_REJECTED", "PERMANENTLY_REJECTED — blocked from any future resubmission", "10_my_questions_all_desktop.png", 0.70, 0.35),
    (363, "Single Reviewer Rule", "Only one reviewer assigned per MCQ at any time", "28_assign_reviewer_dialog_desktop.png", 0.50, 0.55),
    (364, "Creator Excluded",     "MCQ creator is excluded from the reviewer assignment list", "28_assign_reviewer_dialog_desktop.png", 0.50, 0.60),
    (365, "Stack-Match Filter",   "Reviewer must be assigned to MCQ's tech stack to be eligible", "28_assign_reviewer_dialog_desktop.png", 0.50, 0.52),
    (366, "No Backward Status",   "Status cannot go backward — e.g., APPROVED → DRAFT is blocked", "21_kanban_board_desktop.png", 0.64, 0.15),
    (367, "Lifecycle in Kanban",  "Full 5-stage lifecycle visible and trackable in Kanban board", "21_kanban_board_desktop.png", 0.50, 0.10),

    # ── i18n + MOBILE + SECURITY ──────────────────────────────────────────────
    (368, "Language Menu",        "Language switcher menu — 7 languages with flag icons", "83_language_menu_open_desktop.png", 0.76, 0.03),
    (369, "English Locale",       "English (EN) locale — default language for the app", "83_language_menu_open_desktop.png", 0.50, 0.25),
    (370, "Hindi Locale",         "Hindi (HI) locale — full UI translated to Hindi", "83_language_menu_open_desktop.png", 0.50, 0.30),
    (371, "Tamil Locale",         "Tamil (TA) locale — full UI in Tamil language", "83_language_menu_open_desktop.png", 0.50, 0.35),
    (372, "Telugu Locale",        "Telugu (TE) locale — full UI in Telugu language", "83_language_menu_open_desktop.png", 0.50, 0.40),
    (373, "Kannada Locale",       "Kannada (KN) locale — full UI in Kannada language", "83_language_menu_open_desktop.png", 0.50, 0.45),
    (374, "Malayalam Locale",     "Malayalam (ML) locale — full UI in Malayalam language", "83_language_menu_open_desktop.png", 0.50, 0.50),
    (375, "Japanese Locale",      "Japanese (JA) locale — full UI in Japanese language", "83_language_menu_open_desktop.png", 0.50, 0.55),
    (376, "Language Persists",    "Selected language persisted in localStorage across all sessions", "83_language_menu_open_desktop.png", 0.50, 0.60),
    (377, "i18n RTL Considered",  "RTL layout considered in i18n design (Arabic/Urdu ready)", "83_language_menu_open_desktop.png", 0.50, 0.65),
    (378, "Mobile Login Page",    "Login page fully responsive on mobile devices", "01_login_mobile.png", 0.50, 0.50),
    (379, "Mobile Dashboard",     "Dashboard mobile layout — hamburger menu + stacked cards", "05_admin_dashboard_mobile.png", 0.05, 0.10),
    (380, "Mobile MCQ Form",      "MCQ create form fully responsive on mobile", "14_create_mcq_blank_mobile.png", 0.50, 0.40),
    (381, "Mobile My Questions",  "My Questions mobile table with horizontal scroll", "10_my_questions_all_mobile.png", 0.50, 0.40),
    (382, "Mobile Pending Reviews","Pending Reviews mobile card layout", "25_pending_reviews_mobile.png", 0.50, 0.40),
    (383, "Mobile Question Bank", "Question Bank admin page — mobile responsive layout", "26_question_bank_mobile.png", 0.50, 0.40),
    (384, "Mobile Kanban",        "Kanban board on mobile — horizontal scroll between columns", "21_kanban_board_mobile.png", 0.50, 0.40),
    (385, "Mobile Dark Mode",     "Dark mode works on mobile devices — fully responsive", "85_dark_mode_mobile.png", 0.73, 0.03),
    (386, "XSS Prevention",       "XSS prevention — all user input sanitized with DOMPurify", "14_create_mcq_blank_desktop.png", 0.50, 0.22),
    (387, "SQL Injection Guard",  "SQL injection prevented — all DB queries use JPA parameterized queries", "29_user_management_desktop.png", 0.50, 0.30),
    (388, "CORS Configuration",   "CORS configured — only localhost:3000 allowed in development", "60_admin_settings_desktop.png", 0.50, 0.40),
    (389, "CSRF Protection",      "CSRF protection via Spring Security stateless JWT — no cookies", "01_login_desktop.png", 0.30, 0.52),
    (390, "API Rate Limiting",    "API rate limiting prevents abuse and DoS attempts", "01_login_desktop.png", 0.30, 0.38),
    (391, "HTTPS Ready",          "HTTPS-ready configuration — SSL/TLS support via Spring Boot config", "60_admin_settings_desktop.png", 0.50, 0.50),

    # ── BACKEND + INFRA + DATA MODEL ──────────────────────────────────────────
    (392, "Spring Boot 3.2.5",    "Spring Boot 3.2.5 backend with Java 17 runtime", "60_admin_settings_desktop.png", 0.50, 0.20),
    (393, "MySQL Schema",         "MySQL 8.x database with 'quizhub' schema — 19 database entities", "60_admin_settings_desktop.png", 0.50, 0.30),
    (394, "H2 In-Memory Tests",   "H2 in-memory DB used for all backend tests — no MySQL needed for CI", "60_admin_settings_desktop.png", 0.50, 0.35),
    (395, "data.sql Auto-Seed",   "data.sql auto-seeds: 6 tech stacks, 5 users, 40+ topics on first startup", "60_admin_settings_desktop.png", 0.50, 0.40),
    (396, "Spring AI + OpenAI",   "Spring AI + OpenAI GPT-4o-mini integration for all AI features", "35_ai_studio_main_desktop.png", 0.80, 0.50),
    (397, "Ollama Local Profile", "Ollama local AI profile — spring.profiles.active=ollama for offline AI", "60_admin_settings_desktop.png", 0.50, 0.45),
    (398, "pgVector Embeddings",  "pgVector profile for semantic search using vector embeddings", "60_admin_settings_desktop.png", 0.50, 0.50),
    (399, "140+ API Endpoints",   "140+ REST API endpoints covering all platform features", "60_admin_settings_desktop.png", 0.50, 0.55),
    (400, "Swagger API Docs",     "Swagger/OpenAPI documentation at /swagger-ui/index.html", "60_admin_settings_desktop.png", 0.50, 0.60),
    (401, "Actuator Health",      "Spring Actuator /actuator/health endpoint returning status UP", "60_admin_settings_desktop.png", 0.50, 0.65),
    (402, "JaCoCo 92.5%",         "JaCoCo test coverage — 92.5% instruction coverage achieved", "60_admin_settings_desktop.png", 0.50, 0.70),
    (403, "1072 Backend Tests",   "1,072 backend tests — all passing with 0 failures", "60_admin_settings_desktop.png", 0.50, 0.75),
    (404, "Jest 80.37% Coverage", "Jest frontend coverage — 80.37% statement coverage achieved", "60_admin_settings_desktop.png", 0.50, 0.78),
    (405, "957 Frontend Tests",   "957 frontend Jest tests — all passing with 0 failures", "60_admin_settings_desktop.png", 0.50, 0.80),
    (406, "PWA Service Worker",   "PWA service worker — offline support and installable app", "60_admin_settings_desktop.png", 0.50, 0.82),
    (407, "localStorage Cache",   "localStorage cache system for instant page loads (getCacheSync)", "60_admin_settings_desktop.png", 0.50, 0.85),
    (408, "prefetchAll System",   "prefetchAll() preloads data for 8 pages at login — zero spinners", "05_admin_dashboard_desktop.png", 0.50, 0.65),
    (409, "User Entity",          "User entity: id, enterpriseId, name, email, role, status, techStacks list", "29_user_management_desktop.png", 0.50, 0.30),
    (410, "MCQ Entity",           "MCQ entity: id, stem, options[], correctAnswer, difficulty, bloomsLevel, status", "18_mcq_detail_desktop.png", 0.50, 0.25),
    (411, "TechStack Entity",     "TechStack entity: id, name, topics (one-to-many relationship)", "30_master_data_desktop.png", 0.50, 0.30),
    (412, "Topic Entity",         "Topic entity: id, name, linked TechStack foreign key", "31_master_data_topics_desktop.png", 0.50, 0.55),
    (413, "Review Entity",        "Review entity: reviewer, MCQ, decision, comment, timestamps", "57_reviewer_dashboard_desktop.png", 0.50, 0.40),
    (414, "Notification Entity",  "Notification entity: type, message, isRead, createdAt, user FK", "74_notifications_panel_desktop.png", 0.50, 0.40),
    (415, "InboxMessage Entity",  "InboxMessage entity: sender, receiver, subject, body, status, timestamps", "52_inbox_all_desktop.png", 0.50, 0.45),
    (416, "QuizSession Entity",   "QuizSession entity: title, techStack, questions[], timeLimit, expiresAt, status", "44_quiz_builder_desktop.png", 0.50, 0.40),
    (417, "QuizAttempt Entity",   "QuizAttempt entity: user/email, session, score, violations, status (TERMINATED)", "46_quiz_list_desktop.png", 0.50, 0.55),
    (418, "LiveSession Entity",   "LiveSession entity: pin, host, co-host, participants, questions, status", "61_live_quiz_home_desktop.png", 0.50, 0.60),
    (419, "AuditLog Entity",      "AuditLog entity: action type, actorId, targetId, timestamp, details", "56_audit_log_table_desktop.png", 0.50, 0.45),
    (420, "McqComment Entity",    "McqComment entity: mcqId, authorId, text, replyToId, createdAt", "18_mcq_detail_desktop.png", 0.50, 0.74),
    (421, "DuplicateCheck Config","Duplicate MCQ detection config — semantic similarity threshold setting", "39_rulebook_duplicate_desktop.png", 0.35, 0.50),

    # ── LATEST UI/UX ENHANCEMENTS ─────────────────────────────────────────────
    (422, "Transformers Topbar",  "Transformers-inspired animated topbar strip — cinematic story cycling every ~33s", "05_admin_dashboard_desktop.png", 0.50, 0.03),
    (423, "Bumblebee Car Anim",   "Bumblebee Camaro CSS animation drives full topbar width with spinning wheels", "05_admin_dashboard_desktop.png", 0.35, 0.03),
    (424, "Explosion Burst FX",   "Explosion burst with 6 sparkle emoji particles when car reaches end", "05_admin_dashboard_desktop.png", 0.60, 0.03),
    (425, "CSS Robot Rise",       "CSS robot materialises from explosion with arm-throw animation", "05_admin_dashboard_desktop.png", 0.65, 0.03),
    (426, "Saucer Launch",        "Saucer launch trailing '✦ BumbleBee Team ✦' label across screen", "05_admin_dashboard_desktop.png", 0.70, 0.03),
    (427, "Floating Orbs + Stars","Floating orbs + twinkling stars as live animated background accents", "05_admin_dashboard_desktop.png", 0.80, 0.03),
    (428, "Drag Reorder Widgets", "Drag-to-reorder dashboard widgets — order persisted in localStorage", "05_admin_dashboard_desktop.png", 0.17, 0.27),
    (429, "Sound Effects Toggle", "Sound effects system — useSoundEffects hook; toggle on/off + volume slider", "05_admin_dashboard_desktop.png", 0.72, 0.03),
    (430, "Keyboard Shortcuts",   "Press '?' anywhere → overlay modal listing all keyboard shortcuts", "73_keyboard_shortcuts_desktop.png", 0.50, 0.40),
    (431, "Wellness Reminders",   "Wellness reminder popups — periodic non-intrusive break reminders", "05_admin_dashboard_desktop.png", 0.50, 0.50),
    (432, "SW Cache Fix",         "Service worker cache fix — JS/CSS bundles excluded so new deploys load instantly", "60_admin_settings_desktop.png", 0.50, 0.82),
    (433, "Zero Spinner Loads",   "Instant zero-spinner page loads via getCacheSync() + prefetchAll() system", "05_admin_dashboard_desktop.png", 0.50, 0.65),
    (434, "SLA Stuck + Limit",    "SLA breach table shows days stuck AND '/ Nd limit' inline per assignment", "59_reviewer_metrics_sla_table_desktop.png", 0.80, 0.55),

    # ── SME SPECIFIC FEATURES ─────────────────────────────────────────────────
    (435, "SME Dashboard",        "SME-specific dashboard showing own stats and assigned review count", "76_sme_dashboard_desktop.png", 0.50, 0.20),
    (436, "SME My Questions",     "SME My Questions — only that SME's own MCQs visible (scoped)", "77_sme_my_questions_desktop.png", 0.50, 0.40),
    (437, "SME Pending Reviews",  "SME Pending Reviews — MCQs assigned to that SME as reviewer", "78_sme_pending_reviews_desktop.png", 0.50, 0.40),
    (438, "SME RBAC Enforced",    "SME completely blocked from User Management and Master Data pages", "82_sme_rbac_blocked_user_mgmt_desktop.png", 0.50, 0.50),
    (439, "Dark Mode Full App",   "Dark mode applied consistently across every page and component", "85_dark_mode_desktop.png", 0.73, 0.03),
    (440, "Dark Mode Mobile",     "Full dark theme on mobile — consistent with desktop dark mode", "85_dark_mode_mobile.png", 0.73, 0.03),
]

# ─── main ─────────────────────────────────────────────────────────────────────
if __name__ == "__main__":
    print(f"Total features: {len(FEATURES)}")
    split = 220

    doc1_features = FEATURES[:split]
    doc2_features = FEATURES[split:]

    print(f"\n=== Building Doc 1 ({len(doc1_features)} features) ===")
    build_doc(doc1_features, OUT1, 1)

    print(f"\n=== Building Doc 2 ({len(doc2_features)} features) ===")
    build_doc(doc2_features, OUT2, 2)

    print("\n🎉 Both documents complete!")
    print(f"   Doc1: {OUT1}")
    print(f"   Doc2: {OUT2}")
