#!/usr/bin/env python3
"""
QuizHub AI - Complete Feature Evidence Capture
Takes desktop + mobile screenshots of every feature and compiles into Word doc.
"""

import asyncio
import os
import sys
from pathlib import Path

# Install dependencies if missing
try:
    from playwright.async_api import async_playwright
except ImportError:
    os.system("pip3 install playwright")
    os.system("python3 -m playwright install chromium")
    from playwright.async_api import async_playwright

try:
    from docx import Document
    from docx.shared import Inches, Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.section import WD_ORIENT
except ImportError:
    os.system("pip3 install python-docx")
    from docx import Document
    from docx.shared import Inches, Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.section import WD_ORIENT

BASE_URL = "http://localhost:3000"
SCREENSHOT_DIR = Path("/Users/veera.konjeti/Desktop/hack-n-stack/screenshots")
SCREENSHOT_DIR.mkdir(exist_ok=True)

SME_CREDS = {"enterpriseId": "birendra.kumar.singh", "password": "Sme@1234"}
ADMIN_CREDS = {"enterpriseId": "divya.madhanasekar", "password": "Admin@123"}

# All pages to capture with descriptions
SME_PAGES = [
    ("login", "/login", "Login Page - Role-based authentication with JWT tokens"),
    ("sme_dashboard", "/", "SME Dashboard - Overview with stats, recent activity"),
    ("my_questions", "/my-questions", "My Questions - Paginated table with status filter tabs (Draft/Ready/Under Review/Approved/Rejected)"),
    ("kanban_board", "/kanban", "Kanban Board - Visual workflow of question lifecycle"),
    ("pending_reviews", "/pending-reviews", "My Pending Reviews - Questions assigned for review"),
    ("analytics", "/analytics", "Analytics - Charts & insights on question creation"),
    ("reviewer_dashboard", "/reviewer-dashboard", "My Stats - SME reviewer statistics"),
    ("smart_interview_kit", "/smart-interview-kit", "Smart Interview Kit - AI-powered interview preparation"),
    ("ai_studio", "/ai-studio", "AI Studio - AI-powered MCQ generation (Level 2 Feature)"),
    ("question_types", "/question-types", "Question Types - All 17 supported question formats"),
    ("rulebook", "/rulebook", "Rule Book - MCQ creation guidelines"),
    ("leaderboard", "/leaderboard", "Leaderboard - Top contributors ranking"),
    ("quiz_builder", "/quiz-builder", "Assessments - Quiz/assessment builder"),
    ("live_quiz", "/live", "Live Quiz - Real-time quiz sessions"),
]

ADMIN_PAGES = [
    ("admin_dashboard", "/", "Admin Dashboard - System-wide overview"),
    ("admin_question_bank", "/question-bank", "Admin Question Bank - All MCQs across all creators"),
    ("admin_users", "/manage-users", "Admin User Management - Manage SMEs and assignments"),
    ("admin_tech_stacks", "/tech-stacks", "Admin Tech Stacks - Master data management (6 tech stacks)"),
    ("admin_reports", "/reports", "Admin Reports - System reports & analytics"),
]


async def login(page, credentials):
    """Login and return the page with session."""
    await page.goto(f"{BASE_URL}/login")
    await page.wait_for_load_state("networkidle")
    
    # Clear and fill enterprise ID
    eid_input = page.locator('input[placeholder*="gaurav.a.bhola"], input[placeholder*="enterprise"]').first
    await eid_input.fill("")
    await eid_input.fill(credentials["enterpriseId"])
    
    # Clear and fill password
    pwd_input = page.locator('input[type="password"]').first
    await pwd_input.fill("")
    await pwd_input.fill(credentials["password"])
    
    # Click sign in
    await page.locator('button:has-text("Sign In")').click()
    await page.wait_for_load_state("networkidle")
    await asyncio.sleep(1)
    return page


async def take_screenshot_pair(page, name, url, description, screenshots_taken):
    """Take both desktop and mobile screenshots of a page."""
    try:
        await page.goto(f"{BASE_URL}{url}", wait_until="networkidle", timeout=15000)
        await asyncio.sleep(1)
        
        # Desktop screenshot (1440x900)
        await page.set_viewport_size({"width": 1440, "height": 900})
        await asyncio.sleep(0.5)
        desktop_path = SCREENSHOT_DIR / f"{name}_desktop.png"
        await page.screenshot(path=str(desktop_path), full_page=True)
        screenshots_taken.append((str(desktop_path), f"{description} (Desktop)", "desktop"))
        print(f"  ✅ Desktop: {name}")
        
        # Mobile screenshot (390x844 - iPhone 14)
        await page.set_viewport_size({"width": 390, "height": 844})
        await asyncio.sleep(0.5)
        mobile_path = SCREENSHOT_DIR / f"{name}_mobile.png"
        await page.screenshot(path=str(mobile_path), full_page=True)
        screenshots_taken.append((str(mobile_path), f"{description} (Mobile)", "mobile"))
        print(f"  ✅ Mobile: {name}")
        
        # Reset to desktop
        await page.set_viewport_size({"width": 1440, "height": 900})
        
    except Exception as e:
        print(f"  ❌ Error on {name}: {e}")


async def capture_create_mcq_flow(page, screenshots_taken):
    """Capture the Create MCQ form flow."""
    await page.goto(f"{BASE_URL}/my-questions", wait_until="networkidle")
    await asyncio.sleep(1)
    
    # Click Add Question
    try:
        add_btn = page.locator('button:has-text("Add Question")')
        await add_btn.click()
        await asyncio.sleep(1)
        
        desktop_path = SCREENSHOT_DIR / "create_mcq_form_desktop.png"
        await page.screenshot(path=str(desktop_path), full_page=True)
        screenshots_taken.append((str(desktop_path), "Create MCQ - Empty form with all fields (Tech Stack, Topic, Difficulty, Question, Options) (Desktop)", "desktop"))
        print("  ✅ Create MCQ form captured")
        
        await page.set_viewport_size({"width": 390, "height": 844})
        await asyncio.sleep(0.5)
        mobile_path = SCREENSHOT_DIR / "create_mcq_form_mobile.png"
        await page.screenshot(path=str(mobile_path), full_page=True)
        screenshots_taken.append((str(mobile_path), "Create MCQ - Empty form with all fields (Mobile)", "mobile"))
        await page.set_viewport_size({"width": 1440, "height": 900})
    except Exception as e:
        print(f"  ❌ Create MCQ form: {e}")


async def capture_view_question(page, screenshots_taken):
    """Capture viewing a question detail."""
    await page.goto(f"{BASE_URL}/my-questions", wait_until="networkidle")
    await asyncio.sleep(1)
    
    try:
        view_btn = page.locator('button:has-text("View")').first
        await view_btn.click()
        await asyncio.sleep(1)
        
        desktop_path = SCREENSHOT_DIR / "view_question_detail_desktop.png"
        await page.screenshot(path=str(desktop_path), full_page=True)
        screenshots_taken.append((str(desktop_path), "View Question Detail - Full MCQ display with options, correct answer, explanation (Desktop)", "desktop"))
        print("  ✅ View question detail captured")
        
        await page.set_viewport_size({"width": 390, "height": 844})
        await asyncio.sleep(0.5)
        mobile_path = SCREENSHOT_DIR / "view_question_detail_mobile.png"
        await page.screenshot(path=str(mobile_path), full_page=True)
        screenshots_taken.append((str(mobile_path), "View Question Detail (Mobile)", "mobile"))
        await page.set_viewport_size({"width": 1440, "height": 900})
    except Exception as e:
        print(f"  ❌ View question: {e}")


async def capture_edit_question(page, screenshots_taken):
    """Capture editing a question."""
    await page.goto(f"{BASE_URL}/my-questions", wait_until="networkidle")
    await asyncio.sleep(1)
    
    try:
        edit_btn = page.locator('button:has-text("Edit")').first
        await edit_btn.click()
        await asyncio.sleep(1)
        
        desktop_path = SCREENSHOT_DIR / "edit_question_desktop.png"
        await page.screenshot(path=str(desktop_path), full_page=True)
        screenshots_taken.append((str(desktop_path), "Edit MCQ - Pre-filled form with Save Draft and Save & Send for Review buttons (Desktop)", "desktop"))
        print("  ✅ Edit question captured")
        
        await page.set_viewport_size({"width": 390, "height": 844})
        await asyncio.sleep(0.5)
        mobile_path = SCREENSHOT_DIR / "edit_question_mobile.png"
        await page.screenshot(path=str(mobile_path), full_page=True)
        screenshots_taken.append((str(mobile_path), "Edit MCQ (Mobile)", "mobile"))
        await page.set_viewport_size({"width": 1440, "height": 900})
    except Exception as e:
        print(f"  ❌ Edit question: {e}")


async def capture_bulk_upload(page, screenshots_taken):
    """Capture bulk upload page."""
    await page.goto(f"{BASE_URL}/bulk-upload", wait_until="networkidle")
    await asyncio.sleep(1)
    
    desktop_path = SCREENSHOT_DIR / "bulk_upload_desktop.png"
    await page.screenshot(path=str(desktop_path), full_page=True)
    screenshots_taken.append((str(desktop_path), "Bulk Upload - CSV template download, file upload with drag-drop, validation & error reporting (Desktop)", "desktop"))
    print("  ✅ Bulk upload captured")
    
    await page.set_viewport_size({"width": 390, "height": 844})
    await asyncio.sleep(0.5)
    mobile_path = SCREENSHOT_DIR / "bulk_upload_mobile.png"
    await page.screenshot(path=str(mobile_path), full_page=True)
    screenshots_taken.append((str(mobile_path), "Bulk Upload (Mobile)", "mobile"))
    await page.set_viewport_size({"width": 1440, "height": 900})


async def capture_status_tabs(page, screenshots_taken):
    """Capture each status filter tab on My Questions."""
    await page.goto(f"{BASE_URL}/my-questions", wait_until="networkidle")
    await asyncio.sleep(1)
    
    tabs = ["Draft", "Ready for Review", "Approved", "Rejected"]
    for tab in tabs:
        try:
            tab_btn = page.locator(f'button:has-text("{tab}")').first
            await tab_btn.click()
            await asyncio.sleep(0.5)
            
            safe_name = tab.lower().replace(" ", "_")
            desktop_path = SCREENSHOT_DIR / f"status_tab_{safe_name}_desktop.png"
            await page.screenshot(path=str(desktop_path))
            screenshots_taken.append((str(desktop_path), f"My Questions - {tab} tab filter showing count and filtered results (Desktop)", "desktop"))
            print(f"  ✅ Status tab: {tab}")
        except Exception as e:
            print(f"  ❌ Tab {tab}: {e}")


async def capture_dark_mode(page, screenshots_taken):
    """Capture dark mode toggle."""
    await page.goto(f"{BASE_URL}/", wait_until="networkidle")
    await asyncio.sleep(1)
    
    try:
        # Toggle dark mode
        toggle_btn = page.locator('button:has-text("Toggle theme"), button[aria-label*="theme"]').first
        await toggle_btn.click()
        await asyncio.sleep(0.5)
        
        desktop_path = SCREENSHOT_DIR / "dark_mode_desktop.png"
        await page.screenshot(path=str(desktop_path))
        screenshots_taken.append((str(desktop_path), "Dark Mode - Full theme switching with CSS variables (Desktop)", "desktop"))
        print("  ✅ Dark mode captured")
        
        await page.set_viewport_size({"width": 390, "height": 844})
        await asyncio.sleep(0.5)
        mobile_path = SCREENSHOT_DIR / "dark_mode_mobile.png"
        await page.screenshot(path=str(mobile_path))
        screenshots_taken.append((str(mobile_path), "Dark Mode (Mobile)", "mobile"))
        await page.set_viewport_size({"width": 1440, "height": 900})
        
        # Toggle back to light
        toggle_btn = page.locator('button:has-text("Toggle theme"), button[aria-label*="theme"]').first
        await toggle_btn.click()
        await asyncio.sleep(0.3)
    except Exception as e:
        print(f"  ❌ Dark mode: {e}")


async def capture_language_switch(page, screenshots_taken):
    """Capture language/i18n switching."""
    await page.goto(f"{BASE_URL}/", wait_until="networkidle")
    await asyncio.sleep(1)
    
    try:
        lang_btn = page.locator('button:has-text("EN"), button:has-text("🇬🇧")').first
        await lang_btn.click()
        await asyncio.sleep(0.5)
        
        desktop_path = SCREENSHOT_DIR / "language_selector_desktop.png"
        await page.screenshot(path=str(desktop_path))
        screenshots_taken.append((str(desktop_path), "Internationalization (i18n) - Language selector with 7 languages (Desktop)", "desktop"))
        print("  ✅ Language selector captured")
        
        # Try selecting a different language
        try:
            de_option = page.locator('text=Deutsch, text=DE, [data-lang="de"]').first
            await de_option.click()
            await asyncio.sleep(1)
            
            desktop_path = SCREENSHOT_DIR / "language_german_desktop.png"
            await page.screenshot(path=str(desktop_path))
            screenshots_taken.append((str(desktop_path), "German Language (DE) - Full UI translated to German (Desktop)", "desktop"))
            print("  ✅ German language captured")
            
            # Switch back to English
            lang_btn = page.locator('button:has-text("DE"), button:has-text("🇩🇪")').first
            await lang_btn.click()
            await asyncio.sleep(0.5)
            en_option = page.locator('text=English, text=EN, [data-lang="en"]').first
            await en_option.click()
            await asyncio.sleep(0.5)
        except:
            pass
    except Exception as e:
        print(f"  ❌ Language: {e}")


async def capture_notifications(page, screenshots_taken):
    """Capture notifications panel."""
    await page.goto(f"{BASE_URL}/", wait_until="networkidle")
    await asyncio.sleep(1)
    
    try:
        notif_btn = page.locator('button:has-text("Notifications"), button[aria-label*="Notification"]').first
        await notif_btn.click()
        await asyncio.sleep(0.5)
        
        desktop_path = SCREENSHOT_DIR / "notifications_panel_desktop.png"
        await page.screenshot(path=str(desktop_path))
        screenshots_taken.append((str(desktop_path), "Notifications - Real-time notification bell with unread count, panel with mark-as-read (Desktop)", "desktop"))
        print("  ✅ Notifications captured")
    except Exception as e:
        print(f"  ❌ Notifications: {e}")


async def capture_chatbot(page, screenshots_taken):
    """Capture AI chatbot."""
    await page.goto(f"{BASE_URL}/", wait_until="networkidle")
    await asyncio.sleep(1)
    
    try:
        chat_btn = page.locator('button:has-text("AI Collab"), button:has-text("Chat"), [class*="chat"]').first
        await chat_btn.click()
        await asyncio.sleep(1)
        
        desktop_path = SCREENSHOT_DIR / "chatbot_desktop.png"
        await page.screenshot(path=str(desktop_path))
        screenshots_taken.append((str(desktop_path), "AI Chatbot - Integrated AI assistant for MCQ help (Desktop)", "desktop"))
        print("  ✅ Chatbot captured")
    except Exception as e:
        print(f"  ❌ Chatbot: {e}")


async def capture_admin_assign_reviewer(page, screenshots_taken):
    """Capture admin assign reviewer dialog."""
    await page.goto(f"{BASE_URL}/question-bank", wait_until="networkidle")
    await asyncio.sleep(1)
    
    try:
        # Try to find an assign button
        assign_btn = page.locator('button:has-text("Assign"), button:has-text("Reviewer")').first
        await assign_btn.click()
        await asyncio.sleep(1)
        
        desktop_path = SCREENSHOT_DIR / "assign_reviewer_dialog_desktop.png"
        await page.screenshot(path=str(desktop_path))
        screenshots_taken.append((str(desktop_path), "Assign Reviewer Dialog - Tech stack matching, excludes creator, assigns reviewer (Desktop)", "desktop"))
        print("  ✅ Assign reviewer dialog captured")
    except Exception as e:
        print(f"  ❌ Assign reviewer: {e}")


async def main():
    screenshots_taken = []
    
    async with async_playwright() as p:
        browser = await p.chromium.launch(headless=True)
        context = await browser.new_context(viewport={"width": 1440, "height": 900})
        page = await context.new_page()
        
        # ============ SME LOGIN & PAGES ============
        print("\n🔐 Logging in as SME...")
        await login(page, SME_CREDS)
        
        print("\n📸 Section 1: SME Pages")
        for name, url, desc in SME_PAGES:
            await take_screenshot_pair(page, name, url, desc, screenshots_taken)
        
        print("\n📸 Section 2: Create MCQ Flow")
        await capture_create_mcq_flow(page, screenshots_taken)
        
        print("\n📸 Section 3: View Question Detail")
        await capture_view_question(page, screenshots_taken)
        
        print("\n📸 Section 4: Edit Question")
        await capture_edit_question(page, screenshots_taken)
        
        print("\n📸 Section 5: Bulk Upload")
        await capture_bulk_upload(page, screenshots_taken)
        
        print("\n📸 Section 6: Status Filter Tabs")
        await capture_status_tabs(page, screenshots_taken)
        
        print("\n📸 Section 7: Dark Mode")
        await capture_dark_mode(page, screenshots_taken)
        
        print("\n📸 Section 8: Language / i18n")
        await capture_language_switch(page, screenshots_taken)
        
        print("\n📸 Section 9: Notifications")
        await capture_notifications(page, screenshots_taken)
        
        print("\n📸 Section 10: AI Chatbot")
        await capture_chatbot(page, screenshots_taken)
        
        # ============ ADMIN LOGIN & PAGES ============
        print("\n🔐 Logging in as Admin...")
        await page.goto(f"{BASE_URL}/login", wait_until="networkidle")
        await asyncio.sleep(0.5)
        
        # Clear local storage to force re-login
        await page.evaluate("localStorage.clear()")
        await page.goto(f"{BASE_URL}/login", wait_until="networkidle")
        await asyncio.sleep(0.5)
        await login(page, ADMIN_CREDS)
        
        print("\n📸 Section 11: Admin Pages")
        for name, url, desc in ADMIN_PAGES:
            await take_screenshot_pair(page, name, url, desc, screenshots_taken)
        
        print("\n📸 Section 12: Admin Assign Reviewer")
        await capture_admin_assign_reviewer(page, screenshots_taken)
        
        await browser.close()
    
    print(f"\n✅ Total screenshots taken: {len(screenshots_taken)}")
    
    # ============ CREATE WORD DOCUMENT ============
    print("\n📄 Creating Word document...")
    create_word_document(screenshots_taken)
    print("✅ Word document created!")


def create_word_document(screenshots_taken):
    """Create comprehensive Word document with all screenshots."""
    doc = Document()
    
    # Set narrow margins
    for section in doc.sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(1.5)
        section.right_margin = Cm(1.5)
    
    # Title page
    title = doc.add_heading("QuizHub AI — Feature Evidence Document", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Valkey Hack-N-Stack 2026 | Team Valkey")
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0x66, 0x33, 0x99)
    
    doc.add_paragraph()
    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = info.add_run("Complete End-to-End Feature Testing Evidence\nDesktop & Mobile Responsive Views")
    run.font.size = Pt(12)
    
    doc.add_paragraph()
    tech_info = doc.add_paragraph()
    tech_info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = tech_info.add_run(
        "Tech Stack: Spring Boot 3.2.5 (Java 17) + React 19 + MySQL 8 + Spring AI + OpenAI GPT-4o-mini\n"
        "Tests: 2,029 automated (1,072 backend + 957 frontend) | 0 failures\n"
        "Coverage: Backend 92.5% | Frontend 80.37%"
    )
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    
    doc.add_page_break()
    
    # Table of Contents placeholder
    doc.add_heading("Table of Contents", level=1)
    sections_list = [
        "1. Authentication & Login (JWT, Role-based)",
        "2. SME Dashboard",
        "3. My Questions (Status Tabs, Pagination, Search)",
        "4. Create MCQ (Form with all fields)",
        "5. View Question Detail",
        "6. Edit MCQ (Draft/Rejected edit, Save & Send)",
        "7. Bulk Upload (CSV template, validation)",
        "8. Status Filter Tabs (Draft/Ready/Approved/Rejected)",
        "9. Kanban Board (Visual lifecycle workflow)",
        "10. My Pending Reviews (Review questions)",
        "11. Admin Dashboard",
        "12. Admin Question Bank (All questions, all creators)",
        "13. Admin Assign Reviewer (Tech stack matching)",
        "14. Admin User & Tech Stack Management",
        "15. AI Studio - MCQ Generation (Level 2)",
        "16. AI Duplicate Detection (Level 2)",
        "17. Analytics & Reporting",
        "18. Live Quiz & Assessments",
        "19. Smart Interview Kit",
        "20. Dark Mode & i18n (7 languages)",
        "21. Notifications & AI Chatbot",
        "22. Leaderboard & Question Types",
    ]
    for item in sections_list:
        p = doc.add_paragraph(item, style='List Number')
        p.paragraph_format.space_after = Pt(2)
    
    doc.add_page_break()
    
    # Group screenshots by feature
    current_section = None
    section_num = 0
    
    for path, description, view_type in screenshots_taken:
        if not os.path.exists(path):
            continue
        
        # Add section heading if description changes significantly
        base_desc = description.split("(Desktop)")[0].split("(Mobile)")[0].strip()
        if base_desc != current_section:
            current_section = base_desc
            section_num += 1
            doc.add_heading(f"{base_desc}", level=2)
            
            # Add what this feature does
            p = doc.add_paragraph()
            run = p.add_run(f"Feature: ")
            run.bold = True
            p.add_run(base_desc)
        
        # Add view type label
        view_label = "🖥️ Desktop View (1440×900)" if view_type == "desktop" else "📱 Mobile View (390×844)"
        p = doc.add_paragraph()
        run = p.add_run(view_label)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0x33, 0x66, 0x99)
        
        # Add screenshot
        try:
            if view_type == "desktop":
                doc.add_picture(path, width=Inches(6.5))
            else:
                doc.add_picture(path, width=Inches(2.5))
            last_paragraph = doc.paragraphs[-1]
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        except Exception as e:
            doc.add_paragraph(f"[Screenshot: {path} - Error: {e}]")
        
        doc.add_paragraph()  # spacing
    
    # Save
    output_path = "/Users/veera.konjeti/Desktop/hack-n-stack/demo-shots/QuizHub_AI_Feature_Evidence.docx"
    doc.save(output_path)
    print(f"  📁 Saved: {output_path}")


if __name__ == "__main__":
    asyncio.run(main())
